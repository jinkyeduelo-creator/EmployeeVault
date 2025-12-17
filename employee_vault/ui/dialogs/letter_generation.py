"""
Letter Generation Dialog
"""

import os
import logging
from datetime import datetime
from typing import Optional, List, Dict, Any
from pathlib import Path

from PySide6.QtWidgets import *
from PySide6.QtCore import *
from PySide6.QtGui import *

from employee_vault.config import *
from employee_vault.database import DB
from employee_vault.validators import *
from employee_vault.utils import *
from employee_vault.models import *
from employee_vault.ui.widgets import *
from employee_vault.ui.modern_ui_helper import show_success_toast, show_error_toast, show_warning_toast, show_info_toast
from employee_vault.ui.widgets import ModernAnimatedButton, PulseButton, SmoothAnimatedDialog
from employee_vault.ui.ios_button_styles import apply_ios_style

class LetterGenerationDialog(SmoothAnimatedDialog):
    """Dialog for generating excuse letters"""
    def __init__(self, db, current_user, employee=None, parent=None):
        # v4.5.0: Use slide-up animation for modern mobile-like experience
        super().__init__(parent, animation_style="slide_up")
        self.db = db
        self.current_user = current_user
        self.setWindowTitle("Generate Excuse Letter")
        self.setMinimumSize(700, 600)

        layout = QVBoxLayout(self)

        # Title
        title = QLabel("<h2>📝 Generate Excuse Letter</h2>")
        layout.addWidget(title)

        # Form
        form = QFormLayout()

        # Employee selection (name only, no ID shown)
        self.employee_combo = NeumorphicGradientComboBox("Select Employee")
        self.employee_combo.setMinimumHeight(70)
        employees = self.db.all_employees()
        for emp in employees:
            self.employee_combo.addItem(emp['name'], emp['emp_id'])
        if employee:
            idx = self.employee_combo.findData(employee['emp_id'])
            if idx >= 0:
                self.employee_combo.combo_box.setCurrentIndex(idx)
        form.addRow("Employee:", self.employee_combo)

        # Store/Branch (editable text field with title case formatting)
        self.store_branch = NeumorphicGradientLineEdit("e.g., SM North Edsa, Robinson's Galleria")
        self.store_branch.setMinimumHeight(70)
        self.store_branch.line_edit.editingFinished.connect(lambda: self._format_title_case(self.store_branch.line_edit))
        form.addRow("Store/Branch:", self.store_branch)

        # Supervisor fields (editable with title case)
        self.supervisor_name = NeumorphicGradientLineEdit("e.g., Juan Dela Cruz")
        self.supervisor_name.setMinimumHeight(70)
        self.supervisor_name.line_edit.editingFinished.connect(lambda: self._format_title_case(self.supervisor_name.line_edit))
        form.addRow("Supervisor Name:", self.supervisor_name)

        self.supervisor_title = NeumorphicGradientLineEdit("e.g., Branch Manager, HR Manager")
        self.supervisor_title.setMinimumHeight(70)
        self.supervisor_title.line_edit.editingFinished.connect(lambda: self._format_title_case(self.supervisor_title.line_edit))
        form.addRow("Supervisor Title:", self.supervisor_title)

        # Company Name (recipient) — editable, with title case
        self.company_name = NeumorphicGradientLineEdit("e.g., INTERNATIONAL TOYWORLD INC.")
        self.company_name.setMinimumHeight(70)
        # Use editingFinished to avoid the spacing/cursor glitch
        self.company_name.line_edit.editingFinished.connect(lambda: self._format_title_case(self.company_name.line_edit))
        form.addRow("Company Name (Recipient):", self.company_name)

        # Date options
        date_group = QGroupBox("Letter Date")
        date_layout = QVBoxLayout()

        self.date_option = NeumorphicGradientComboBox("Today")
        self.date_option.setMinimumHeight(70)
        self.date_option.combo_box.setFocusPolicy(Qt.ClickFocus)
        self.date_option.addItems(["Today", "Yesterday", "Specific Date", "Date Range"])
        self.date_option.combo_box.currentTextChanged.connect(self._on_date_option_changed)
        date_layout.addWidget(self.date_option)

        self.date_edit = DatePicker()
        self.date_edit.setDate(QDate.currentDate())
        self.date_edit.setVisible(False)
        date_layout.addWidget(self.date_edit)

        self.date_range_layout = QHBoxLayout()
        self.start_date = DatePicker()
        self.start_date.setDate(QDate.currentDate())
        self.end_date = DatePicker()
        self.end_date.setDate(QDate.currentDate())
        self.date_range_layout.addWidget(QLabel("From:"))
        self.date_range_layout.addWidget(self.start_date)
        self.date_range_layout.addWidget(QLabel("To:"))
        self.date_range_layout.addWidget(self.end_date)
        date_range_widget = QWidget()
        date_range_widget.setLayout(self.date_range_layout)
        date_range_widget.setVisible(False)
        self.date_range_widget = date_range_widget
        date_layout.addWidget(date_range_widget)

        self._wheel_guard = WheelGuard()
        self.date_option.combo_box.installEventFilter(self._wheel_guard)
        self.date_edit.installEventFilter(self._wheel_guard)
        self.start_date.installEventFilter(self._wheel_guard)
        self.end_date.installEventFilter(self._wheel_guard)

        date_group.setLayout(date_layout)
        form.addRow(date_group)

        # Reason
        self.reason_edit = NeumorphicGradientLineEdit("e.g., personal reasons, family emergency, illness")
        self.reason_edit.setMinimumHeight(70)
        form.addRow("Reason:", self.reason_edit)

        layout.addLayout(form)

        # Buttons - Phase 3: iOS frosted glass
        btn_layout = QHBoxLayout()
        preview_btn = ModernAnimatedButton("👁️ Preview")
        apply_ios_style(preview_btn, 'blue')
        preview_btn.clicked.connect(self.preview_letter)
        generate_btn = ModernAnimatedButton("📄 Generate Letter")
        apply_ios_style(generate_btn, 'green')
        generate_btn.clicked.connect(self.generate_letter)
        cancel_btn = ModernAnimatedButton("Cancel")
        apply_ios_style(cancel_btn, 'gray')
        cancel_btn.clicked.connect(self.reject)

        btn_layout.addStretch()
        btn_layout.addWidget(preview_btn)
        btn_layout.addWidget(generate_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)


    def _format_title_case(self, line_edit):
        """Format text to title case (first letter of each word capitalized)"""
        cursor_pos = line_edit.cursorPosition()
        text = line_edit.text()
        formatted = text.title()  # Use Python's built-in title() method
        if text != formatted:
            line_edit.blockSignals(True)
            line_edit.setText(formatted)
            line_edit.setCursorPosition(min(cursor_pos, len(formatted)))
            line_edit.blockSignals(False)

    def _on_date_option_changed(self, option):
        self.date_edit.setVisible(option == "Specific Date")
        self.date_range_widget.setVisible(option == "Date Range")

    def _get_letter_date_string(self):
        option = self.date_option.combo_box.currentText()
        if option == "Today":
            return datetime.now().strftime("%B %d, %Y")
        elif option == "Yesterday":
            return (datetime.now() - timedelta(days=1)).strftime("%B %d, %Y")
        elif option == "Specific Date":
            return self.date_edit.date().toString("MMMM dd, yyyy")
        elif option == "Date Range":
            start = self.start_date.date().toString("MMMM dd, yyyy")
            end = self.end_date.date().toString("MMMM dd, yyyy")
            return f"{start} to {end}"
        return ""

    def _generate_letter_content(self):
        # Get employee
        emp_id = self.employee_combo.combo_box.currentData()
        employees = self.db.all_employees()
        employee = next((e for e in employees if e['emp_id'] == emp_id), None)
        if not employee:
            return None

        # Get store/branch from text field
        store_branch_text = self.store_branch.line_edit.text().strip()
        if not store_branch_text:
            store_branch_text = "Office"

        # Use defaults for missing fields (don't stop generation for preview)
        company = self.company_name.line_edit.text().strip()
        if not company:
            company = "INTERNATIONAL TOYWORLD INC."

        supervisor_name = self.supervisor_name.line_edit.text().strip()
        if not supervisor_name:
            supervisor_name = "[Supervisor Name]"

        supervisor_title = self.supervisor_title.line_edit.text().strip()
        if not supervisor_title:
            supervisor_title = "[Supervisor Title]"
        
        # v3.9: Get template or use default
        template = self.db.get_letter_template(1)
        if not template or template.strip() == "":
            # Use default template if database template is empty
            template = """[DATE]

The Store Manager
[COMPANY_NAME]
[BRANCH_NAME]

Dear Sir/Madam,

This is to inform you that [EMPLOYEE_NAME] (ID: [EMPLOYEE_ID]) was unable to report for work on [LETTER_DATE] due to [REASON].

We kindly request your understanding regarding this matter.

Thank you for your consideration.

Respectfully yours,

_________________________________
[SUPERVISOR_NAME]
[SUPERVISOR_TITLE]"""

        logging.info(f"Letter template length: {len(template)}, starts with: {template[:50] if template else 'EMPTY'}")
        
        # Replace placeholders
        employee['name'] = employee['name'].split(' (Employee ID')[0]
        letter_date = self._get_letter_date_string()
        content = template.replace("[DATE]", datetime.now().strftime("%B %d, %Y"))

        content = content.replace("[COMPANY_NAME]", company)
        content = content.replace("[BRANCH_NAME]", store_branch_text)
        content = content.replace("[ADDRESS]", "")
        content = content.replace("[EMPLOYEE_NAME]", employee['name'])
        content = content.replace("[EMPLOYEE_ID]", employee['emp_id'])
        content = content.replace("[LETTER_DATE]", letter_date)
        content = re.sub(r'\s*\(Employee ID:.*?\)', '', content)
        content = content.replace(
            "[REASON]",
            self.reason_edit.line_edit.text().strip().title() or "Personal Reasons"
        )
        content = content.replace("[SUPERVISOR_NAME]", supervisor_name)
        content = content.replace("[SUPERVISOR_TITLE]", supervisor_title)
        
        # Clean up
        content = re.sub(r'(?mi)^\s*INTERNATIONAL TOYWORLD INC\.\s*$', '', content).strip()
        content = re.sub(r'(?m)^(The Store Manager)\s*\n\s*\n', r'\1\n', content)
        content = re.sub(r'(?mi)^\s*re:\s*.*\r?\n?', '', content)
        content = re.sub(r'\n{3,}', '\n\n', content).strip()
        
        return content

    def _build_letter_context(self, employee: dict, store_branch: str, company: str,
                              supervisor_name: str, supervisor_title: str,
                              letter_date: str, reason_text: str) -> Dict[str, str]:
        """Construct placeholder context for templated DOC/DOCX exports."""
        safe_name = employee.get('name', '').split(' (Employee ID')[0]
        context = {
            "[DATE]": datetime.now().strftime("%B %d, %Y"),
            "[COMPANY_NAME]": company,
            "[BRANCH_NAME]": store_branch,
            "[STORE_BRANCH]": store_branch,
            "[ADDRESS]": "",
            "[EMPLOYEE_NAME]": safe_name,
            "[EMPLOYEE_ID]": employee.get('emp_id', ''),
            "[LETTER_DATE]": letter_date,
            "[REASON]": reason_text,
            "[SUPERVISOR_NAME]": supervisor_name,
            "[SUPERVISOR_TITLE]": supervisor_title,
        }
        # Friendly aliases for brace-based templates
        alias_context = {}
        for key, value in context.items():
            token = key.strip("[]")
            alias_context[f"{{{{{token}}}}}"] = value  # {{TOKEN}}
            alias_context[token] = value              # TOKEN
        context.update(alias_context)
        return context

    def _find_letter_template_file(self) -> Optional[str]:
        """Locate a DOCX/DOC template in assets/, preferring DOCX with 'cuddly' or 'letter'."""
        assets_dir = resource_path("assets")
        if not os.path.isdir(assets_dir):
            return None

        docx_files = []
        doc_files = []
        for entry in os.listdir(assets_dir):
            lower = entry.lower()
            full_path = os.path.join(assets_dir, entry)
            if not os.path.isfile(full_path):
                continue
            if lower.endswith(".docx"):
                docx_files.append(full_path)
            elif lower.endswith(".doc"):
                doc_files.append(full_path)

        def _sort_key(path: str):
            name = Path(path).name.lower()
            score = 0
            if "cuddly" in name:
                score -= 2
            if "letter" in name or "excuse" in name:
                score -= 1
            return score

        docx_files = sorted(docx_files, key=_sort_key)
        doc_files = sorted(doc_files, key=_sort_key)

        if docx_files:
            return docx_files[0]
        if doc_files:
            # We do not parse .doc directly, but surface it so users can convert to DOCX
            return doc_files[0]
        return None

    def _replace_placeholders_in_doc(self, doc, context: Dict[str, str]):
        """Replace placeholder tokens inside a docx Document (paragraphs + tables)."""
        def replace_in_runs(runs):
            for run in runs:
                for key, value in context.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)

        for paragraph in doc.paragraphs:
            replace_in_runs(paragraph.runs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_runs(cell.paragraphs[0].runs if cell.paragraphs else [])
                    for para in cell.paragraphs:
                        replace_in_runs(para.runs)

    def preview_letter(self):
        """Preview the generated letter with proper formatting"""
        try:
            content = self._generate_letter_content()

            # DEBUG: Print content to console
           # print(f"=== LETTER PREVIEW DEBUG ===")
           # print(f"Content length: {len(content) if content else 0}")
           # print(f"Content type: {type(content)}")
           # print(f"Content preview (first 200 chars): {content[:200] if content else 'NONE'}")
           # print(f"=== END DEBUG ===")

            if not content or content.strip() == "":
                show_error_toast(self, f"The letter content is empty!\n\nDebugging info:\n"
                    f"Content: {repr(content)}\n"
                    f"Please check all fields are filled.")
                return

            # Create preview dialog
            preview_dialog = QDialog(self)
            preview_dialog.setWindowTitle("Letter Preview")
            preview_dialog.setMinimumSize(850, 700)

            layout = QVBoxLayout(preview_dialog)
            layout.setContentsMargins(20, 20, 20, 20)

            # Build a scrollable letter preview using a QLabel inside a QScrollArea.  Using
            # QLabel avoids global QSS size constraints that can shrink QTextEdit to a
            # single line.  The QLabel is wrapped in a QScrollArea to allow scrolling
            # for long letters.
            scroll_area = QScrollArea()
            scroll_area.setWidgetResizable(True)
            scroll_area.setFrameShape(QFrame.NoFrame)

            letter_widget = QWidget()
            letter_layout = QVBoxLayout(letter_widget)
            letter_layout.setContentsMargins(0, 0, 0, 0)

            # Escape HTML-sensitive characters and convert newlines to <br> for the QLabel
            def _escape_html(text: str) -> str:
                return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

            # Build HTML for the preview.  Surround the body of the letter with horizontal rules
            # above and below to approximate the single lines shown in the reference letter.  The
            # <hr> elements provide thin gray lines spanning the width of the preview.
            # Define a horizontal rule style that spans nearly the full width of the letter.
            # By using negative left/right margins equal to the letter padding (25px),
            # the line extends almost to the edges of the white page.  Adjust the
            # margin values if you change the padding on the label.
            hr_style = (
                "border: none; border-top: 1px solid #888; "
                "margin-top: 8px; margin-bottom: 8px; "
                "margin-left: -25px; margin-right: -25px; "
                "width: 100%;"
            )
            body_html = _escape_html(content).replace("\n", "<br>")
            html_text = (
                f"<div><hr style=\"{hr_style}\"></div>"
                f"<p>{body_html}</p>"
                f"<div><hr style=\"{hr_style}\"></div>"
            )

            letter_label = QLabel()
            letter_label.setTextFormat(Qt.RichText)
            letter_label.setText(html_text)
            letter_label.setWordWrap(True)
            letter_label.setAlignment(Qt.AlignTop | Qt.AlignLeft)
            # Style for the letter content.  Keep the same white background and
            # typography as before.
            letter_label.setStyleSheet("""
                background-color: white;
                color: #212121;
                border: 2px solid #ddd;
                border-radius: 8px;
                padding: 25px;
                font-family: 'Times New Roman', 'Georgia', serif;
                font-size: 12pt;
                line-height: 1.8;
            """)
            letter_layout.addWidget(letter_label)

            scroll_area.setWidget(letter_widget)
            layout.addWidget(scroll_area, 1)

            # Close button with professional styling
            btn_layout = QHBoxLayout()
            close_btn = ModernAnimatedButton("Close")
            close_btn.setStyleSheet("""
                QPushButton {
                    background: #2196F3;
                    color: white;
                    padding: 10px 30px;
                    border: none;
                    border-radius: 6px;
                    font-weight: bold;
                    font-size: 14px;
                }
                QPushButton:hover {
                    background: #1976D2;
                }
            """)
            close_btn.clicked.connect(preview_dialog.accept)

            btn_layout.addStretch()
            btn_layout.addWidget(close_btn)
            btn_layout.addStretch()

            # Fix: Add stretch factor of 0 so button layout only takes space it needs
            layout.addLayout(btn_layout, 0)

            preview_dialog.exec()

        except Exception as e:
            show_error_toast(self, f"Failed to generate preview:\n{str(e)}\n\nPlease check your template and employee data.")
            logging.error(f"Letter preview error: {e}", exc_info=True)

    def generate_letter(self):
        # Validate inputs
        if not self.store_branch.line_edit.text().strip():
            show_warning_toast(self, "Please enter Store/Branch name")
            self.store_branch.setFocus()
            return

        if not self.supervisor_name.line_edit.text().strip():
            show_warning_toast(self, "Please enter Supervisor Name")
            self.supervisor_name.setFocus()
            return

        if not self.supervisor_title.line_edit.text().strip():
            show_warning_toast(self, "Please enter Supervisor Title")
            self.supervisor_title.setFocus()
            return

        if not self.reason_edit.line_edit.text().strip():
            show_warning_toast(self, "Please enter a reason for the excuse letter")
            self.reason_edit.setFocus()
            return

        content = self._generate_letter_content()
        if not content:
            show_warning_toast(self, "Failed to generate letter content")
            return

        # Get employee info
        emp_id = self.employee_combo.combo_box.currentData()

        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename_docx = f"excuse_letter_{emp_id}_{timestamp}.docx"
        filepath = os.path.join(LETTERS_DIR, filename_docx)

        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            employees = self.db.all_employees()
            employee = next((e for e in employees if e['emp_id'] == emp_id), {"emp_id": emp_id, "name": ""})
            store_branch_text = self.store_branch.line_edit.text().strip() or "Office"
            company = self.company_name.line_edit.text().strip() or "INTERNATIONAL TOYWORLD INC."
            supervisor_name = self.supervisor_name.line_edit.text().strip() or "[Supervisor Name]"
            supervisor_title = self.supervisor_title.line_edit.text().strip() or "[Supervisor Title]"
            letter_date = self._get_letter_date_string()
            reason_text = self.reason_edit.line_edit.text().strip().title() or "Personal Reasons"

            context = self._build_letter_context(
                employee,
                store_branch_text,
                company,
                supervisor_name,
                supervisor_title,
                letter_date,
                reason_text,
            )

            template_path = self._find_letter_template_file()
            template_ext = Path(template_path).suffix.lower() if template_path else ""
            if template_ext == ".docx":
                doc = Document(template_path)
            else:
                if template_ext == ".doc":
                    logging.warning(
                        f"Template {template_path} is .doc; please convert to .docx for placeholder support. Using generated body instead."
                    )
                doc = Document()

            # Set margins - more space for header and footer
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.5)  # More space for header
                section.bottom_margin = Inches(1.2)  # More space for footer
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                
                # Add header with company logo
                        # ==================== ENHANCED HEADER WITH PROPER FORMATTING ====================
                header = section.header
                header_para = header.paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Try to add header image first
                header_image_path = resource_path("cuddly_header.png")
                if not os.path.exists(header_image_path):
                    for possible_path in ["cuddly_header.png", 
                                          os.path.join(os.path.dirname(__file__), "cuddly_header.png"),
                                          os.path.join(os.getcwd(), "cuddly_header.png")]:
                        if os.path.exists(possible_path):
                            header_image_path = possible_path
                            break
                
                image_loaded_successfully = False
                if os.path.exists(header_image_path):
                    try:
                        # Best solution: User provides 'cuddly_header.png' with logo, text, and lines
                        run = header_para.add_run()
                        run.add_picture(header_image_path, width=Inches(6.2))
                        image_loaded_successfully = True # Image loaded
                    except Exception as img_error:
                        logging.warning(f"Could not add header image: {img_error}")
                        # Fallback to text header
                        pass # Will fall through to the 'if not image_loaded_successfully' block
                
                if not image_loaded_successfully:
                    # Text-based header (fallback) - REPLICATING SCREENSHOT
                    # NOTE: This is a centered approximation. A logo on the left
                    # and text on the right is not feasible with this text fallback.
                    run = header_para.add_run("Cuddly International Corp.\n")
                    run.font.bold = True
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(50, 50, 50) # Dark gray
                    
                    run = header_para.add_run(
                        "650 Jesus Ext. cor. Beata, Pandacan, Manila, PHILIPPINES\n"
                        "Tel: (632) 588-0324 to 25  Fax: (632) 588-0327\n"
                        "email: sales@cuddlyinternational.com\n"
                    )
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(80, 80, 80)
                
                # --- The old single horizontal line was here ---
                
                # --- ADD SINGLE HORIZONTAL LINE AFTER HEADER ---
                # A single thin line is added beneath the header to separate it from the body,
                # matching the style shown in the sample letter.  The line uses underscores
                # and a small font size to appear as a continuous rule.
                line_para = header.add_paragraph()
                line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set some space before and after to separate the line from the header text
                line_para.paragraph_format.space_before = Pt(4)
                line_para.paragraph_format.space_after = Pt(4)
                run_line = line_para.add_run("_" * 150)
                run_line.font.size = Pt(5)
                # Use a medium gray to mimic the subtle line in the reference letter
                run_line.font.color.rgb = RGBColor(128, 128, 128)
                
                # ==================== ENHANCED FOOTER WITH PROPER FORMATTING ====================
                footer = section.footer
                
                # Add horizontal line before footer.  Use a single thin line with the same
                # styling as the header line so the top and bottom lines match.  The number
                # of underscores is increased to better span the page width when rendered.
                footer_line_para = footer.add_paragraph()
                footer_line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                footer_line_para.paragraph_format.space_before = Pt(4)
                footer_line_para.paragraph_format.space_after = Pt(4)
                run = footer_line_para.add_run("_" * 150)
                run.font.size = Pt(5)
                run.font.color.rgb = RGBColor(128, 128, 128)
                
                footer_para = footer.add_paragraph()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Try to add footer image first
                footer_image_path = resource_path("cuddly_footer.png")
                if not os.path.exists(footer_image_path):
                    for possible_path in ["cuddly_footer.png",
                                          os.path.join(os.path.dirname(__file__), "cuddly_footer.png"),
                                          os.path.join(os.getcwd(), "cuddly_footer.png")]:
                        if os.path.exists(possible_path):
                            footer_image_path = possible_path
                            break
                
                if os.path.exists(footer_image_path):
                    try:
                        run = footer_para.add_run()
                        run.add_picture(footer_image_path, width=Inches(5.8))
                    except Exception as img_error:
                        logging.warning(f"Could not add footer image: {img_error}")
                        # Fallback to text footer
                        run = footer_para.add_run("Our Brands:\n")
                        run.font.bold = True
                        run.font.size = Pt(10)
                        
                        run = footer_para.add_run("APRUVA  |  APRUVA FOR MOM & BABY  |  POSH\n")
                        run.font.bold = True
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(192, 0, 0)
                        
                        run = footer_para.add_run(
                            "Quality Products for Your Family\n"
                            "© 2025 Cuddly International Corp. All Rights Reserved."
                        )
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(96, 96, 96)
                else:
                    # Text-based footer (fallback)
                    run = footer_para.add_run("Our Brands:\n")
                    run.font.bold = True
                    run.font.size = Pt(10)
                    
                    run = footer_para.add_run("APRUVA  |  APRUVA FOR MOM & BABY  |  POSH\n")
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(192, 0, 0)
                    
                    run = footer_para.add_run(
                        "Quality Products for Your Family\n"
                        "© 2025 Cuddly International Corp. All Rights Reserved."
                    )
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(96, 96, 96)


            # Apply placeholder replacement when a DOCX template is present
            if template_ext == ".docx":
                self._replace_placeholders_in_doc(doc, context)

            # ==================== ADD LETTER BODY CONTENT ====================
            # Add the actual letter content to the document body
            lines = content.split('\n')
            for line in lines:
                para = doc.add_paragraph(line)
                para_format = para.paragraph_format
                para_format.space_after = Pt(0)
                para_format.space_before = Pt(0)

                # Apply appropriate styling based on content
                if line.strip():  # Non-empty line
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)

                        # Bold for certain lines
                        if any(keyword in line for keyword in ['Dear', 'Respectfully', 'The Store Manager']):
                            run.font.bold = True

            # Save document
            doc.save(filepath)

            # Save to history
            self.db.save_letter_history(
                emp_id, "excuse", letter_date, None,
                reason_text,
                supervisor_name,
                supervisor_title,
                filepath, self.current_user
            )

            reply = QMessageBox.question(
                self, "Success",
                f"Letter generated successfully!\n\nSaved to: {filepath}\n\nOpen the letter now?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )

            if reply == QMessageBox.StandardButton.Yes:
                QDesktopServices.openUrl(QUrl.fromLocalFile(filepath))

            self.accept()

        except Exception as e:
            show_error_toast(self, f"Failed to save letter:\n{str(e)}")
