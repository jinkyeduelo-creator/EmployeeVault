"""
Dashboard Page
Main dashboard with statistics and overview
"""

import os
import json
import shutil
import sqlite3
import logging
from datetime import datetime, timedelta
from typing import List, Dict

from PySide6.QtWidgets import *
from PySide6.QtCore import *
from PySide6.QtGui import *


from employee_vault.ui.widgets import ModernAnimatedButton, PulseButton, titlecase, GlassCard, GlassPanelDark, CountUpLabel, DatePicker
from employee_vault.ui.modern_ui_helper import show_success_toast, show_error_toast, show_warning_toast, show_info_toast
from employee_vault.config import *
from employee_vault.glassmorphism_theme import *
from employee_vault.database import DB

# Import animation manager and particle effects
try:
    from employee_vault.animation_manager import get_animation_manager
    from employee_vault.ui.particle_effects import create_celebration_effect
    ANIMATIONS_AVAILABLE = True
except ImportError:
    ANIMATIONS_AVAILABLE = False
    get_animation_manager = None
    create_celebration_effect = None


class SimpleBarChart(QWidget):
    """Simple bar chart widget using QPainter"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self._data = {}  # {label: value}
        self._colors = ['#4CAF50', '#2196F3', '#FF9800', '#E91E63', '#9C27B0', '#00BCD4', '#FFC107', '#795548']
        self._animated_values = {}
        self._animation_timer = None
        self.setMinimumHeight(200)
        self.setMinimumWidth(300)
        
    def set_data(self, data: dict):
        """Set chart data as {label: value}"""
        self._data = data
        # Start animation
        self._animated_values = {k: 0 for k in data.keys()}
        self._start_animation()
        
    def _start_animation(self):
        """Animate bar heights"""
        if self._animation_timer:
            self._animation_timer.stop()
        self._animation_timer = QTimer(self)
        self._animation_timer.timeout.connect(self._animate_step)
        self._animation_timer.start(16)  # ~60fps
        self._animation_progress = 0
        
    def _animate_step(self):
        """Animation step"""
        self._animation_progress += 0.05
        if self._animation_progress >= 1.0:
            self._animation_progress = 1.0
            if self._animation_timer:
                self._animation_timer.stop()
        
        # Ease out animation
        eased = 1 - (1 - self._animation_progress) ** 3
        for key in self._data:
            self._animated_values[key] = self._data[key] * eased
        self.update()
        
    def paintEvent(self, event):
        """Custom paint for bar chart"""
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        
        if not self._data:
            painter.setPen(QColor(150, 150, 150))
            painter.drawText(self.rect(), Qt.AlignCenter, "No data")
            return
            
        rect = self.rect()
        padding = 40
        chart_rect = rect.adjusted(padding, 20, -padding, -40)
        
        max_val = max(self._data.values()) if self._data.values() else 1
        bar_count = len(self._data)
        if bar_count == 0:
            return
            
        bar_width = min(50, (chart_rect.width() - (bar_count - 1) * 10) / bar_count)
        total_bars_width = bar_count * bar_width + (bar_count - 1) * 10
        start_x = chart_rect.x() + (chart_rect.width() - total_bars_width) / 2
        
        for i, (label, value) in enumerate(self._data.items()):
            animated_val = self._animated_values.get(label, value)
            bar_height = (animated_val / max_val) * (chart_rect.height() - 20) if max_val > 0 else 0
            
            x = start_x + i * (bar_width + 10)
            y = chart_rect.bottom() - bar_height
            
            # Draw bar with gradient
            color = QColor(self._colors[i % len(self._colors)])
            gradient = QLinearGradient(x, y, x, chart_rect.bottom())
            gradient.setColorAt(0, color.lighter(120))
            gradient.setColorAt(1, color)
            
            painter.setBrush(QBrush(gradient))
            painter.setPen(Qt.NoPen)
            painter.drawRoundedRect(int(x), int(y), int(bar_width), int(bar_height), 4, 4)
            
            # Draw value on top
            painter.setPen(QColor(255, 255, 255))
            painter.setFont(QFont("Segoe UI", 9, QFont.Bold))
            val_rect = QRect(int(x), int(y - 20), int(bar_width), 20)
            painter.drawText(val_rect, Qt.AlignCenter, str(int(animated_val)))
            
            # Draw label below
            painter.setFont(QFont("Segoe UI", 8))
            label_rect = QRect(int(x - 10), chart_rect.bottom() + 5, int(bar_width + 20), 30)
            # Truncate long labels
            fm = painter.fontMetrics()
            elided = fm.elidedText(label, Qt.ElideRight, int(bar_width + 20))
            painter.drawText(label_rect, Qt.AlignHCenter | Qt.AlignTop, elided)


class QuickStatsCard(QFrame):
    """Modern stats card with solid gradient and glow effects"""

    clicked = Signal(str)  # Signal emitted when card is clicked, passes card_type

    def __init__(self, title, value, icon, color, card_type="", parent=None):
        super().__init__(parent)

        # Store card type for filtering (total, active, expiring)
        self.card_type = card_type
        self.card_color = color

        # 1. ASSIGN AN ID - crucial so background style applies ONLY to container
        self.setObjectName("statsCard")

        self.setFixedHeight(100)
        self.target_value = value
        self.current_value = 0

        # Make card interactive
        self.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.setMouseTracking(True)

        # Hover lift animation
        self.lift_anim = QPropertyAnimation(self, b"geometry")
        self.lift_anim.setDuration(200)
        self.lift_anim.setEasingCurve(QEasingCurve.Type.OutCubic)

        # Shadow effect for depth - colored glow
        self.shadow = QGraphicsDropShadowEffect(self)
        self.shadow.setBlurRadius(20)
        self.shadow.setColor(QColor(color))
        self.shadow.setOffset(0, 4)
        self.setGraphicsEffect(self.shadow)

        # Get RGB values for gradient
        rgb = self._get_rgb(color)
        
        # SOLID GRADIENT CARDS - Modern with colored glow on hover
        self.setStyleSheet(f"""
            #statsCard {{
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                                           stop:0 rgba({rgb}, 0.85),
                                           stop:0.5 rgba({rgb}, 0.70),
                                           stop:1 rgba({rgb}, 0.55));
                border: 2px solid rgba({rgb}, 0.9);
                border-top: 2px solid rgba(255, 255, 255, 0.3);
                border-left: 2px solid rgba(255, 255, 255, 0.2);
                border-radius: 22px;
            }}
            #statsCard:hover {{
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                                           stop:0 rgba({rgb}, 0.95),
                                           stop:0.5 rgba({rgb}, 0.80),
                                           stop:1 rgba({rgb}, 0.65));
                border: 2px solid rgba({rgb}, 1.0);
                border-top: 2px solid rgba(255, 255, 255, 0.5);
                border-left: 2px solid rgba(255, 255, 255, 0.3);
            }}
            QLabel {{
                color: white;
                background: transparent;
                border: none;
                font-family: 'Segoe UI', sans-serif;
            }}
        """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 15, 20, 15)  
        layout.setSpacing(10)

        top_layout = QHBoxLayout()
        icon_label = QLabel(icon)
        icon_label.setStyleSheet("font-size: 26px;")
        icon_label.setAlignment(Qt.AlignTop)
        top_layout.addWidget(icon_label)
        top_layout.addStretch()

        # v4.5.0: Use CountUpLabel for animated statistics
        if CountUpLabel is not None:
            self.value_label = CountUpLabel()
            self.value_label.setStyleSheet("font-size: 32px; font-weight: 800; background: transparent;")
            self.value_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
            # Set target value with animation
            QTimer.singleShot(100, lambda: self.value_label.set_target_value(value, animate=True))
        else:
            # Fallback to old animation method
            self.value_label = QLabel("0")
            self.value_label.setStyleSheet("font-size: 32px; font-weight: 800;")
            self.value_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
            
            # FEATURE: Animated counter with easing
            self.animation = QPropertyAnimation(self, b"animatedValue")
            self.animation.setDuration(1500)
            self.animation.setStartValue(0)
            self.animation.setEndValue(value)
            self.animation.setEasingCurve(QEasingCurve.OutCubic)
            self.animation.start()
            
        top_layout.addWidget(self.value_label)
        layout.addLayout(top_layout)

        layout.addStretch()  # Push title to bottom

        title_label = QLabel(title)
        title_label.setStyleSheet("font-size: 13px; font-weight: 600; opacity: 0.8;")
        layout.addWidget(title_label)

        # FEATURE: Animated counter with easing
        self.animation = QPropertyAnimation(self, b"animatedValue")
        self.animation.setDuration(1500)
        self.animation.setStartValue(0)
        self.animation.setEndValue(value)
        self.animation.setEasingCurve(QEasingCurve.OutCubic)
        self.animation.start()

    def get_animatedValue(self):
        return self.current_value

    def set_animatedValue(self, value):
        self.current_value = int(value)
        self.value_label.setText(f"{self.current_value:,}")  # Added comma formatting

    animatedValue = Property(int, get_animatedValue, set_animatedValue)

    def _get_rgb(self, hex_color):
        """Convert hex color to RGB string for rgba()"""
        hex_color = hex_color.lstrip('#')
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return f"{r}, {g}, {b}"

    def _lighten(self, color):
        """Lighten a color"""
        colors = {
            "#2196F3": "#42A5F5",
            "#4CAF50": "#66BB6A",
            "#ff9800": "#FFA726",
            "#9C27B0": "#AB47BC"
        }
        return colors.get(color, color)

    def _darken(self, color):
        """Darken a color"""
        colors = {
            "#2196F3": "#1976D2",
            "#4CAF50": "#388E3C",
            "#ff9800": "#F57C00",
            "#9C27B0": "#7B1FA2"
        }
        return colors.get(color, color)

    def mousePressEvent(self, event):
        """Handle card click - emit signal with card type"""
        if event.button() == Qt.MouseButton.LeftButton and self.card_type:
            self.clicked.emit(self.card_type)
        super().mousePressEvent(event)

    def enterEvent(self, event):
        """Lift effect on hover with enhanced colored glow and elevation shadow"""
        # Store original geometry for animation
        self.original_geometry = self.geometry()

        # Lift card up by 6px for more pronounced effect
        new_geometry = self.geometry()
        new_geometry.moveTop(new_geometry.top() - 6)
        self.lift_anim.setStartValue(self.geometry())
        self.lift_anim.setEndValue(new_geometry)
        self.lift_anim.start()

        # Enhanced shadow - increased blur, offset and glow intensity for elevation effect
        self.shadow.setBlurRadius(45)
        self.shadow.setOffset(0, 12)
        # Make glow color more vibrant with higher alpha
        glow_color = QColor(self.card_color)
        glow_color.setAlpha(200)
        self.shadow.setColor(glow_color)

        super().enterEvent(event)

    def leaveEvent(self, event):
        """Reset position and shadow on mouse leave"""
        # Return to original position
        self.lift_anim.setStartValue(self.geometry())
        self.lift_anim.setEndValue(self.original_geometry)
        self.lift_anim.start()

        # Reset shadow with subtle glow
        self.shadow.setBlurRadius(20)
        self.shadow.setOffset(0, 4)
        # Reset glow color to subtle
        glow_color = QColor(self.card_color)
        glow_color.setAlpha(100)
        self.shadow.setColor(glow_color)

        super().leaveEvent(event)

# RecentActivityItem removed - using simple text list instead


class EnhancedDashboardPage(QWidget):
    """Enhanced dashboard with interactive KPI cards and recent activity"""

    # Signal to request navigation to employee page with filter
    navigate_to_employees = Signal(str)  # Passes filter type: "total", "active", "expiring"
    open_employee_detail = Signal(str)  # Passes employee ID
    open_employee_edit = Signal(str)  # Passes employee ID to open edit dialog

    def __init__(self, parent=None):
        super().__init__(parent)
        layout = QVBoxLayout(self)

        title = QLabel("<h1>📊 Dashboard</h1>")
        layout.addWidget(title)

        # Row 1: Quick stats cards
        self.stats_container = QWidget()
        self.stats_layout = QHBoxLayout(self.stats_container)
        self.stats_layout.setSpacing(15)
        layout.addWidget(self.stats_container)

        # Row 2: Data Completeness + Recent Activity side by side
        row2 = QHBoxLayout()
        row2.setSpacing(10)
        
        self.completeness_card = self._create_completeness_card()
        self.completeness_card.setMinimumHeight(160)
        self.completeness_card.setMaximumHeight(220)
        row2.addWidget(self.completeness_card)
        
        self.recent_activity_card = self._create_recent_activity_card()
        self.recent_activity_card.setMinimumHeight(160)
        self.recent_activity_card.setMaximumHeight(220)
        row2.addWidget(self.recent_activity_card)
        
        layout.addLayout(row2)
        
        # Row 3: Department + Agency Distribution side by side
        row3 = QHBoxLayout()
        row3.setSpacing(10)
        
        self.dept_chart_card = self._create_department_chart()
        self.dept_chart_card.setMinimumHeight(160)
        self.dept_chart_card.setMaximumHeight(220)
        row3.addWidget(self.dept_chart_card)
        
        self.agency_chart_card = self._create_agency_chart()
        self.agency_chart_card.setMinimumHeight(160)
        self.agency_chart_card.setMaximumHeight(220)
        row3.addWidget(self.agency_chart_card)
        
        layout.addLayout(row3)

        layout.addStretch(1)

    def _create_completeness_card(self):
        """Create the data completeness card widget - Full iOS frosted glass with expandable view"""
        # Assign unique ID for proper styling
        card = QGroupBox("📋 Data Completeness")
        card.setObjectName("dataCompletenessCard")

        # Full iOS frosted glass styling matching QuickStats cards
        # Diagonal gradient for authentic glassmorphism effect - SOFTER GRADIENTS
        card.setStyleSheet(f"""
            #dataCompletenessCard {{
                font-weight: bold;
                font-size: 15px;
                color: white;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                                           stop:0 rgba(255, 255, 255, 0.15),
                                           stop:1 rgba(255, 152, 0, 0.28));
                border: {GLASS_BORDER_WIDTH}px solid rgba(255, 255, 255, 0.3);
                border-bottom: 1.5px solid rgba(255, 255, 255, 0.15);
                border-right: 1.5px solid rgba(255, 255, 255, 0.15);
                border-radius: {GLASS_CARD_RADIUS}px;
                margin-top: 10px;
                padding: 15px;
                padding-top: 10px;
            }}
            #dataCompletenessCard::title {{
                subcontrol-origin: padding;
                subcontrol-position: top left;
                left: 10px;
                top: 8px;
                padding: 0 5px;
                color: white;
                background: transparent;
            }}
            QLabel {{
                color: white;
                background: transparent;
                border: none;
                font-size: 12px;
                font-family: 'Segoe UI', sans-serif;
                line-height: 1.5;
            }}
        """)

        layout = QVBoxLayout(card)
        layout.setContentsMargins(15, 25, 15, 15)
        layout.setSpacing(5)

        # Skeleton loader for completeness data
        from employee_vault.ui.widgets import SkeletonLine
        self.completeness_skeleton = QWidget()
        skel_layout = QVBoxLayout(self.completeness_skeleton)
        skel_layout.setSpacing(5)
        for _ in range(3):
            skel_layout.addWidget(SkeletonLine(250))
        skel_layout.addStretch()
        self.completeness_skeleton.setVisible(True)
        layout.addWidget(self.completeness_skeleton)

        # Scroll area for completeness content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.NoFrame)
        scroll_area.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll_area.setMaximumHeight(250)

        # Content widget inside scroll area
        scroll_content = QWidget()
        scroll_content.setStyleSheet("background: transparent;")
        scroll_content_layout = QVBoxLayout(scroll_content)
        scroll_content_layout.setContentsMargins(0, 0, 0, 0)

        # Single label for all completeness content
        self.completeness_label = QLabel("")
        self.completeness_label.setWordWrap(True)
        self.completeness_label.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        self.completeness_label.setVisible(False)
        scroll_content_layout.addWidget(self.completeness_label)
        scroll_content_layout.addStretch()

        scroll_area.setWidget(scroll_content)
        layout.addWidget(scroll_area)

        # Store scroll area reference
        self.completeness_scroll = scroll_area

        return card

    def _create_department_chart(self):
        """Create department distribution chart card"""
        card = QGroupBox("🏢 Department Distribution")
        card.setObjectName("deptChartCard")
        
        card.setStyleSheet(f"""
            #deptChartCard {{
                font-weight: bold;
                font-size: 15px;
                color: white;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                                           stop:0 rgba(255, 255, 255, 0.15),
                                           stop:1 rgba(76, 175, 80, 0.28));
                border: {GLASS_BORDER_WIDTH}px solid rgba(255, 255, 255, 0.3);
                border-radius: {GLASS_CARD_RADIUS}px;
                margin-top: 10px;
                padding: 15px;
                padding-top: 10px;
            }}
            #deptChartCard::title {{
                subcontrol-origin: padding;
                subcontrol-position: top left;
                left: 10px;
                top: 8px;
                padding: 0 5px;
                color: white;
                background: transparent;
            }}
        """)
        
        layout = QVBoxLayout(card)
        layout.setContentsMargins(10, 25, 10, 10)
        
        # Scroll area for chart content - invisible scrollbar but still scrollable
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll.setStyleSheet("""
            QScrollArea { background: transparent; border: none; }
            QScrollBar:vertical { width: 0px; background: transparent; }
        """)
        
        # Container for bar chart
        self.dept_chart_widget = QWidget()
        self.dept_chart_widget.setStyleSheet("background: transparent;")
        self.dept_chart_layout = QVBoxLayout(self.dept_chart_widget)
        self.dept_chart_layout.setSpacing(4)
        self.dept_chart_layout.setContentsMargins(0, 0, 0, 0)
        
        scroll.setWidget(self.dept_chart_widget)
        layout.addWidget(scroll)
        
        return card

    def _create_agency_chart(self):
        """Create agency distribution chart card"""
        card = QGroupBox("🏛️ Agency Distribution")
        card.setObjectName("agencyChartCard")
        
        card.setStyleSheet(f"""
            #agencyChartCard {{
                font-weight: bold;
                font-size: 15px;
                color: white;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                                           stop:0 rgba(255, 255, 255, 0.15),
                                           stop:1 rgba(156, 39, 176, 0.28));
                border: {GLASS_BORDER_WIDTH}px solid rgba(255, 255, 255, 0.3);
                border-radius: {GLASS_CARD_RADIUS}px;
                margin-top: 10px;
                padding: 15px;
                padding-top: 10px;
            }}
            #agencyChartCard::title {{
                subcontrol-origin: padding;
                subcontrol-position: top left;
                left: 10px;
                top: 8px;
                padding: 0 5px;
                color: white;
                background: transparent;
            }}
        """)
        
        layout = QVBoxLayout(card)
        layout.setContentsMargins(10, 25, 10, 10)
        
        # Scroll area for chart content - invisible scrollbar but still scrollable
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll.setStyleSheet("""
            QScrollArea { background: transparent; border: none; }
            QScrollBar:vertical { width: 0px; background: transparent; }
        """)
        
        # Container for bar chart
        self.agency_chart_widget = QWidget()
        self.agency_chart_widget.setStyleSheet("background: transparent;")
        self.agency_chart_layout = QVBoxLayout(self.agency_chart_widget)
        self.agency_chart_layout.setSpacing(4)
        self.agency_chart_layout.setContentsMargins(0, 0, 0, 0)
        
        scroll.setWidget(self.agency_chart_widget)
        layout.addWidget(scroll)
        
        return card

    def _normalize_department(self, dept):
        """Normalize department names - combine all stores into 'Store'"""
        if not dept or dept == 'Unknown':
            return 'Unknown'
        # Check if department contains "Store" (case-insensitive)
        if 'store' in dept.lower():
            return 'Store'  # Combine all stores
        return dept

    def _update_charts(self, employees):
        """Update department and agency distribution charts"""
        from collections import Counter

        # Department distribution - normalize store departments and show all
        dept_counts = Counter(
            self._normalize_department(e.get('department', 'Unknown'))
            for e in employees if not e.get('resign_date')
        )
        self._render_bar_chart(self.dept_chart_layout, dept_counts, "#4CAF50", max_items=None)
        
        # Agency distribution - show all agencies
        agency_counts = Counter(e.get('agency', 'Direct Hire') or 'Direct Hire' for e in employees if not e.get('resign_date'))
        self._render_bar_chart(self.agency_chart_layout, agency_counts, "#9C27B0", max_items=None)

    def _render_bar_chart(self, layout, data_counts, color, max_items=None):
        """Render a simple horizontal bar chart"""
        # Clear existing items
        while layout.count():
            item = layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        
        if not data_counts:
            no_data = QLabel("<i>No data</i>")
            no_data.setStyleSheet("color: rgba(255,255,255,0.5);")
            layout.addWidget(no_data)
            return
        
        # Get items sorted by count (show all if max_items is None)
        sorted_items = data_counts.most_common(max_items) if max_items else data_counts.most_common()
        max_value = max(v for _, v in sorted_items) if sorted_items else 1
        
        for name, count in sorted_items:
            row = QWidget()
            row.setStyleSheet("background: transparent;")
            row_layout = QHBoxLayout(row)
            row_layout.setContentsMargins(0, 0, 0, 0)
            row_layout.setSpacing(8)
            
            # Label (truncate if too long)
            display_name = name[:15] + "..." if len(name) > 15 else name
            label = QLabel(display_name)
            label.setFixedWidth(100)
            label.setStyleSheet("color: white; font-size: 11px; background: transparent;")
            label.setToolTip(name)
            row_layout.addWidget(label)
            
            # Bar
            bar_container = QWidget()
            bar_container.setFixedHeight(16)
            bar_container.setStyleSheet("background: transparent;")
            bar_layout = QHBoxLayout(bar_container)
            bar_layout.setContentsMargins(0, 0, 0, 0)
            bar_layout.setSpacing(0)
            
            bar_width = int((count / max_value) * 150)
            bar = QFrame()
            bar.setFixedSize(max(bar_width, 5), 12)
            bar.setStyleSheet(f"""
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 {color}, stop:1 {color}88);
                border-radius: 6px;
            """)
            bar_layout.addWidget(bar)
            bar_layout.addStretch()
            
            row_layout.addWidget(bar_container, 1)
            
            # Count
            count_label = QLabel(str(count))
            count_label.setFixedWidth(30)
            count_label.setStyleSheet("color: white; font-size: 11px; font-weight: bold; background: transparent;")
            count_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
            row_layout.addWidget(count_label)
            
            layout.addWidget(row)
        
        # Add stretch at the end
        layout.addStretch()

    def _create_recent_activity_card(self):
        """Create the recent activity glass card with simple text list"""
        card = QGroupBox("📋 Recent Activity")
        card.setObjectName("recentActivityCard")

        # Glass card styling matching other cards - title inside box
        card.setStyleSheet(f"""
            #recentActivityCard {{
                font-weight: bold;
                font-size: 15px;
                color: white;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                                           stop:0 rgba(255, 255, 255, 0.15),
                                           stop:1 rgba(33, 150, 243, 0.28));
                border: {GLASS_BORDER_WIDTH}px solid rgba(255, 255, 255, 0.3);
                border-bottom: 1.5px solid rgba(255, 255, 255, 0.15);
                border-right: 1.5px solid rgba(255, 255, 255, 0.15);
                border-radius: {GLASS_CARD_RADIUS}px;
                margin-top: 0px;
                padding: 15px 15px 10px 15px;
            }}
            #recentActivityCard::title {{
                subcontrol-origin: padding;
                subcontrol-position: top left;
                left: 12px;
                top: 8px;
                padding: 0 8px;
                color: white;
                background: transparent;
            }}
        """)

        layout = QVBoxLayout(card)
        layout.setContentsMargins(10, 25, 10, 10)
        layout.setSpacing(5)

        # Skeleton loader
        from employee_vault.ui.widgets import SkeletonLine
        self.recent_skeleton = QWidget()
        skel_layout = QVBoxLayout(self.recent_skeleton)
        skel_layout.setSpacing(5)
        for _ in range(3):
            skel_layout.addWidget(SkeletonLine(400))
        self.recent_skeleton.setVisible(True)
        layout.addWidget(self.recent_skeleton)

        # Scroll area for recent activity content - invisible scrollbar but still scrollable
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll.setStyleSheet("""
            QScrollArea { background: transparent; border: none; }
            QScrollBar:vertical { width: 0px; background: transparent; }
        """)
        
        # Container for activity content
        self.recent_content_widget = QWidget()
        self.recent_content_widget.setStyleSheet("background: transparent;")
        recent_content_layout = QVBoxLayout(self.recent_content_widget)
        recent_content_layout.setSpacing(4)
        recent_content_layout.setContentsMargins(0, 0, 0, 0)

        # Simple text list for activity (no clickability)
        self.recent_text_label = QLabel("")
        self.recent_text_label.setWordWrap(True)
        self.recent_text_label.setStyleSheet("""
            QLabel {
                color: white;
                background: transparent;
                border: none;
                font-size: 12px;
                line-height: 1.4;
                padding: 0px;
            }
        """)
        self.recent_text_label.setVisible(False)
        recent_content_layout.addWidget(self.recent_text_label)

        # "No activity" message
        self.no_activity_label = QLabel("<i>No recent activity</i>")
        self.no_activity_label.setStyleSheet("color: rgba(255, 255, 255, 0.5); font-size: 12px; background: transparent;")
        self.no_activity_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.no_activity_label.setVisible(False)
        recent_content_layout.addWidget(self.no_activity_label)
        
        recent_content_layout.addStretch()
        scroll.setWidget(self.recent_content_widget)
        layout.addWidget(scroll)

        return card

    def _check_data_completeness(self, employees):
        """QUICK WIN #8: Check for incomplete employee profiles"""
        missing_photos = 0
        missing_emergency = 0
        missing_ids = 0
        incomplete_employees = []

        for emp in employees:
            issues = []
            emp_id = emp.get('emp_id', '')

            # Check photo
            photo_path = os.path.join(PHOTOS_DIR, f"{emp_id}.png")
            if not os.path.exists(photo_path):
                missing_photos += 1
                issues.append("photo")

            # Check emergency contact
            if not emp.get('emergency_contact_name'):
                missing_emergency += 1
                issues.append("emergency contact")

            # Check government IDs
            if not (emp.get('sss_number') or emp.get('tin_number')):
                missing_ids += 1
                issues.append("government IDs")

            if issues:
                incomplete_employees.append({
                    'name': emp.get('name', 'Unknown'),
                    'emp_id': emp_id,
                    'issues': issues
                })

        return {
            'missing_photos': missing_photos,
            'missing_emergency': missing_emergency,
            'missing_ids': missing_ids,
            'incomplete_employees': incomplete_employees,
            'total_employees': len(employees),
            'complete_count': len(employees) - len(incomplete_employees)
        }

    def _update_completeness_display(self, completeness):
        """Update the completeness card with current data - simple static display"""
        total = completeness['total_employees']
        complete = completeness['complete_count']
        incomplete = len(completeness['incomplete_employees'])

        if incomplete == 0:
            # All profiles complete - show success message
            html = "<p style='color: #4CAF50; font-size: 16px;'><b>✓ All profiles complete!</b></p>"
            html += f"<p>All {total} employee profiles have complete information.</p>"
        else:
            # Some profiles incomplete - show summary and list
            percentage = (complete / total * 100) if total > 0 else 0
            
            html = f"<p style='font-size: 16px;'><b>{complete}/{total} profiles complete ({percentage:.1f}%)</b></p>"
            html += f"<p style='color: #ff9800;'>⚠️ <b>{incomplete} employees need attention</b></p>"
            
            # Quick stats
            issues_list = []
            if completeness['missing_photos']:
                issues_list.append(f"📷 {completeness['missing_photos']} missing photos")
            if completeness['missing_emergency']:
                issues_list.append(f"📞 {completeness['missing_emergency']} missing emergency contacts")
            if completeness['missing_ids']:
                issues_list.append(f"🪪 {completeness['missing_ids']} missing government IDs")
            
            html += "<p style='margin-top: 5px;'>" + " • ".join(issues_list) + "</p>"
            
            # List all incomplete employees
            html += "<p style='font-weight: bold; margin-top: 10px; margin-bottom: 5px;'>Incomplete Profiles:</p>"
            
            for emp in completeness['incomplete_employees']:
                issues_str = ", ".join(emp['issues'])
                html += f"<p style='margin: 3px 0 3px 5px;'>• <b>{emp['name']}</b> ({emp['emp_id']})<br/>"
                html += f"<span style='color: #ffcc80; font-size: 11px; margin-left: 12px;'>Missing: {issues_str}</span></p>"
        
        self.completeness_label.setText(html)
        self.completeness_label.setVisible(True)

    def refresh(self, employees):
        # Clear previous stats
        for i in reversed(range(self.stats_layout.count())):
            self.stats_layout.itemAt(i).widget().setParent(None)

        total = len(employees)
        active = sum(1 for e in employees if not e.get('resign_date'))

        # Count expiring contracts and check for reminders
        expiring = 0
        expiring_employees = []
        missing_photos = 0
        missing_photos_employees = []
        from datetime import datetime, timedelta
        today = datetime.now().date()
        alert_date = today + timedelta(days=30)

        for emp in employees:
            emp_id = emp.get('emp_id', '')
            
            # Check for missing photos
            photo_path = os.path.join(PHOTOS_DIR, f"{emp_id}.png")
            if not os.path.exists(photo_path):
                missing_photos += 1
                missing_photos_employees.append(emp)
            
            if emp.get('contract_expiry'):
                try:
                    expiry_date = datetime.strptime(emp['contract_expiry'], "%m-%d-%Y").date()
                    days = (expiry_date - today).days
                    if 0 <= days <= 30:
                        expiring += 1
                        expiring_employees.append(emp)
                except ValueError:
                    # Invalid date format, skip this employee
                    pass

        # Store missing photos list for click handler
        self._missing_photos_employees = missing_photos_employees

        # 4 stats cards with card types for filtering
        total_card = QuickStatsCard("Total Employees", total, "👥", "#2196F3", "total")
        total_card.clicked.connect(self._handle_card_click)
        total_card.setStyleSheet("opacity: 0;")
        self.stats_layout.addWidget(total_card)

        active_card = QuickStatsCard("Active", active, "✅", "#4CAF50", "active")
        active_card.clicked.connect(self._handle_card_click)
        active_card.setStyleSheet("opacity: 0;")
        self.stats_layout.addWidget(active_card)

        expiring_card = QuickStatsCard("Expiring Soon", expiring, "⚠️", "#ff9800", "expiring")
        expiring_card.clicked.connect(self._handle_card_click)
        expiring_card.setStyleSheet("opacity: 0;")
        self.stats_layout.addWidget(expiring_card)

        # NEW: Missing Photos card
        missing_photos_card = QuickStatsCard("Missing Photos", missing_photos, "📷", "#E91E63", "missing_photos")
        missing_photos_card.clicked.connect(self._handle_card_click)
        missing_photos_card.setStyleSheet("opacity: 0;")
        self.stats_layout.addWidget(missing_photos_card)
        
        # Staggered entrance animation for cards
        self._animate_cards_entrance([total_card, active_card, expiring_card, missing_photos_card])

        # Show contract renewal reminder if there are expiring contracts
        if expiring_employees:
            self._show_contract_reminder(expiring_employees)

        # QUICK WIN #8: Update data completeness widget
        completeness = self._check_data_completeness(employees)
        self._update_completeness_display(completeness)

        # Hide skeleton, show actual content for completeness
        if hasattr(self, 'completeness_skeleton'):
            self.completeness_skeleton.setVisible(False)
        
        # Update department and agency distribution charts
        if hasattr(self, 'dept_chart_layout'):
            self._update_charts(employees)

        # Recent activity - Simple text list (no clickability)
        recent = sorted(employees, key=lambda e: e.get('modified', ''), reverse=True)[:5]
        if recent and recent[0].get('modified'):
            # Build simple text lines
            lines = []
            for e in recent:
                status_icon = "🟢" if not e.get('resign_date') else "🔴"
                name = e.get('name', 'Unknown')
                position = e.get('position', 'N/A')
                modified = e.get('modified', 'N/A')

                lines.append(f"{status_icon} <b>{name}</b> - {position} <span style='color: rgba(255, 255, 255, 0.6);'><i>(Modified: {modified})</i></span>")

            # Set text content
            self.recent_text_label.setText("<br>".join(lines))
            self.recent_text_label.setVisible(True)
            self.no_activity_label.setVisible(False)
        else:
            # Show "no activity" message
            self.recent_text_label.setVisible(False)
            self.no_activity_label.setVisible(True)

        # Hide skeleton
        if hasattr(self, 'recent_skeleton'):
            self.recent_skeleton.setVisible(False)
    
    def _animate_cards_entrance(self, cards):
        """Animate dashboard cards with staggered entrance"""
        for i, card in enumerate(cards):
            # Simple opacity animation with delay
            card.setStyleSheet("opacity: 0;")
            QTimer.singleShot(i * 80, lambda c=card: c.setStyleSheet(""))

    def _show_contract_reminder(self, expiring_employees):
        """Show contract renewal reminder dialog"""
        # Check if reminder was skipped for today
        from datetime import date
        import os
        skip_file = 'contract_reminder_skip.txt'

        if os.path.exists(skip_file):
            try:
                with open(skip_file, 'r') as f:
                    skip_date = date.fromisoformat(f.read().strip())
                    if skip_date == date.today():
                        return  # Skip reminder for today
            except (IOError, OSError, ValueError):
                # If file is corrupted or doesn't exist, show reminder
                pass

        # Show the reminder dialog
        try:
            from employee_vault.ui.dialogs.contract_reminder import ContractReminderDialog
            dialog = ContractReminderDialog(self, expiring_employees)
            dialog.exec()
        except Exception as e:
            logging.error(f"Failed to show contract reminder: {e}")

    def _handle_card_click(self, card_type: str):
        """Handle KPI card click - navigate to employees with filter"""
        # Handle missing_photos specially - show a dialog with the list
        if card_type == "missing_photos":
            self._show_missing_photos_dialog()
            return
        
        # Emit signal to request navigation (parent will handle actual navigation)
        self.navigate_to_employees.emit(card_type)

        # Show feedback toast
        messages = {
            "total": "Showing all employees",
            "active": "Showing active employees only",
            "expiring": "Showing employees with contracts expiring soon"
        }
        if card_type in messages:
            show_info_toast(self, messages[card_type])

    def _show_missing_photos_dialog(self):
        """Show dialog listing all employees without photos - clickable to open edit dialog"""
        from employee_vault.ui.widgets import AnimatedDialogBase, ModernAnimatedButton
        from employee_vault.ui.ios_button_styles import apply_ios_style
        
        dlg = AnimatedDialogBase(self, animation_style="fade")
        dlg.setWindowTitle("📷 Employees Missing Photos")
        dlg.resize(550, 450)
        
        layout = QVBoxLayout(dlg)
        
        # Header
        header = QLabel(f"<h2>📷 Employees Without Photos ({len(self._missing_photos_employees)})</h2>")
        layout.addWidget(header)
        
        if not self._missing_photos_employees:
            success_label = QLabel("<p style='color: #4CAF50; font-size: 16px;'>✅ All employees have photos!</p>")
            layout.addWidget(success_label)
        else:
            # Instructions
            instructions = QLabel("<p>Click on an employee to open their profile and upload a photo:</p>")
            layout.addWidget(instructions)
            
            # Scroll area with clickable employee list
            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setFrameShape(QFrame.NoFrame)
            scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")
            
            scroll_content = QWidget()
            scroll_layout = QVBoxLayout(scroll_content)
            scroll_layout.setContentsMargins(5, 5, 5, 5)
            scroll_layout.setSpacing(8)
            
            def create_click_handler(employee, dialog):
                """Create a click handler for the employee button"""
                def handler():
                    dialog.close()
                    # Emit signal to open employee edit dialog
                    self.open_employee_edit.emit(employee.get('emp_id', ''))
                    show_info_toast(self, f"Opening {employee.get('name', 'employee')} for editing...")
                return handler
            
            for emp in self._missing_photos_employees:
                emp_id = emp.get('emp_id', 'N/A')
                name = emp.get('name', 'Unknown')
                dept = emp.get('department', 'N/A')
                position = emp.get('position', 'N/A')
                
                # Create a clickable button for each employee
                emp_btn = QPushButton(f"📷  {name}  •  {emp_id}  •  {dept}  •  {position}")
                emp_btn.setStyleSheet("""
                    QPushButton {
                        background-color: rgba(255, 255, 255, 0.1);
                        border: 1px solid rgba(255, 255, 255, 0.2);
                        border-radius: 8px;
                        color: white;
                        padding: 12px 15px;
                        text-align: left;
                        font-size: 13px;
                    }
                    QPushButton:hover {
                        background-color: rgba(33, 150, 243, 0.3);
                        border-color: #2196F3;
                    }
                    QPushButton:pressed {
                        background-color: rgba(33, 150, 243, 0.5);
                    }
                """)
                emp_btn.setCursor(Qt.PointingHandCursor)
                emp_btn.clicked.connect(create_click_handler(emp, dlg))
                scroll_layout.addWidget(emp_btn)
            
            scroll_layout.addStretch()
            scroll.setWidget(scroll_content)
            layout.addWidget(scroll, 1)
        
        # Close button
        close_btn = ModernAnimatedButton("Close")
        apply_ios_style(close_btn, 'gray')
        close_btn.clicked.connect(dlg.close)
        layout.addWidget(close_btn)
        
        dlg.exec()

    def _handle_activity_item_click(self, emp_id: str):
        """Handle recent activity item click - open employee detail"""
        # Emit signal to open employee detail dialog
        self.open_employee_detail.emit(emp_id)
        show_info_toast(self, f"Opening employee {emp_id}")

# ============================================================================
# ID CARD GENERATOR
# ============================================================================



# ============================================================================
# ID CARD GENERATOR BACKEND CLASS (Embedded)
# ============================================================================

class IDCardGeneratorBackend:
    """
    Professional ID Card Generator for Employee Vault
    Supports QR codes, barcodes, templates, and direct printing
    Works 100% offline for LAN deployment
    """

    def __init__(self, db_path: str, template_dir: str = "templates"):
        """
        Initialize ID Card Generator

        Args:
            db_path: Path to SQLite database
            template_dir: Directory containing card templates
        """
        self.db_path = db_path
        self.template_dir = template_dir
        self.card_size = (638, 1011)  # CR80 card size at 300 DPI - PORTRAIT (53.98mm x 85.6mm)
        self.secret_key = "EmployeeVault2024"  # For QR code security checksums

        # Ensure template directory exists
        os.makedirs(template_dir, exist_ok=True)

        logging.info(f"IDCardGenerator initialized with database: {db_path}")

    def generate_card(self, employee_id: int, template: str = "standard", side: str = "front"):
        """
        Generate ID card for a single employee

        Args:
            employee_id: Database ID of employee
            template: Template name to use (standard, visitor, contractor, management)
            side: "front" or "back" to generate specific side

        Returns:
            PIL Image object of the generated card (or dict with both sides if side="both")

        Raises:
            ValueError: If employee not found or invalid data
        """
        try:
            from PIL import Image, ImageDraw, ImageFont
        except ImportError:
            raise ImportError("Pillow library required. Install with: pip install Pillow")

        # Fetch employee data from database
        employee = self._get_employee_data(employee_id)

        if side == "both":
            # Generate both sides
            front = self._generate_front(employee, template)
            back = self._generate_back(employee, template)
            logging.info(f"Both card sides generated for employee {employee_id}: {employee.get('name', 'Unknown')}")
            return {'front': front, 'back': back}
        elif side == "back":
            # Generate only back
            card = self._generate_back(employee, template)
            logging.info(f"Back card generated for employee {employee_id}: {employee.get('name', 'Unknown')}")
            return card
        else:
            # Generate only front (default)
            card = self._generate_front(employee, template)
            logging.info(f"Front card generated for employee {employee_id}: {employee.get('name', 'Unknown')}")
            return card

    def _generate_front(self, employee: dict, template: str):
        """Generate front side of ID card"""
        # Load base template with colorful background
        card = self._load_template(template, employee.get('department', 'Default'))

        # Add employee photo
        card = self._add_photo(card, employee.get('photo'))

        # Add text fields
        card = self._add_text_fields(card, employee)

        # Add QR code
        card = self._add_qr_code(card, employee)

        # Barcode removed as per requirement

        # Add company logo
        card = self._add_logo(card)

        return card

    def _generate_back(self, employee: dict, template: str):
        """Generate back side of ID card (PORTRAIT orientation) with colorful background"""
        from PIL import Image, ImageDraw, ImageFont

        # Try to load the colorful background image
        background_path = "id_card_background.png"
        if os.path.exists(background_path):
            try:
                # Load and resize background to fit card size
                back_card = Image.open(background_path).copy()
                back_card = back_card.resize(self.card_size, Image.Resampling.LANCZOS)
                logging.debug(f"Loaded colorful background for back card from {background_path}")
            except Exception as e:
                logging.warning(f"Could not load background image: {e}, using white background")
                back_card = Image.new('RGB', self.card_size, color='white')
        else:
            # Create back card with white background if image not found
            back_card = Image.new('RGB', self.card_size, color='white')
        
        draw = ImageDraw.Draw(back_card)

        # Get department color for header
        header_color = self._get_department_color(employee.get('department', 'Default'))

        # Draw colored header bar (top 80 pixels)
        draw.rectangle([(0, 0), (self.card_size[0], 80)], fill=header_color)

        # Draw border
        draw.rectangle([(0, 0), (self.card_size[0]-1, self.card_size[1]-1)],
                      outline='black', width=3)

        # Load fonts
        try:
            font_large = ImageFont.truetype("C:\\Windows\\Fonts\\arial.ttf", 28)
            font_medium = ImageFont.truetype("C:\\Windows\\Fonts\\arial.ttf", 22)
            font_small = ImageFont.truetype("C:\\Windows\\Fonts\\arial.ttf", 16)
            font_tiny = ImageFont.truetype("C:\\Windows\\Fonts\\arial.ttf", 13)
        except (IOError, OSError):
            # Font file not found, use default
            font_large = ImageFont.load_default()
            font_medium = ImageFont.load_default()
            font_small = ImageFont.load_default()
            font_tiny = ImageFont.load_default()

        # Header text (centered)
        header_text = "EMPLOYEE ID CARD"
        bbox = draw.textbbox((0, 0), header_text, font=font_medium)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, 28), header_text, fill='white', font=font_medium)

        # Important Information Section (CENTERED)
        y_pos = 120
        title_text = "IMPORTANT INFORMATION"
        bbox = draw.textbbox((0, 0), title_text, font=font_medium)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), title_text, fill='black', font=font_medium)
        y_pos += 45

        # Email (centered)
        if employee.get('email'):
            label_text = "Email:"
            bbox = draw.textbbox((0, 0), label_text, font=font_small)
            text_width = bbox[2] - bbox[0]
            x_centered = (638 - text_width) // 2
            draw.text((x_centered, y_pos), label_text, fill=(80, 80, 80), font=font_small)
            y_pos += 25
            
            email_text = employee['email']
            bbox = draw.textbbox((0, 0), email_text, font=font_small)
            text_width = bbox[2] - bbox[0]
            x_centered = (638 - text_width) // 2
            draw.text((x_centered, y_pos), email_text, fill='black', font=font_small)
            y_pos += 40

        # Emergency Contact Information (centered)
        label_text = "Emergency Contact:"
        bbox = draw.textbbox((0, 0), label_text, font=font_small)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), label_text, fill=(80, 80, 80), font=font_small)
        y_pos += 25
        
        contact_name = employee.get('emergency_contact_name', 'N/A')
        bbox = draw.textbbox((0, 0), contact_name, font=font_small)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), contact_name, fill='black', font=font_small)
        y_pos += 25
        
        contact_phone = employee.get('emergency_contact_phone', 'N/A')
        bbox = draw.textbbox((0, 0), contact_phone, font=font_small)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), contact_phone, fill='black', font=font_small)
        y_pos += 40

        # Card Validity (centered)
        label_text = "Card Issued:"
        bbox = draw.textbbox((0, 0), label_text, font=font_small)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), label_text, fill=(80, 80, 80), font=font_small)
        y_pos += 25
        
        issue_text = employee['issue_date']
        bbox = draw.textbbox((0, 0), issue_text, font=font_small)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), issue_text, fill='black', font=font_small)
        y_pos += 35

        label_text = "Valid Until:"
        bbox = draw.textbbox((0, 0), label_text, font=font_small)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), label_text, fill=(80, 80, 80), font=font_small)
        y_pos += 25
        
        expiry_text = employee['expiry_date']
        bbox = draw.textbbox((0, 0), expiry_text, font=font_small)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), expiry_text, fill=(200, 0, 0), font=font_small)
        y_pos += 50

        # Instructions (centered)
        instr_title = "INSTRUCTIONS:"
        bbox = draw.textbbox((0, 0), instr_title, font=font_medium)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), instr_title, fill='black', font=font_medium)
        y_pos += 35

        instructions = [
            "1. This card must be worn at all",
            "   times while on premises",
            "2. Report lost or stolen cards",
            "   immediately",
            "3. This card is non-transferable",
            "4. Return card upon termination",
            "   of employment"
        ]

        for instruction in instructions:
            bbox = draw.textbbox((0, 0), instruction, font=font_tiny)
            text_width = bbox[2] - bbox[0]
            x_centered = (638 - text_width) // 2
            draw.text((x_centered, y_pos), instruction, fill=(60, 60, 60), font=font_tiny)
            y_pos += 23

        # Footer with QR code reference
        y_pos = 760
        qr_text = "Scan QR code on front for verification"
        bbox = draw.textbbox((0, 0), qr_text, font=font_tiny)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), qr_text, fill=(100, 100, 100), font=font_tiny)

        # Signature line (centered)
        y_pos = 820
        sig_start = 150
        sig_end = 488
        draw.line([(sig_start, y_pos), (sig_end, y_pos)], fill='black', width=2)
        
        sig_text = "Authorized Signature"
        bbox = draw.textbbox((0, 0), sig_text, font=font_tiny)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos + 10), sig_text, fill=(100, 100, 100), font=font_tiny)

        # Company info at bottom (centered) - CHANGED TO CUDDLY INTERNATIONAL CORPORATION
        y_pos = 910
        company_text = "Cuddly International Corporation"
        bbox = draw.textbbox((0, 0), company_text, font=font_tiny)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), company_text, fill=(150, 150, 150), font=font_tiny)
        
        y_pos += 20
        id_text = f"ID: {employee['employee_id']}"
        bbox = draw.textbbox((0, 0), id_text, font=font_tiny)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        draw.text((x_centered, y_pos), id_text, fill=(150, 150, 150), font=font_tiny)

        return back_card

    def _get_employee_data(self, employee_id: int) -> dict:
        """
        Fetch employee data from database

        Args:
            employee_id: Database ID

        Returns:
            Dictionary with employee information
        """
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()

        # Query employee data
        # Note: Database uses 'emp_id' and 'name' fields, not 'first_name' and 'last_name'
        # Note: Database does NOT have a 'photo' column - photos are stored separately
        cursor.execute("""
            SELECT emp_id, name, department, position,
                   hire_date, email, emergency_contact_name, emergency_contact_phone
            FROM employees
            WHERE emp_id = ?
        """, (employee_id,))

        row = cursor.fetchone()
        conn.close()

        if not row:
            raise ValueError(f"Employee with ID {employee_id} not found in database")

        # Calculate expiry date (1 year from now)
        issue_date = datetime.now()
        expiry_date = issue_date + timedelta(days=365)

        # Use the emp_id from database
        emp_id_str = str(row['emp_id'])

        # Parse name (database stores full name in single field)
        full_name = row['name'] or 'Unknown'
        name_parts = full_name.split()
        first_name = name_parts[0] if name_parts else 'Unknown'
        last_name = name_parts[-1] if len(name_parts) > 1 else ''

        employee_data = {
            'id': employee_id,
            'employee_id': emp_id_str,
            'name': full_name,
            'first_name': first_name,
            'last_name': last_name,
            'department': row['department'] or 'General',
            'position': row['position'] or 'Employee',
            'hire_date': row['hire_date'] or issue_date.strftime("%Y-%m-%d"),
            'photo': None,  # Photos stored separately in this system
            'email': row['email'],
            'emergency_contact_name': row['emergency_contact_name'] or 'N/A',
            'emergency_contact_phone': row['emergency_contact_phone'] or 'N/A',
            'issue_date': issue_date.strftime("%Y-%m-%d"),
            'expiry_date': expiry_date.strftime("%Y-%m-%d")
        }

        return employee_data

    def _load_template(self, template: str, department: str):
        """
        Load card template based on template name and department

        Args:
            template: Template type (standard, visitor, contractor)
            department: Employee department for color coding

        Returns:
            PIL Image object with base template
        """
        from PIL import Image, ImageDraw

        # Try to load the colorful background image
        background_path = "id_card_background.png"
        if os.path.exists(background_path):
            try:
                # Load and resize background to fit card size
                background = Image.open(background_path).copy()
                background = background.resize(self.card_size, Image.Resampling.LANCZOS)
                card = background
                logging.debug(f"Loaded colorful background from {background_path}")
            except Exception as e:
                logging.warning(f"Could not load background image: {e}, creating plain template")
                card = self._create_plain_template(department)
        else:
            # Check for specific template file
            template_path = os.path.join(self.template_dir, f"{template}.png")
            
            if os.path.exists(template_path):
                # Load existing template
                card = Image.open(template_path).copy()
                logging.debug(f"Loaded template from {template_path}")
            else:
                # Create plain template
                card = self._create_plain_template(department)

        # Add colored header bar for department identification
        draw = ImageDraw.Draw(card)
        header_color = self._get_department_color(department)
        draw.rectangle([(0, 0), (self.card_size[0], 80)], fill=header_color)
        
        # Draw border
        draw.rectangle([(0, 0), (self.card_size[0]-1, self.card_size[1]-1)],
                      outline='black', width=3)

        return card
    
    def _create_plain_template(self, department: str):
        """Create a plain white template with department color header"""
        from PIL import Image, ImageDraw
        
        card = Image.new('RGB', self.card_size, color='white')
        logging.debug(f"Created plain template")
        return card

    def _get_department_color(self, department: str) -> tuple:
        """
        Get RGB color tuple for department

        Args:
            department: Department name

        Returns:
            RGB tuple (r, g, b)
        """
        colors = {
            'IT': (0, 102, 204),        # Blue
            'HR': (0, 170, 0),          # Green
            'Finance': (255, 215, 0),   # Gold
            'Operations': (255, 102, 0), # Orange
            'Management': (102, 0, 153), # Purple
            'Sales': (220, 20, 60),     # Crimson
            'Marketing': (255, 20, 147), # Deep Pink
            'Default': (128, 128, 128)  # Gray
        }

        return colors.get(department, colors['Default'])

    def _add_photo(self, card, photo_blob: bytes):
        """
        Add employee photo to card (PORTRAIT orientation)

        Args:
            card: Card image
            photo_blob: Photo data from database (BLOB)

        Returns:
            Card with photo added
        """
        if not photo_blob:
            logging.debug("No photo available, skipping")
            return card

        try:
            from PIL import Image
            from io import BytesIO

            # Convert BLOB to image
            photo = Image.open(BytesIO(photo_blob))

            # Resize to standard size for PORTRAIT (220x280 pixels - smaller for portrait)
            photo = photo.resize((220, 280), Image.Resampling.LANCZOS)

            # Paste photo onto card - centered horizontally, below header
            # Card width is 638, photo is 220, so center at (638-220)/2 = 209
            card.paste(photo, (209, 100))

            logging.debug("Photo added successfully (portrait)")

        except Exception as e:
            logging.warning(f"Could not add photo: {e}")

        return card

    def _add_text_fields(self, card, employee: dict):
        """
        Add text fields to card (PORTRAIT orientation)
        - Company name above photo
        - Employee information centered below photo
        
        Args:
            card: Card image
            employee: Employee data dictionary

        Returns:
            Card with text fields added
        """
        from PIL import ImageDraw, ImageFont

        draw = ImageDraw.Draw(card)

        # Try to load TrueType fonts with larger sizes for better readability
        try:
            # v4.0: Larger fonts for better readability
            font_xlarge = ImageFont.truetype("arial.ttf", 48)  # Name (was 42)
            font_large = ImageFont.truetype("arial.ttf", 36)   # Company (was 32)
            font_medium = ImageFont.truetype("arial.ttf", 30)  # ID and department (was 26)
            font_small = ImageFont.truetype("arial.ttf", 24)   # Position (was 20)
            font_tiny = ImageFont.truetype("arial.ttf", 18)    # Dates (was 16)
        except (IOError, OSError):
            try:
                # Try Arial on Windows
                # v4.0: Larger fallback fonts
                font_xlarge = ImageFont.truetype("C:\\Windows\\Fonts\\arial.ttf", 48)
                font_large = ImageFont.truetype("C:\\Windows\\Fonts\\arial.ttf", 36)
                font_medium = ImageFont.truetype("C:\\Windows\\Fonts\\arial.ttf", 30)
                font_small = ImageFont.truetype("C:\\Windows\\Fonts\\arial.ttf", 24)
                font_tiny = ImageFont.truetype("C:\\Windows\\Fonts\\arial.ttf", 16)
            except (IOError, OSError):
                # Fall back to default
                font_xlarge = ImageFont.load_default()
                font_large = ImageFont.load_default()
                font_medium = ImageFont.load_default()
                font_small = ImageFont.load_default()
                font_tiny = ImageFont.load_default()

        # Company name in header (white text, centered) - REMOVED from header
        # Will be added above photo instead

        # Add "Cuddly International Corporation" above the photo (centered, bold look)
        company_text = "Cuddly International Corporation"
        bbox = draw.textbbox((0, 0), company_text, font=font_large)
        text_width = bbox[2] - bbox[0]
        x_centered = (638 - text_width) // 2
        # Position just below header, above photo (photo starts at y=100)
        draw.text((x_centered, 88), company_text, fill='black', font=font_large)

        # Employee name (large, centered, below photo)
        y_pos = 400
        name_bbox = draw.textbbox((0, 0), employee['name'], font=font_xlarge)
        name_width = name_bbox[2] - name_bbox[0]
        name_x = max(30, (638 - name_width) // 2)  # Center, but min 30px from edge
        draw.text((name_x, y_pos), employee['name'], fill='black', font=font_xlarge)

        # Employee ID (blue, centered) - NO "ID:" prefix
        y_pos += 55
        id_text = employee['employee_id']  # Just the ID, no label
        id_bbox = draw.textbbox((0, 0), id_text, font=font_medium)
        id_width = id_bbox[2] - id_bbox[0]
        id_x = (638 - id_width) // 2
        draw.text((id_x, y_pos), id_text, fill=(0, 0, 255), font=font_medium)

        # Department (centered, NO label)
        y_pos += 50
        dept_bbox = draw.textbbox((0, 0), employee['department'], font=font_medium)
        dept_width = dept_bbox[2] - dept_bbox[0]
        dept_x = (638 - dept_width) // 2
        draw.text((dept_x, y_pos), employee['department'], fill='black', font=font_medium)

        # Position (centered, NO label)
        y_pos += 45
        pos_bbox = draw.textbbox((0, 0), employee['position'], font=font_small)
        pos_width = pos_bbox[2] - pos_bbox[0]
        pos_x = (638 - pos_width) // 2
        draw.text((pos_x, y_pos), employee['position'], fill=(80, 80, 80), font=font_small)

        # Dates (small, centered, NO labels)
        y_pos += 55
        hired_text = f"Hired: {employee['hire_date']}"
        hired_bbox = draw.textbbox((0, 0), hired_text, font=font_tiny)
        hired_width = hired_bbox[2] - hired_bbox[0]
        hired_x = (638 - hired_width) // 2
        draw.text((hired_x, y_pos), hired_text, fill=(80, 80, 80), font=font_tiny)
        
        y_pos += 25
        issued_text = f"Issued: {employee['issue_date']}"
        issued_bbox = draw.textbbox((0, 0), issued_text, font=font_tiny)
        issued_width = issued_bbox[2] - issued_bbox[0]
        issued_x = (638 - issued_width) // 2
        draw.text((issued_x, y_pos), issued_text, fill=(80, 80, 80), font=font_tiny)
        
        y_pos += 25
        expires_text = f"Expires: {employee['expiry_date']}"
        expires_bbox = draw.textbbox((0, 0), expires_text, font=font_tiny)
        expires_width = expires_bbox[2] - expires_bbox[0]
        expires_x = (638 - expires_width) // 2
        draw.text((expires_x, y_pos), expires_text, fill=(200, 0, 0), font=font_tiny)

        logging.debug("Text fields added (portrait, centered, with company name)")

        return card


        # Department (centered, NO label)
        y_pos += 50
        dept_bbox = draw.textbbox((0, 0), employee['department'], font=font_medium)
        dept_width = dept_bbox[2] - dept_bbox[0]
        dept_x = (638 - dept_width) // 2
        draw.text((dept_x, y_pos), employee['department'], fill='black', font=font_medium)

        # Position (centered, NO label)
        y_pos += 40
        pos_bbox = draw.textbbox((0, 0), employee['position'], font=font_small)
        pos_width = pos_bbox[2] - pos_bbox[0]
        pos_x = (638 - pos_width) // 2
        draw.text((pos_x, y_pos), employee['position'], fill=(80, 80, 80), font=font_small)

        # Dates (small, centered, NO labels)
        y_pos += 50
        hired_text = f"Hired: {employee['hire_date']}"
        hired_bbox = draw.textbbox((0, 0), hired_text, font=font_tiny)
        hired_width = hired_bbox[2] - hired_bbox[0]
        hired_x = (638 - hired_width) // 2
        draw.text((hired_x, y_pos), hired_text, fill=(80, 80, 80), font=font_tiny)
        
        y_pos += 25
        issued_text = f"Issued: {employee['issue_date']}"
        issued_bbox = draw.textbbox((0, 0), issued_text, font=font_tiny)
        issued_width = issued_bbox[2] - issued_bbox[0]
        issued_x = (638 - issued_width) // 2
        draw.text((issued_x, y_pos), issued_text, fill=(80, 80, 80), font=font_tiny)
        
        y_pos += 25
        expires_text = f"Expires: {employee['expiry_date']}"
        expires_bbox = draw.textbbox((0, 0), expires_text, font=font_tiny)
        expires_width = expires_bbox[2] - expires_bbox[0]
        expires_x = (638 - expires_width) // 2
        draw.text((expires_x, y_pos), expires_text, fill=(200, 0, 0), font=font_tiny)

        logging.debug("Text fields added (portrait, no labels)")

        return card

    def _add_qr_code(self, card, employee: dict):
        """
        Add QR code with employee data in readable format

        Args:
            card: Card image
            employee: Employee data dictionary

        Returns:
            Card with QR code added
        """
        try:
            import qrcode
        except ImportError:
            logging.warning("qrcode library not installed, skipping QR code")
            return card

        # Create readable text format for QR code
        qr_text = f"""EMPLOYEE ID CARD
━━━━━━━━━━━━━━━━
Name: {employee['name']}
ID: {employee['employee_id']}
Department: {employee['department']}
Position: {employee['position']}
━━━━━━━━━━━━━━━━
Hired: {employee['hire_date']}
Issued: {employee['issue_date']}
Valid Until: {employee['expiry_date']}
━━━━━━━━━━━━━━━━
Cuddly International Corporation
Verified ✓"""

        # Generate QR code
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_M,
            box_size=8,
            border=1
        )
        qr.add_data(qr_text)
        qr.make(fit=True)

        qr_img = qr.make_image(fill_color="black", back_color="white")
        qr_img = qr_img.resize((150, 150))

        # Paste QR code onto card (bottom center for portrait)
        # Card width is 638, QR is 150, so center at (638-150)/2 = 244
        card.paste(qr_img, (244, 840))

        logging.debug("QR code added with readable text format (portrait)")

        return card

    def _add_barcode(self, card, employee_id: str):
        """
        Add barcode of employee ID

        Args:
            card: Card image
            employee_id: Employee ID string

        Returns:
            Card with barcode added
        """
        try:
            import barcode
            from barcode.writer import ImageWriter
            from io import BytesIO
            from PIL import Image
        except ImportError:
            logging.warning("python-barcode library not installed, skipping barcode")
            return card

        try:
            # Generate Code128 barcode (supports alphanumeric)
            CODE128 = barcode.get_barcode_class('code128')

            # Remove "EMP" prefix for cleaner barcode
            barcode_data = employee_id.replace("EMP", "")

            # Generate barcode
            barcode_img = CODE128(barcode_data, writer=ImageWriter())

            # Render to BytesIO
            buffer = BytesIO()
            barcode_img.write(buffer, options={'write_text': False})
            buffer.seek(0)

            # Convert to PIL Image
            barcode_pil = Image.open(buffer)

            # Resize to fit card
            barcode_pil = barcode_pil.resize((450, 80))

            # Paste barcode onto card (bottom left)
            card.paste(barcode_pil, (50, 540))

            logging.debug("Barcode added")

        except Exception as e:
            logging.warning(f"Could not add barcode: {e}")

        return card

    def _add_logo(self, card):
        """
        Add company logo to card (Apruva logo in top left corner)

        Args:
            card: Card image

        Returns:
            Card with logo added
        """
        from PIL import Image

        # Try multiple logo paths
        logo_paths = [
            "company_logo.png",  # Same directory as app
            os.path.join(self.template_dir, "logo.png"),
            os.path.join(self.template_dir, "company_logo.png")
        ]
        
        logo_path = None
        for path in logo_paths:
            if os.path.exists(path):
                logo_path = path
                break

        if not logo_path:
            logging.debug("Logo not found, skipping")
            return card

        try:
            logo = Image.open(logo_path)

            # Resize logo to fit in header (smaller, proportional)
            # Original logo is wider, so we'll make it fit nicely in the corner
            logo_width = 120
            logo_height = int(logo_width * logo.size[1] / logo.size[0])  # Maintain aspect ratio
            if logo_height > 70:  # Don't let it get too tall
                logo_height = 70
                logo_width = int(logo_height * logo.size[0] / logo.size[1])
            
            logo = logo.resize((logo_width, logo_height), Image.Resampling.LANCZOS)

            # Paste logo in top left corner of header with some padding
            # Position: 15px from left, centered vertically in 80px header
            logo_y = (80 - logo_height) // 2
            if logo.mode == 'RGBA':
                card.paste(logo, (15, logo_y), logo)  # Use alpha channel as mask
            else:
                card.paste(logo, (15, logo_y))

            logging.debug(f"Logo added from {logo_path}")

        except Exception as e:
            logging.warning(f"Could not add logo: {e}")

        return card

    def save_card(self, card, output_path: str):
        """
        Save card image to file

        Args:
            card: PIL Image object
            output_path: File path to save (PNG format)
        """
        try:
            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

            # Save as high-quality PNG
            card.save(output_path, 'PNG', dpi=(300, 300), optimize=True)

            file_size = os.path.getsize(output_path) / 1024  # KB
            logging.info(f"Card saved to {output_path} ({file_size:.1f} KB)")

            return output_path

        except Exception as e:
            logging.error(f"Error saving card: {e}")
            raise

    def print_card(self, card, printer_name: str = None, scale_factor: float = 1.0):
        """
        Print card to Windows printer with size control

        Args:
            card: PIL Image object
            printer_name: Name of printer (None = default printer)
            scale_factor: Print size multiplier (1.0 = actual size, 0 = full page, 1.5 = quarter page, 2.5 = half page)

        Note:
            Requires pywin32 library on Windows
        """
        try:
            import win32print
            import win32ui
            from PIL import ImageWin
        except ImportError:
            logging.error("pywin32 library required for printing. Install with: pip install pywin32")
            raise ImportError("pywin32 required for Windows printing")

        try:
            # Get default printer if not specified
            if not printer_name:
                printer_name = win32print.GetDefaultPrinter()

            logging.info(f"Printing to: {printer_name} (scale: {scale_factor})")

            # Create device context
            hdc = win32ui.CreateDC()
            hdc.CreatePrinterDC(printer_name)

            # Start print job
            hdc.StartDoc("Employee ID Card")
            hdc.StartPage()

            # Get printer capabilities
            printable_width = hdc.GetDeviceCaps(110)  # HORZRES
            printable_height = hdc.GetDeviceCaps(111)  # VERTRES
            dpi_x = hdc.GetDeviceCaps(88)  # LOGPIXELSX
            dpi_y = hdc.GetDeviceCaps(90)  # LOGPIXELSY

            # Calculate scaling
            card_width, card_height = card.size
            
            if scale_factor == 0:
                # Full page auto-scale
                scale = min(printable_width / card_width, printable_height / card_height) * 0.9
            else:
                # Fixed scale based on actual card size
                # Card is at 300 DPI, printer uses its DPI
                scale = (dpi_x / 300.0) * scale_factor

            scaled_width = int(card_width * scale)
            scaled_height = int(card_height * scale)

            # Center position
            x = (printable_width - scaled_width) // 2
            y = (printable_height - scaled_height) // 2

            # Print the image
            dib = ImageWin.Dib(card)
            dib.draw(hdc.GetHandleOutput(), (x, y, x + scaled_width, y + scaled_height))

            # End print job
            hdc.EndPage()
            hdc.EndDoc()
            hdc.DeleteDC()

            logging.info(f"Card printed successfully to {printer_name} at scale {scale_factor}")

        except Exception as e:
            logging.error(f"Printing error: {e}")
            raise

    def batch_generate(self, employee_ids: list = None, department: str = None,
                      template: str = "standard", output_dir: str = "cards", side: str = "front") -> dict:
        """
        Generate multiple ID cards at once

        Args:
            employee_ids: List of employee database IDs (None = all employees)
            department: Filter by department name (None = no filter)
            template: Template to use for all cards
            output_dir: Directory to save generated cards
            side: "front", "back", or "both" - which side(s) to generate

        Returns:
            Dictionary with 'success', 'errors', and 'total' keys
        """
        os.makedirs(output_dir, exist_ok=True)

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        # Build query based on filters
        if employee_ids:
            placeholders = ','.join('?' * len(employee_ids))
            query = f"SELECT emp_id FROM employees WHERE emp_id IN ({placeholders})"
            cursor.execute(query, employee_ids)
        elif department:
            cursor.execute("SELECT emp_id FROM employees WHERE department = ?", (department,))
        else:
            cursor.execute("SELECT emp_id FROM employees")

        employee_list = [row[0] for row in cursor.fetchall()]
        conn.close()

        logging.info(f"Batch generating {len(employee_list)} cards ({side} side(s))...")

        generated_files = []
        errors = []

        for emp_id in employee_list:
            try:
                # Generate card with specified side(s)
                result = self.generate_card(emp_id, template=template, side=side)

                if side == "both":
                    # Save both sides
                    front_path = os.path.join(output_dir, f"card_emp_{emp_id}_front.png")
                    back_path = os.path.join(output_dir, f"card_emp_{emp_id}_back.png")

                    self.save_card(result['front'], front_path)
                    self.save_card(result['back'], back_path)

                    generated_files.append(front_path)
                    generated_files.append(back_path)
                else:
                    # Save single side
                    side_suffix = f"_{side}" if side != "front" else ""
                    output_path = os.path.join(output_dir, f"card_emp_{emp_id}{side_suffix}.png")
                    self.save_card(result, output_path)
                    generated_files.append(output_path)

            except Exception as e:
                error_msg = f"Employee {emp_id}: {str(e)}"
                errors.append(error_msg)
                logging.error(error_msg)

        # Summary
        logging.info(f"Batch generation complete: {len(generated_files)} success, {len(errors)} errors")

        if errors:
            logging.warning(f"Errors encountered: {errors}")

        return {
            'success': generated_files,
            'errors': errors,
            'total': len(employee_list)
        }

    def verify_qr_code(self, qr_data_str: str) -> dict:
        """
        Verify scanned QR code checksum and expiry

        Args:
            qr_data_str: JSON string from scanned QR code

        Returns:
            Employee data dictionary if valid

        Raises:
            ValueError: If QR code is invalid or expired
        """
        try:
            qr_data = json.loads(qr_data_str)
        except json.JSONDecodeError:
            raise ValueError("Invalid QR code format")

        # Extract checksum
        provided_checksum = qr_data.pop('checksum', None)

        if not provided_checksum:
            raise ValueError("QR code missing security checksum")

        # Recalculate checksum
        data_str = json.dumps(qr_data, sort_keys=True)
        expected_checksum = hashlib.sha256((data_str + self.secret_key).encode()).hexdigest()[:8]

        # Verify checksum
        if provided_checksum != expected_checksum:
            raise ValueError("Invalid QR code - security checksum mismatch (tampered or fake card)")

        # Check expiry date
        try:
            expiry_date = datetime.strptime(qr_data['valid_until'], "%Y-%m-%d")
            if expiry_date < datetime.now():
                raise ValueError(f"ID card expired on {qr_data['valid_until']}")
        except KeyError:
            raise ValueError("QR code missing expiry date")

        logging.info(f"QR code verified successfully for {qr_data.get('name', 'Unknown')}")

        return qr_data

    def get_card_statistics(self) -> dict:
        """
        Get statistics about ID card generation

        Returns:
            Dictionary with statistics
        """
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute("SELECT COUNT(*) FROM employees")
        total_employees = cursor.fetchone()[0]

        cursor.execute("SELECT COUNT(DISTINCT department) FROM employees WHERE department IS NOT NULL")
        total_departments = cursor.fetchone()[0]

        conn.close()

        templates_available = 0
        if os.path.exists(self.template_dir):
            templates_available = len([f for f in os.listdir(self.template_dir) if f.endswith('.png')])

        return {
            'total_employees': total_employees,
            'total_departments': total_departments,
            'templates_available': templates_available
        }


# --- (This goes between EnhancedDashboardPage and MainWindow) ---
# --- PASTE THIS NEW CLASS, REPLACING THE OLD IDCardGeneratorBackenderator ---

class IDCardGeneratorBackenderator(QDialog):
    """
    Professional ID Card Generator with QR codes, barcodes, and batch generation
    """
    def __init__(self, parent=None, db=None):
        super().__init__(parent)
        self.db = db
        self.setWindowTitle("🆔 Professional ID Card Generator")
        self.resize(900, 700)

        # Initialize the professional card generator
        if not ID_CARD_GEN_AVAILABLE:
            show_error_toast(self, "Professional ID card generator module not found!")
            self.close()
            return

        try:
            db_path = getattr(self.db, 'db_path', DB_FILE)
            self.card_generator = IDCardGeneratorBackend(db_path=db_path, template_dir="templates")
        except Exception as e:
            show_error_toast(self, f"Failed to initialize card generator: {e}")
            self.close()
            return

        # Main layout
        layout = QVBoxLayout(self)

        # Title (fixed at top)
        title = QLabel("<h2>🆔 Professional ID Card Generator</h2>")
        title.setStyleSheet("background-color: #4a9eff; color: white; padding: 10px;")
        layout.addWidget(title)

        # Create scroll area for main content
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.NoFrame)

        # Widget to hold all scrollable content
        content_widget = QWidget()
        content_layout = QVBoxLayout(content_widget)

        # Info
        info = QLabel("Generate professional employee ID cards with photos, QR codes, and barcodes.")
        info.setWordWrap(True)
        info.setStyleSheet("padding: 10px; background-color: #f0f0f0; border-radius: 5px;")
        content_layout.addWidget(info)

        # === SINGLE CARD GENERATION ===
        single_group = QGroupBox("🎴 Generate Single Card")
        single_layout = QVBoxLayout()

        # Employee selection
        emp_label = QLabel("<b>Select Employee:</b>")
        emp_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9);")
        single_layout.addWidget(emp_label)

        self.employee_combo = NeumorphicGradientComboBox("— Select an Active Employee —")
        self.employee_combo.setMinimumHeight(70)
        self.employees = {e['emp_id']: e for e in self.db.all_employees() if not e.get('resign_date')}
        self.employee_combo.addItem("— Select an Active Employee —", None)
        for emp_id, emp in sorted(self.employees.items(), key=lambda item: item[1]['name']):
            self.employee_combo.addItem(f"{emp['name']} (ID: {emp_id})", emp_id)
        single_layout.addWidget(self.employee_combo)

        # Template removed - using default only
        
        # Side selection (Front/Back/Both)
        side_row = QHBoxLayout()
        side_row.addWidget(QLabel("<b>Card Side:</b>"))

        self.side_group = QButtonGroup()
        self.front_radio = QRadioButton("Front Only")
        self.back_radio = QRadioButton("Back Only")
        self.both_radio = QRadioButton("Both Sides")

        self.front_radio.setChecked(True)
        self.side_group.addButton(self.front_radio)
        self.side_group.addButton(self.back_radio)
        self.side_group.addButton(self.both_radio)

        side_row.addWidget(self.front_radio)
        side_row.addWidget(self.back_radio)
        side_row.addWidget(self.both_radio)
        side_row.addStretch()
        single_layout.addLayout(side_row)

        # Buttons
        single_btn_row = QHBoxLayout()
        self.generate_btn = ModernAnimatedButton("📄 Generate & Preview")
        self.generate_btn.clicked.connect(self._generate_single_card)
        single_btn_row.addWidget(self.generate_btn)

        self.save_btn = PulseButton("💾 Save to File")
        save_btn.start_pulse()
        self.save_btn.clicked.connect(self._save_single_card)
        self.save_btn.setEnabled(False)
        single_btn_row.addWidget(self.save_btn)

        self.print_btn = ModernAnimatedButton("🖨️ Print Card")
        self.print_btn.clicked.connect(self._print_single_card)
        self.print_btn.setEnabled(False)
        single_btn_row.addWidget(self.print_btn)
        single_layout.addLayout(single_btn_row)

        # Preview
        self.preview_label = QLabel("<i>Card preview will appear here</i>")
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setMinimumHeight(200)
        self.preview_label.setStyleSheet("background-color: #f5f5f5; border: 2px dashed #ccc; border-radius: 5px;")
        single_layout.addWidget(self.preview_label)

        single_group.setLayout(single_layout)
        content_layout.addWidget(single_group)

        # === BATCH GENERATION ===
        batch_group = QGroupBox("📦 Batch Card Generation")
        batch_layout = QVBoxLayout()

        # Batch options
        self.batch_all_radio = QRadioButton("All Active Employees")
        self.batch_all_radio.setChecked(True)
        batch_layout.addWidget(self.batch_all_radio)

        self.batch_dept_radio = QRadioButton("Specific Department:")
        batch_layout.addWidget(self.batch_dept_radio)

        self.dept_combo = NeumorphicGradientComboBox("Select Department")
        self.dept_combo.setMinimumHeight(70)
        departments = set(emp.get('department') for emp in self.employees.values() if emp.get('department'))
        self.dept_combo.addItems(sorted(departments))
        self.dept_combo.setEnabled(False)
        batch_layout.addWidget(self.dept_combo)

        self.batch_dept_radio.toggled.connect(lambda checked: self.dept_combo.setEnabled(checked))

        # Batch side selection
        batch_side_row = QHBoxLayout()
        batch_side_row.addWidget(QLabel("<b>Card Side:</b>"))

        self.batch_side_group = QButtonGroup()
        self.batch_front_radio = QRadioButton("Front Only")
        self.batch_back_radio = QRadioButton("Back Only")
        self.batch_both_radio = QRadioButton("Both Sides")

        self.batch_front_radio.setChecked(True)
        self.batch_side_group.addButton(self.batch_front_radio)
        self.batch_side_group.addButton(self.batch_back_radio)
        self.batch_side_group.addButton(self.batch_both_radio)

        batch_side_row.addWidget(self.batch_front_radio)
        batch_side_row.addWidget(self.batch_back_radio)
        batch_side_row.addWidget(self.batch_both_radio)
        batch_side_row.addStretch()
        batch_layout.addLayout(batch_side_row)

        # Output directory
        output_row = QHBoxLayout()
        output_row.addWidget(QLabel("<b>Output Directory:</b>"))
        self.output_path_label = QLabel("cards_output")
        # Fix: Add color and border-radius to make text visible and match smooth edges
        self.output_path_label.setStyleSheet("border: 1px solid #ccc; padding: 5px; background-color: white; color: #000000; border-radius: 6px;")
        output_row.addWidget(self.output_path_label, 1)

        browse_btn = ModernAnimatedButton("📁 Browse...")
        browse_btn.clicked.connect(self._browse_output_dir)
        output_row.addWidget(browse_btn)
        batch_layout.addLayout(output_row)

        # Batch button
        batch_btn_row = QHBoxLayout()
        self.batch_generate_btn = ModernAnimatedButton("🚀 Generate All Cards")
        self.batch_generate_btn.clicked.connect(self._batch_generate)
        batch_btn_row.addWidget(self.batch_generate_btn)
        
        self.batch_print_btn = ModernAnimatedButton("🖨️ Print Multiple Cards")
        self.batch_print_btn.clicked.connect(self._batch_print_dialog)
        self.batch_print_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold;")
        batch_btn_row.addWidget(self.batch_print_btn)
        
        batch_layout.addLayout(batch_btn_row)

        batch_group.setLayout(batch_layout)
        content_layout.addWidget(batch_group)

        # Add stretch at bottom of scrollable content
        content_layout.addStretch()

        # Set the content widget to scroll area
        scroll_area.setWidget(content_widget)
        layout.addWidget(scroll_area)

        # Close button (fixed at bottom, not scrollable)
        close_layout = QHBoxLayout()
        close_layout.addStretch()
        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(self.close)
        close_btn.setFixedHeight(35)
        close_layout.addWidget(close_btn)
        layout.addLayout(close_layout)

        self.current_card = None

    def _generate_single_card(self):
        emp_id = self.employee_combo.currentData()
        if not emp_id:
            show_warning_toast(self, "Please select an employee first.")
            return

        # v3.1: Use default template only
        template = "standard"

        # Determine which side to generate
        if self.front_radio.isChecked():
            side = "front"
        elif self.back_radio.isChecked():
            side = "back"
        else:  # both_radio is checked
            side = "both"

        try:
            result = self.card_generator.generate_card(emp_id, template=template, side=side)

            if side == "both":
                # Handle both sides - create a composite image showing both
                from PIL import Image

                # Get the front and back images
                front_img = result['front']
                back_img = result['back']

                # Create a composite image with both cards side by side
                # Add 20 pixels spacing between cards
                spacing = 20
                composite_width = front_img.width + back_img.width + spacing
                composite_height = max(front_img.height, back_img.height)

                composite = Image.new('RGB', (composite_width, composite_height), color='white')
                composite.paste(front_img, (0, 0))
                composite.paste(back_img, (front_img.width + spacing, 0))

                # Save composite for preview
                temp_path = "temp_card_preview_both.png"
                composite.save(temp_path)

                pixmap = QPixmap(temp_path)
                scaled_pixmap = pixmap.scaled(
                    self.preview_label.width() - 20,
                    self.preview_label.height() - 20,
                    Qt.KeepAspectRatio,
                    Qt.SmoothTransformation
                )
                self.preview_label.setPixmap(scaled_pixmap)

                self.current_card = result  # Store both sides
                self.current_emp_id = emp_id
                self.save_btn.setEnabled(True)
                self.print_btn.setEnabled(True)

                show_success_toast(self, "Both sides generated successfully!\nPreview shows front (left) and back (right).")
            else:
                # Handle single side
                temp_path = "temp_card_preview.png"
                self.card_generator.save_card(result, temp_path)

                pixmap = QPixmap(temp_path)
                scaled_pixmap = pixmap.scaled(
                    self.preview_label.width() - 20,
                    self.preview_label.height() - 20,
                    Qt.KeepAspectRatio,
                    Qt.SmoothTransformation
                )
                self.preview_label.setPixmap(scaled_pixmap)

                self.save_btn.setEnabled(True)
                self.print_btn.setEnabled(True)
                self.current_card = result
                self.current_emp_id = emp_id

                side_text = "Front" if side == "front" else "Back"
                show_success_toast(self, f"ID card ({side_text} side) generated successfully!")

        except Exception as e:
            show_error_toast(self, f"Failed to generate card:\n{str(e)}")

    def _save_single_card(self):
        if not self.current_card:
            show_warning_toast(self, "Please generate a card first.")
            return

        # Check if we have both sides
        if isinstance(self.current_card, dict) and 'front' in self.current_card and 'back' in self.current_card:
            # Save both sides
            default_name = f"id_card_emp_{self.current_emp_id}_front.png"
            file_path, _ = QFileDialog.getSaveFileName(self, "Save ID Card (Front)", default_name, "PNG Images (*.png)")

            if file_path:
                try:
                    # Save front
                    self.card_generator.save_card(self.current_card['front'], file_path)

                    # Generate back filename
                    back_path = file_path.replace("_front.png", "_back.png")
                    if back_path == file_path:  # If didn't contain "_front"
                        back_path = file_path.replace(".png", "_back.png")

                    # Save back
                    self.card_generator.save_card(self.current_card['back'], back_path)

                    show_success_toast(self, f"Both sides saved successfully:\n\nFront: {file_path}\nBack: {back_path}")
                except Exception as e:
                    show_error_toast(self, f"Failed to save card:\n{str(e)}")
        else:
            # Save single side
            default_name = f"id_card_emp_{self.current_emp_id}.png"
            file_path, _ = QFileDialog.getSaveFileName(self, "Save ID Card", default_name, "PNG Images (*.png)")

            if file_path:
                try:
                    self.card_generator.save_card(self.current_card, file_path)
                    show_success_toast(self, f"ID card saved successfully to:\n{file_path}")
                except Exception as e:
                    show_error_toast(self, f"Failed to save card:\n{str(e)}")

    def _print_single_card(self):
        if not self.current_card:
            show_warning_toast(self, "Please generate a card first.")
            return

        # Show print size selection dialog
        # v4.4.1: Animated dialog for print size selection
        from employee_vault.ui.widgets import QuickAnimatedDialog
        size_dialog = QuickAnimatedDialog(self, animation_style="fade")
        size_dialog.setWindowTitle("Select Print Size")
        size_dialog.setMinimumWidth(400)
        
        layout = QVBoxLayout(size_dialog)
        
        layout.addWidget(QLabel("<h3>📐 Select ID Card Print Size</h3>"))
        layout.addWidget(QLabel("Choose the size for printing your ID card:"))
        
        # Size options with radio buttons
        size_group = QButtonGroup(size_dialog)
        
        # Actual size (recommended)
        actual_radio = QRadioButton("✓ Actual Size (Recommended)")
        actual_radio.setChecked(True)
        actual_info = QLabel("   → Prints at CR80 standard size (53.98mm × 85.6mm / 2.125\" × 3.375\")")
        actual_info.setStyleSheet("color: gray; font-size: 12px; margin-left: 20px;")
        
        # Half page
        half_radio = QRadioButton("Half Page Size")
        half_info = QLabel("   → Prints at ~105mm × 148mm (fits 2 per A4 page)")
        half_info.setStyleSheet("color: gray; font-size: 12px; margin-left: 20px;")
        
        # Quarter page  
        quarter_radio = QRadioButton("Quarter Page Size")
        quarter_info = QLabel("   → Prints at ~75mm × 105mm (fits 4 per A4 page)")
        quarter_info.setStyleSheet("color: gray; font-size: 12px; margin-left: 20px;")
        
        # Full page
        full_radio = QRadioButton("Full Page Size")
        full_info = QLabel("   → Prints across entire page (not recommended)")
        full_info.setStyleSheet("color: gray; font-size: 12px; margin-left: 20px;")
        
        size_group.addButton(actual_radio, 1)
        size_group.addButton(quarter_radio, 2)
        size_group.addButton(half_radio, 3)
        size_group.addButton(full_radio, 4)
        
        layout.addWidget(actual_radio)
        layout.addWidget(actual_info)
        layout.addWidget(quarter_radio)
        layout.addWidget(quarter_info)
        layout.addWidget(half_radio)
        layout.addWidget(half_info)
        layout.addWidget(full_radio)
        layout.addWidget(full_info)
        
        layout.addWidget(QLabel("\n💡 Tip: Choose 'Actual Size' for standard ID cards that fit in card holders."))
        
        # Buttons
        btn_layout = QHBoxLayout()
        print_btn = ModernAnimatedButton("🖨️ Print")
        cancel_btn = ModernAnimatedButton("Cancel")
        
        print_btn.clicked.connect(size_dialog.accept)
        cancel_btn.clicked.connect(size_dialog.reject)
        
        btn_layout.addStretch()
        btn_layout.addWidget(print_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)
        
        if size_dialog.exec() != QDialog.Accepted:
            return
        
        # Get selected size
        selected_id = size_group.checkedId()
        if selected_id == 1:
            scale_factor = 1.0  # Actual size
        elif selected_id == 2:
            scale_factor = 1.5  # Quarter page
        elif selected_id == 3:
            scale_factor = 2.5  # Half page
        else:
            scale_factor = 0  # Full page (auto-scale)

        # Check if we have both sides
        if isinstance(self.current_card, dict) and 'front' in self.current_card and 'back' in self.current_card:
            # v3.1: Print both sides on single paper
            try:
                from PIL import Image
                
                # Create side-by-side layout
                front_img = self.current_card['front']
                back_img = self.current_card['back']
                
                # Add spacing between cards
                spacing = 40
                combined_width = front_img.width + back_img.width + spacing
                combined_height = max(front_img.height, back_img.height)
                
                # Create white background
                combined = Image.new('RGB', (combined_width, combined_height), color='white')
                combined.paste(front_img, (0, 0))
                combined.paste(back_img, (front_img.width + spacing, 0))
                
                # Print the combined image
                self.card_generator.print_card(combined, scale_factor=scale_factor)
                show_success_toast(self, "Both card sides sent to printer on single page!")
            except ImportError:
                show_error_toast(self, "Windows printing requires pywin32 library.")
            except Exception as e:
                show_error_toast(self, f"Failed to print card:\n{str(e)}")
        else:
            # Print single side
            try:
                self.card_generator.print_card(self.current_card, scale_factor=scale_factor)
                show_success_toast(self, "Card sent to printer successfully!")
            except ImportError:
                show_error_toast(self, "Windows printing requires pywin32 library.")
            except Exception as e:
                show_error_toast(self, f"Failed to print card:\n{str(e)}")

    def _browse_output_dir(self):
        directory = QFileDialog.getExistingDirectory(self, "Select Output Directory", self.output_path_label.text())
        if directory:
            self.output_path_label.setText(directory)

    def _batch_generate(self):
        output_dir = self.output_path_label.text()

        if self.batch_all_radio.isChecked():
            target = "all employees"
            department = None
        else:
            department = self.dept_combo.currentText()
            target = f"department: {department}"

        # Determine which side to generate for batch
        if self.batch_front_radio.isChecked():
            side = "front"
            side_text = "Front Only"
        elif self.batch_back_radio.isChecked():
            side = "back"
            side_text = "Back Only"
        else:  # batch_both_radio is checked
            side = "both"
            side_text = "Both Sides"

        reply = QMessageBox.question(self, "Batch Generation",
                                     f"Generate ID cards ({side_text}) for {target}?\n\nOutput: {output_dir}",
                                     QMessageBox.Yes | QMessageBox.No)

        if reply == QMessageBox.No:
            return

        try:
            from PySide6.QtWidgets import QProgressDialog
            progress = QProgressDialog("Generating ID cards...", "Cancel", 0, 100, self)
            progress.setWindowModality(Qt.WindowModal)
            progress.setValue(0)
            progress.show()

            if department:
                result = self.card_generator.batch_generate(department=department, output_dir=output_dir, side=side)
            else:
                result = self.card_generator.batch_generate(output_dir=output_dir, side=side)

            progress.setValue(100)
            progress.close()

            success_count = len(result['success'])
            error_count = len(result['errors'])

            result_msg = (f"Batch generation complete!\n\n"
                         f"Total: {result['total']}\n"
                         f"Successful: {success_count}\n"
                         f"Errors: {error_count}\n\n"
                         f"Cards saved to:\n{output_dir}")

            if error_count > 0:
                result_msg += f"\n\nErrors:\n" + "\n".join(result['errors'][:5])

            # v3.2: Add "Go to Folder" option
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Batch Complete")
            msg_box.setText(result_msg)
            msg_box.setIcon(QMessageBox.Information)
            
            ok_btn = msg_box.addButton("OK", QMessageBox.AcceptRole)
            folder_btn = msg_box.addButton("📁 Go to Folder", QMessageBox.ActionRole)
            
            msg_box.exec()
            
            # Check which button was clicked
            if msg_box.clickedButton() == folder_btn:
                # Open folder in file explorer
                import subprocess
                import platform
                
                abs_output_dir = os.path.abspath(output_dir)
                if platform.system() == "Windows":
                    subprocess.Popen(f'explorer "{abs_output_dir}"')
                elif platform.system() == "Darwin":  # macOS
                    subprocess.Popen(["open", abs_output_dir])
                else:  # Linux
                    subprocess.Popen(["xdg-open", abs_output_dir])

        except Exception as e:
            show_error_toast(self, f"Batch generation failed:\n{str(e)}")

    def _batch_print_dialog(self):
        """Show dialog for batch printing multiple ID cards on one paper"""
        dialog = BatchPrintDialog(self, self.db, self.card_generator)
        dialog.exec()

# --- (The class MainWindow(QMainWindow) should start after this) ---

# ============================================================================
# BATCH PRINT DIALOG - Multiple Cards Per Page
# ============================================================================

class BatchPrintDialog(QDialog):
    """Dialog for printing multiple ID cards on one sheet"""
    def __init__(self, parent, db, card_generator):
        super().__init__(parent)
        self.db = db
        self.card_generator = card_generator
        self.setWindowTitle("🖨️ Batch Print ID Cards")
        self.resize(700, 700)  # v3.2: Increased height
        
        # v3.2: Make entire dialog scrollable
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.NoFrame)
        
        content_widget = QWidget()
        layout = QVBoxLayout(content_widget)
        
        # Title
        title = QLabel("<h2>🖨️ Print Multiple ID Cards</h2>")
        title.setStyleSheet("background-color: #4CAF50; color: white; padding: 10px;")
        layout.addWidget(title)
        
        # Info
        info = QLabel("Print multiple ID cards on a single sheet of paper. Select employees, paper size, and layout.")
        info.setWordWrap(True)
        info.setStyleSheet("padding: 10px; background-color: #f0f0f0; border-radius: 5px; margin: 10px;")
        layout.addWidget(info)
        
        # Employee Selection
        emp_group = QGroupBox("👥 Select Employees to Print")
        emp_layout = QVBoxLayout()
        
        select_row = QHBoxLayout()
        select_all_btn = ModernAnimatedButton("Select All")
        select_all_btn.clicked.connect(self._select_all)
        select_none_btn = ModernAnimatedButton("Select None")
        select_none_btn.clicked.connect(self._select_none)
        select_row.addWidget(select_all_btn)
        select_row.addWidget(select_none_btn)
        select_row.addStretch()
        emp_layout.addLayout(select_row)
        
        # v3.2: Employee list with VISIBLE checkboxes
        self.employee_list = QListWidget()
        self.employee_list.setSelectionMode(QAbstractItemView.MultiSelection)
        # Fix checkbox visibility - make checked items visible
        self.employee_list.setStyleSheet("""
            QListWidget::item {
                padding: 5px;
            }
            QListWidget::item:checked {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
            }
        """)
        
        employees = self.db.all_employees()
        active_employees = [e for e in employees if not e.get('resign_date')]
        
        for emp in sorted(active_employees, key=lambda x: x.get('name', '')):
            item = QListWidgetItem(f"{emp.get('name')} (ID: {emp.get('emp_id')})")
            item.setData(Qt.UserRole, emp.get('emp_id'))
            item.setCheckState(Qt.Unchecked)
            item.setFlags(item.flags() | Qt.ItemIsUserCheckable)  # v3.2: Ensure checkable
            self.employee_list.addItem(item)
        
        emp_layout.addWidget(self.employee_list)
        emp_group.setLayout(emp_layout)
        layout.addWidget(emp_group)
        
        # Print Settings
        settings_group = QGroupBox("⚙️ Print Settings")
        settings_layout = QVBoxLayout()
        
        # Paper size selection
        paper_label = QLabel("<b>Paper Size:</b>")
        paper_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        settings_layout.addWidget(paper_label)
        self.paper_combo = NeumorphicGradientComboBox("Select Paper Size")
        self.paper_combo.setMinimumHeight(70)
        self.paper_combo.addItems(["A4 (210 x 297 mm)", "Letter (8.5 x 11 in)", "Legal (8.5 x 14 in)"])
        self.paper_combo.currentIndexChanged.connect(self._update_layout_preview)
        settings_layout.addWidget(self.paper_combo)

        # Cards per page
        cards_label = QLabel("<b>Cards Per Page:</b>")
        cards_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        settings_layout.addWidget(cards_label)
        self.cards_combo = NeumorphicGradientComboBox("Select Cards Per Page")
        self.cards_combo.setMinimumHeight(70)
        self.cards_combo.addItems(["1 card per page", "2 cards per page", "4 cards per page (2x2)",
                                   "6 cards per page (2x3)", "8 cards per page (2x4)",
                                   "9 cards per page (3x3)", "10 cards per page (2x5)"])
        self.cards_combo.setCurrentIndex(2)  # Default to 4 cards
        self.cards_combo.currentIndexChanged.connect(self._update_layout_preview)
        settings_layout.addWidget(self.cards_combo)

        # Card side selection
        side_label = QLabel("<b>Print Side:</b>")
        side_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        settings_layout.addWidget(side_label)
        self.side_combo = NeumorphicGradientComboBox("Select Print Side")
        self.side_combo.setMinimumHeight(70)
        self.side_combo.addItems(["Front Only", "Back Only", "Both Sides (Front then Back)"])
        settings_layout.addWidget(self.side_combo)
        
        # Layout preview label
        self.preview_label = QLabel()
        self.preview_label.setStyleSheet("padding: 10px; background-color: white; border: 1px solid #ccc; margin: 10px;")
        self.preview_label.setAlignment(Qt.AlignCenter)
        self._update_layout_preview()
        settings_layout.addWidget(self.preview_label)
        
        settings_group.setLayout(settings_layout)
        layout.addWidget(settings_group)
        
        # Set scroll area
        scroll_area.setWidget(content_widget)
        
        # Main layout with scroll area
        main_layout = QVBoxLayout(self)
        main_layout.addWidget(scroll_area)
        
        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        
        preview_btn = ModernAnimatedButton("👁️ Preview")
        preview_btn.clicked.connect(self._preview_print)
        preview_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 8px 20px;")
        btn_layout.addWidget(preview_btn)
        
        print_btn = ModernAnimatedButton("🖨️ Print")
        print_btn.clicked.connect(self._do_print)
        print_btn.setStyleSheet("background-color: #4CAF50; color: white; padding: 8px 20px; font-weight: bold;")
        btn_layout.addWidget(print_btn)
        
        cancel_btn = ModernAnimatedButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        btn_layout.addWidget(cancel_btn)
        
        layout.addLayout(btn_layout)
    
    def _select_all(self):
        for i in range(self.employee_list.count()):
            item = self.employee_list.item(i)
            item.setCheckState(Qt.Checked)
    
    def _select_none(self):
        for i in range(self.employee_list.count()):
            item = self.employee_list.item(i)
            item.setCheckState(Qt.Unchecked)
    
    def _update_layout_preview(self):
        paper = self.paper_combo.currentText()
        cards = self.cards_combo.currentText()
        preview_text = f"<b>Layout Preview:</b><br>{paper}<br>{cards}"
        self.preview_label.setText(preview_text)
    
    def _get_selected_employees(self):
        selected = []
        for i in range(self.employee_list.count()):
            item = self.employee_list.item(i)
            if item.checkState() == Qt.Checked:
                emp_id = item.data(Qt.UserRole)
                selected.append(emp_id)
        return selected
    
    def _preview_print(self):
        selected = self._get_selected_employees()
        if not selected:
            show_warning_toast(self, "Please select at least one employee.")
            return
        
        # Generate preview
        try:
            from PIL import Image
            
            # Get settings
            cards_per_page = int(self.cards_combo.currentIndex())
            cards_map = {0: 1, 1: 2, 2: 4, 3: 6, 4: 8, 5: 9, 6: 10}
            cards_per_page = cards_map[cards_per_page]
            
            side_idx = self.side_combo.currentIndex()  # 0=front, 1=back, 2=both
            
            # v3.2: Handle "Both Sides" preview
            if side_idx == 2:  # Both sides
                # Create front sheet
                front_sheet = self._create_print_sheet(selected, cards_per_page, "front")
                back_sheet = self._create_print_sheet(selected, cards_per_page, "back")
                
                # Save both sheets
                import tempfile
                temp_front = tempfile.NamedTemporaryFile(delete=False, suffix='_front.png')
                temp_back = tempfile.NamedTemporaryFile(delete=False, suffix='_back.png')
                front_sheet.save(temp_front.name)
                back_sheet.save(temp_back.name)
                
                # Show preview dialog with tabs for front and back
                # v4.4.1: Animated dialog for print preview (both sides)
                from employee_vault.ui.widgets import AnimatedDialogBase
                preview_dialog = AnimatedDialogBase(self, animation_style="fade")
                preview_dialog.setWindowTitle("Print Preview - Both Sides")
                preview_dialog.resize(900, 700)
                
                preview_layout = QVBoxLayout(preview_dialog)
                
                # Create tab widget
                tabs = QTabWidget()
                
                # Front tab
                front_tab = QWidget()
                front_layout = QVBoxLayout(front_tab)
                front_label = QLabel()
                front_pixmap = QPixmap(temp_front.name)
                front_scaled = front_pixmap.scaled(850, 600, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                front_label.setPixmap(front_scaled)
                front_label.setAlignment(Qt.AlignCenter)
                front_scroll = QScrollArea()
                front_scroll.setWidget(front_label)
                front_scroll.setWidgetResizable(True)
                front_layout.addWidget(front_scroll)
                tabs.addTab(front_tab, "🎴 Front Side")
                
                # Back tab
                back_tab = QWidget()
                back_layout = QVBoxLayout(back_tab)
                back_label = QLabel()
                back_pixmap = QPixmap(temp_back.name)
                back_scaled = back_pixmap.scaled(850, 600, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                back_label.setPixmap(back_scaled)
                back_label.setAlignment(Qt.AlignCenter)
                back_scroll = QScrollArea()
                back_scroll.setWidget(back_label)
                back_scroll.setWidgetResizable(True)
                back_layout.addWidget(back_scroll)
                tabs.addTab(back_tab, "🎴 Back Side")
                
                preview_layout.addWidget(tabs)
                
                info_label = QLabel(f"Preview: {len(selected)} employee(s), {cards_per_page} cards/page, BOTH sides")
                info_label.setStyleSheet("padding: 10px; background-color: #f0f0f0; font-weight: bold;")
                preview_layout.addWidget(info_label)
                
                close_btn = ModernAnimatedButton("Close Preview")
                close_btn.clicked.connect(preview_dialog.close)
                preview_layout.addWidget(close_btn)
                
                preview_dialog.exec()
                
                # Cleanup
                os.unlink(temp_front.name)
                os.unlink(temp_back.name)
                
            else:  # Front or Back only
                side = "front" if side_idx == 0 else "back"
                sheet = self._create_print_sheet(selected, cards_per_page, side)
                
                # Save preview
                import tempfile
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                sheet.save(temp_file.name)
                
                # Show preview
                # v4.4.1: Animated dialog for print preview (single side)
                from employee_vault.ui.widgets import AnimatedDialogBase
                preview_dialog = AnimatedDialogBase(self, animation_style="fade")
                preview_dialog.setWindowTitle(f"Print Preview - {side.title()}")
                preview_dialog.resize(800, 600)
                
                preview_layout = QVBoxLayout(preview_dialog)
                
                label = QLabel()
                pixmap = QPixmap(temp_file.name)
                scaled_pixmap = pixmap.scaled(750, 550, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                label.setPixmap(scaled_pixmap)
                label.setAlignment(Qt.AlignCenter)
                
                scroll = QScrollArea()
                scroll.setWidget(label)
                scroll.setWidgetResizable(True)
                preview_layout.addWidget(scroll)
                
                info_label = QLabel(f"Preview: {len(selected)} employee(s), {cards_per_page} cards/page, {side.title()} only")
                info_label.setStyleSheet("padding: 10px; background-color: #f0f0f0;")
                preview_layout.addWidget(info_label)
                
                close_btn = ModernAnimatedButton("Close Preview")
                close_btn.clicked.connect(preview_dialog.close)
                preview_layout.addWidget(close_btn)
                
                preview_dialog.exec()
                
                # Cleanup
                os.unlink(temp_file.name)
            
        except Exception as e:
            show_error_toast(self, f"Failed to create preview:\n{str(e)}\n\nPlease check that all selected employees have photos.")
    
    def _create_print_sheet(self, emp_ids, cards_per_page, side):
        """Create a sheet with multiple ID cards arranged in a grid"""
        from PIL import Image
        
        # Paper sizes in pixels at 300 DPI
        paper_sizes = {
            0: (2480, 3508),  # A4: 210 x 297 mm
            1: (2550, 3300),  # Letter: 8.5 x 11 inches
            2: (2550, 4200),  # Legal: 8.5 x 14 inches
        }
        
        paper_idx = self.paper_combo.currentIndex()
        sheet_width, sheet_height = paper_sizes[paper_idx]
        
        # Create white background
        sheet = Image.new('RGB', (sheet_width, sheet_height), 'white')
        
        # Calculate grid layout
        grid_layouts = {
            1: (1, 1), 2: (1, 2), 4: (2, 2), 
            6: (2, 3), 8: (2, 4), 9: (3, 3), 10: (2, 5)
        }
        
        cols, rows = grid_layouts.get(cards_per_page, (2, 2))
        
        # Calculate card size and spacing
        card_width = 638  # CR80 standard at 300 DPI
        card_height = 1011
        
        # Scale cards to fit on sheet with margins
        margin = 100  # pixels
        available_width = (sheet_width - margin * 2) / cols
        available_height = (sheet_height - margin * 2) / rows
        
        scale = min(available_width / card_width, available_height / card_height) * 0.9
        
        scaled_width = int(card_width * scale)
        scaled_height = int(card_height * scale)
        
        # Calculate spacing
        x_spacing = (sheet_width - margin * 2) // cols
        y_spacing = (sheet_height - margin * 2) // rows
        
        # Place cards
        for idx, emp_id in enumerate(emp_ids[:cards_per_page]):
            if idx >= cards_per_page:
                break
            
            # Generate card
            try:
                card = self.card_generator.generate_card(emp_id, side=side)
                card_resized = card.resize((scaled_width, scaled_height), Image.Resampling.LANCZOS)
                
                # Calculate position (centered in each grid cell)
                row = idx // cols
                col = idx % cols
                
                x = margin + col * x_spacing + (x_spacing - scaled_width) // 2
                y = margin + row * y_spacing + (y_spacing - scaled_height) // 2
                
                sheet.paste(card_resized, (x, y))
            except Exception as e:
                logging.error(f"Failed to add card for employee {emp_id}: {e}")
        
        return sheet
    
    def _do_print(self):
        selected = self._get_selected_employees()
        if not selected:
            show_warning_toast(self, "Please select at least one employee.")
            return
        
        try:
            from PIL import Image
            
            # Get settings
            cards_per_page = int(self.cards_combo.currentIndex())
            cards_map = {0: 1, 1: 2, 2: 4, 3: 6, 4: 8, 5: 9, 6: 10}
            cards_per_page = cards_map[cards_per_page]
            
            side_idx = self.side_combo.currentIndex()
            
            # Progress dialog
            progress = QProgressDialog("Generating print sheets...", "Cancel", 0, 100, self)
            progress.setWindowModality(Qt.WindowModal)
            progress.setMinimumDuration(0)
            progress.setValue(0)
            
            sheets = []
            
            # Generate sheets for front
            if side_idx == 0 or side_idx == 2:
                for i in range(0, len(selected), cards_per_page):
                    if progress.wasCanceled():
                        return
                    
                    batch = selected[i:i+cards_per_page]
                    sheet = self._create_print_sheet(batch, cards_per_page, "front")
                    sheets.append(("Front", sheet))
                    
                    progress.setValue(int((i / len(selected)) * 50))
            
            # Generate sheets for back
            if side_idx == 1 or side_idx == 2:
                for i in range(0, len(selected), cards_per_page):
                    if progress.wasCanceled():
                        return
                    
                    batch = selected[i:i+cards_per_page]
                    sheet = self._create_print_sheet(batch, cards_per_page, "back")
                    sheets.append(("Back", sheet))
                    
                    if side_idx == 2:
                        progress.setValue(50 + int((i / len(selected)) * 50))
                    else:
                        progress.setValue(int((i / len(selected)) * 100))
            
            progress.setValue(100)
            progress.close()
            
            # Show print dialog
            if sheets:
                self._show_print_dialog(sheets)
            
        except Exception as e:
            show_error_toast(self, f"Failed to create print sheets:\n{str(e)}")
    
    def _show_print_dialog(self, sheets):
        """Show native print dialog"""
        from PySide6.QtPrintSupport import QPrinter, QPrintDialog
        from PySide6.QtGui import QPainter
        
        printer = QPrinter(QPrinter.HighResolution)
        printer.setPageOrientation(QPrinter.Portrait)
        
        # Set paper size based on selection
        paper_sizes = {
            0: QPrinter.A4,
            1: QPrinter.Letter,
            2: QPrinter.Legal
        }
        printer.setPageSize(paper_sizes[self.paper_combo.currentIndex()])
        
        dialog = QPrintDialog(printer, self)
        
        if dialog.exec() == QPrintDialog.Accepted:
            painter = QPainter()
            painter.begin(printer)
            
            for idx, (side_name, sheet) in enumerate(sheets):
                if idx > 0:
                    printer.newPage()
                
                # Convert PIL image to QPixmap
                import tempfile
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                sheet.save(temp_file.name)
                
                pixmap = QPixmap(temp_file.name)
                
                # Scale to fit printer page
                target_rect = printer.pageRect(QPrinter.DevicePixel)
                scaled_pixmap = pixmap.scaled(
                    target_rect.width(), 
                    target_rect.height(), 
                    Qt.KeepAspectRatio, 
                    Qt.SmoothTransformation
                )
                
                # Center on page
                x = (target_rect.width() - scaled_pixmap.width()) // 2
                y = (target_rect.height() - scaled_pixmap.height()) // 2
                
                painter.drawPixmap(x, y, scaled_pixmap)
                
                # Cleanup
                os.unlink(temp_file.name)
            
            painter.end()
            
            show_success_toast(self, f"Successfully sent {len(sheets)} page(s) to printer!\n\n"
                f"{len(self._get_selected_employees())} employee cards printed.")
            
            self.accept()

# ============================================================================
# FEATURE A: PRINTING SYSTEM
# ============================================================================

class PrintSystemDialog(QDialog):
    """Professional printing system for employee records"""
    def __init__(self, parent, db, employees):
        super().__init__(parent)
        self.db = db
        self.employees = employees
        self.setWindowTitle("🖨️ Print System")
        self.resize(600, 500)

        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("<h2>🖨️ Print System</h2>"))

        options_group = QGroupBox("What would you like to print?")
        options_layout = QVBoxLayout(options_group)

        print_profile_btn = ModernAnimatedButton("📄 Print Employee Profile")
        print_profile_btn.setToolTip("Print detailed profile of selected employee")
        print_profile_btn.clicked.connect(self.print_employee_profile)
        options_layout.addWidget(print_profile_btn)

        print_list_btn = ModernAnimatedButton("📋 Print Employee List")
        print_list_btn.setToolTip("Print table of all employees")
        print_list_btn.clicked.connect(self.print_employee_list)
        options_layout.addWidget(print_list_btn)

        print_contracts_btn = ModernAnimatedButton("📊 Print Contract Report")
        print_contracts_btn.setToolTip("Print report of expiring contracts")
        print_contracts_btn.clicked.connect(self.print_contract_report)
        options_layout.addWidget(print_contracts_btn)

        layout.addWidget(options_group)

        filter_group = QGroupBox("Filter Options")
        filter_layout = QVBoxLayout(filter_group)

        self.include_resigned = QCheckBox("Include resigned employees")
        filter_layout.addWidget(self.include_resigned)

        self.include_photos = QCheckBox("Include photos in prints")
        self.include_photos.setChecked(True)
        filter_layout.addWidget(self.include_photos)

        layout.addWidget(filter_group)

        button_layout = QHBoxLayout()
        button_layout.addStretch()
        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(self.reject)
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)

    def get_filtered_employees(self):
        if self.include_resigned.isChecked():
            return self.employees
        return [e for e in self.employees if not e.get('resign_date')]

    def print_employee_profile(self):
        employees = self.get_filtered_employees()
        if not employees:
            show_warning_toast(self, "No employees to print.")
            return

        names = [f"{e['emp_id']} - {e['name']}" for e in employees]
        name, ok = QInputDialog.getItem(self, "Select Employee",
                                       "Choose employee to print:", names, 0, False)
        if not ok: return

        emp_id = name.split(" - ")[0]
        employee = next((e for e in employees if e['emp_id'] == emp_id), None)
        if not employee: return

        html = self._generate_profile_html(employee)
        self._print_html(html, f"Employee Profile - {employee['name']}")

    def print_employee_list(self):
        employees = self.get_filtered_employees()
        if not employees:
            show_warning_toast(self, "No employees to print.")
            return
        html = self._generate_list_html(employees)
        self._print_html(html, "Employee List")

    def print_contract_report(self):
        employees = self.get_filtered_employees()
        with_contracts = [e for e in employees if e.get('contract_expiry')]
        if not with_contracts:
            show_warning_toast(self, "No employees with contracts.")
            return
        html = self._generate_contract_html(with_contracts)
        self._print_html(html, "Contract Report")

    def _generate_profile_html(self, emp):
        photo_html = ""
        if self.include_photos.isChecked():
            photo_path = os.path.join(PHOTOS_DIR, f"{emp['emp_id']}.png")
            if os.path.exists(photo_path):
                photo_html = f'<img src="file:///{photo_path}" width="150" style="float:right; margin:10px;"/>'

        status = "Active" if not emp.get('resign_date') else f"Resigned ({emp.get('resign_date')})"

        html = f"""<html><head><style>
body {{ font-family: Arial, sans-serif; padding: 20px; }}
h1 {{ color: #2196F3; border-bottom: 2px solid #2196F3; }}
h2 {{ color: #555; margin-top: 20px; }}
table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
td {{ padding: 8px; border-bottom: 1px solid #ddd; }}
td:first-child {{ font-weight: bold; width: 200px; }}
.status {{ padding: 5px 10px; border-radius: 3px; display: inline-block; }}
.active {{ background: #4CAF50; color: white; }}
.resigned {{ background: #f44336; color: white; }}
</style></head><body>
{photo_html}
<h1>Employee Profile</h1>
<div class="status {'active' if not emp.get('resign_date') else 'resigned'}">{status}</div>
<h2>Personal Information</h2>
<table>
<tr><td>Employee ID:</td><td>{emp['emp_id']}</td></tr>
<tr><td>Name:</td><td>{emp['name']}</td></tr>
<tr><td>Email:</td><td>{emp.get('email', 'N/A')}</td></tr>
<tr><td>Phone:</td><td>{emp.get('phone', 'N/A')}</td></tr>
</table>
<h2>Employment Details</h2>
<table>
<tr><td>Department:</td><td>{emp.get('department', 'N/A')}</td></tr>
<tr><td>Position:</td><td>{emp.get('position', 'N/A')}</td></tr>
<tr><td>Hire Date:</td><td>{emp.get('hire_date', 'N/A')}</td></tr>
<tr><td>Salary:</td><td>₱{float(emp.get('salary', 0)):,.2f}</td></tr>
<tr><td>Agency:</td><td>{emp.get('agency') or 'Direct Hire'}</td></tr>
</table>
<h2>Government IDs</h2>
<table>
<tr><td>SSS Number:</td><td>{emp.get('sss_number', 'N/A')}</td></tr>
<tr><td>TIN:</td><td>{emp.get('tin', 'N/A')}</td></tr>
<tr><td>PhilHealth:</td><td>{emp.get('philhealth', 'N/A')}</td></tr>
<tr><td>Pag-IBIG:</td><td>{emp.get('pagibig', 'N/A')}</td></tr>
</table>
<h2>Contract Information</h2>
<table>
<tr><td>Contract Start:</td><td>{emp.get('contract_start_date', 'N/A')}</td></tr>
<tr><td>Contract Duration:</td><td>{emp.get('contract_months', 'N/A')} months</td></tr>
<tr><td>Contract Expiry:</td><td>{emp.get('contract_expiry', 'N/A')}</td></tr>
</table>
<h2>Emergency Contact</h2>
<table>
<tr><td>Name:</td><td>{emp.get('emergency_contact_name', 'N/A')}</td></tr>
<tr><td>Phone:</td><td>{emp.get('emergency_contact_phone', 'N/A')}</td></tr>
</table>
<p style="margin-top:30px; font-size:10px; color:#999;">
Printed on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</p>
</body></html>"""
        return html

    def _generate_list_html(self, employees):
        rows = ""
        for emp in employees:
            status = "Active" if not emp.get('resign_date') else "Resigned"
            rows += f"""<tr>
<td>{emp['emp_id']}</td><td>{emp['name']}</td>
<td>{emp.get('department', 'N/A')}</td><td>{emp.get('position', 'N/A')}</td>
<td>₱{float(emp.get('salary', 0)):,.2f}</td><td>{status}</td>
</tr>"""

        html = f"""<html><head><style>
body {{ font-family: Arial, sans-serif; padding: 20px; }}
h1 {{ color: #2196F3; }}
table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
th {{ background: #2196F3; color: white; padding: 10px; text-align: left; }}
td {{ padding: 8px; border-bottom: 1px solid #ddd; }}
tr:hover {{ background: #f5f5f5; }}
</style></head><body>
<h1>Employee List</h1>
<p><b>Total Employees:</b> {len(employees)}</p>
<table><thead><tr>
<th>ID</th><th>Name</th><th>Department</th><th>Position</th><th>Salary</th><th>Status</th>
</tr></thead><tbody>{rows}</tbody></table>
<p style="margin-top:30px; font-size:10px; color:#999;">
Printed on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</p>
</body></html>"""
        return html

    def _generate_contract_html(self, employees):
        from datetime import datetime as dt
        expiring_soon, future_expiry = [], []

        for emp in employees:
            try:
                expiry = dt.strptime(emp['contract_expiry'], "%m-%d-%Y")
                days_left = (expiry - dt.now()).days
                if days_left <= 30:
                    expiring_soon.append((emp, days_left))
                else:
                    future_expiry.append((emp, days_left))
            except ValueError:
                # Invalid date format, skip this employee
                pass

        expiring_soon.sort(key=lambda x: x[1])
        future_expiry.sort(key=lambda x: x[1])

        rows = ""
        for emp, days in expiring_soon:
            color = "#f44336" if days <= 0 else "#ff9800"
            status = "EXPIRED" if days <= 0 else f"{days} days left"
            rows += f"""<tr style="background:{color}20;">
<td><b>{emp['emp_id']}</b></td><td><b>{emp['name']}</b></td>
<td>{emp.get('agency', 'Direct')}</td><td>{emp['contract_expiry']}</td>
<td style="color:{color}; font-weight:bold;">{status}</td></tr>"""

        for emp, days in future_expiry:
            rows += f"""<tr>
<td>{emp['emp_id']}</td><td>{emp['name']}</td>
<td>{emp.get('agency', 'Direct')}</td><td>{emp['contract_expiry']}</td>
<td>{days} days</td></tr>"""

        html = f"""<html><head><style>
body {{ font-family: Arial, sans-serif; padding: 20px; }}
h1 {{ color: #2196F3; }}
.warning {{ background: #fff3cd; padding: 15px; border-left: 4px solid #ff9800; margin: 20px 0; }}
table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
th {{ background: #2196F3; color: white; padding: 10px; text-align: left; }}
td {{ padding: 8px; border-bottom: 1px solid #ddd; }}
</style></head><body>
<h1>Contract Expiry Report</h1>
<p><b>Total Contracts:</b> {len(employees)}</p>
<p><b>Expiring Soon (≤30 days):</b> {len(expiring_soon)}</p>
{f'<div class="warning">⚠️ <b>WARNING:</b> {len(expiring_soon)} contract(s) expiring within 30 days!</div>' if expiring_soon else ''}
<table><thead><tr>
<th>ID</th><th>Name</th><th>Agency</th><th>Expiry Date</th><th>Status</th>
</tr></thead><tbody>{rows}</tbody></table>
<p style="margin-top:30px; font-size:10px; color:#999;">
Printed on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</p>
</body></html>"""
        return html

    def _print_html(self, html, title):
        printer = QPrinter(QPrinter.HighResolution)
        printer.setDocName(title)
        dialog = QPrintDialog(printer, self)
        if dialog.exec() == QDialog.Accepted:
            document = QTextDocument()
            document.setHtml(html)
            document.print_(printer)
            show_success_toast(self, "Document sent to printer!")

# ============================================================================
# FEATURE B: BULK OPERATIONS
# ============================================================================

class BulkOperationsDialog(QDialog):
    """Bulk edit multiple employees at once"""
    def __init__(self, parent, db, employees):
        super().__init__(parent)
        self.db = db
        self.employees = employees
        self.setWindowTitle("📦 Bulk Operations")
        self.resize(700, 600)

        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("<h2>📦 Bulk Operations</h2>"))

        selection_group = QGroupBox("Select Employees")
        selection_layout = QVBoxLayout(selection_group)

        self.employee_list = QListWidget()
        self.employee_list.setSelectionMode(QListWidget.MultiSelection)

        for emp in employees:
            status = "" if not emp.get('resign_date') else " [Resigned]"
            item = QListWidgetItem(f"{emp['emp_id']} - {emp['name']} ({emp.get('department', 'N/A')}){status}")
            item.setData(Qt.UserRole, emp)
            self.employee_list.addItem(item)

        selection_layout.addWidget(self.employee_list)

        select_btns = QHBoxLayout()
        select_all_btn = ModernAnimatedButton("Select All")
        select_all_btn.clicked.connect(self.employee_list.selectAll)
        select_btns.addWidget(select_all_btn)

        select_none_btn = ModernAnimatedButton("Clear Selection")
        select_none_btn.clicked.connect(self.employee_list.clearSelection)
        select_btns.addWidget(select_none_btn)

        selection_layout.addLayout(select_btns)
        layout.addWidget(selection_group)

        ops_group = QGroupBox("Choose Operation")
        ops_layout = QVBoxLayout(ops_group)

        bulk_dept_btn = ModernAnimatedButton("🏢 Change Department")
        bulk_dept_btn.clicked.connect(self.bulk_change_department)
        ops_layout.addWidget(bulk_dept_btn)

        bulk_position_btn = ModernAnimatedButton("💼 Change Position")
        bulk_position_btn.clicked.connect(self.bulk_change_position)
        ops_layout.addWidget(bulk_position_btn)

        bulk_agency_btn = ModernAnimatedButton("🏛️ Change Agency")
        bulk_agency_btn.clicked.connect(self.bulk_change_agency)
        ops_layout.addWidget(bulk_agency_btn)

        bulk_archive_btn = ModernAnimatedButton("📦 Archive Selected")
        bulk_archive_btn.clicked.connect(self.bulk_archive)
        ops_layout.addWidget(bulk_archive_btn)

        bulk_export_btn = ModernAnimatedButton("📄 Export Selected")
        bulk_export_btn.clicked.connect(self.bulk_export)
        ops_layout.addWidget(bulk_export_btn)

        layout.addWidget(ops_group)

        self.status_label = QLabel("Select employees and choose an operation")
        layout.addWidget(self.status_label)

        button_layout = QHBoxLayout()
        button_layout.addStretch()
        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(self.reject)
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)

    def get_selected_employees(self):
        return [item.data(Qt.UserRole) for item in self.employee_list.selectedItems()]

    def bulk_change_department(self):
        selected = self.get_selected_employees()
        if not selected:
            show_warning_toast(self, "Please select employees first.")
            return

        dept, ok = QInputDialog.getItem(self, "Change Department",
                                       "New department:",
                                       ["Office", "Warehouse", "Store"], 0, False)
        if not ok: return

        reply = QMessageBox.question(self, "Confirm",
                                    f"Change department to '{dept}' for {len(selected)} employee(s)?",
                                    QMessageBox.Yes | QMessageBox.No)

        if reply == QMessageBox.Yes:
            for emp in selected:
                self.db.conn.execute("UPDATE employees SET department=? WHERE emp_id=?",
                    (dept, emp['emp_id']))
            self.db.conn.commit()
            show_success_toast(self, f"Updated {len(selected)} employee(s)!")
            self.status_label.setText(f"✅ Changed department for {len(selected)} employees")

    def bulk_change_position(self):
        selected = self.get_selected_employees()
        if not selected:
            show_warning_toast(self, "Please select employees first.")
            return

        position, ok = QInputDialog.getText(self, "Change Position", "New position:")
        if not ok or not position: return

        reply = QMessageBox.question(self, "Confirm",
                                    f"Change position to '{position}' for {len(selected)} employee(s)?",
                                    QMessageBox.Yes | QMessageBox.No)

        if reply == QMessageBox.Yes:
            for emp in selected:
                self.db.conn.execute("UPDATE employees SET position=? WHERE emp_id=?",
                    (position, emp['emp_id']))
            self.db.conn.commit()
            show_success_toast(self, f"Updated {len(selected)} employee(s)!")
            self.status_label.setText(f"✅ Changed position for {len(selected)} employees")

    def bulk_change_agency(self):
        selected = self.get_selected_employees()
        if not selected:
            show_warning_toast(self, "Please select employees first.")
            return

        agencies = ["Direct Hire"] + self.db.get_agencies()
        agency, ok = QInputDialog.getItem(self, "Change Agency", "New agency:", agencies, 0, False)
        if not ok: return

        agency_value = None if agency == "Direct Hire" else agency

        reply = QMessageBox.question(self, "Confirm",
                                    f"Change agency to '{agency}' for {len(selected)} employee(s)?",
                                    QMessageBox.Yes | QMessageBox.No)

        if reply == QMessageBox.Yes:
            for emp in selected:
                self.db.conn.execute("UPDATE employees SET agency=? WHERE emp_id=?",
                    (agency_value, emp['emp_id']))
            self.db.conn.commit()
            show_success_toast(self, f"Updated {len(selected)} employee(s)!")
            self.status_label.setText(f"✅ Changed agency for {len(selected)} employees")

    def bulk_archive(self):
        selected = self.get_selected_employees()
        if not selected:
            show_warning_toast(self, "Please select employees first.")
            return

        reason, ok = QInputDialog.getText(self, "Archive Reason", "Reason for archiving:")
        if not ok or not reason: return

        reply = QMessageBox.question(self, "Confirm",
                                    f"Archive {len(selected)} employee(s)?",
                                    QMessageBox.Yes | QMessageBox.No)

        if reply == QMessageBox.Yes:
            for emp in selected:
                self.db.archive_employee(emp['emp_id'], "admin", reason)
            show_success_toast(self, f"Archived {len(selected)} employee(s)!")
            self.status_label.setText(f"✅ Archived {len(selected)} employees")

    def bulk_export(self):
        selected = self.get_selected_employees()
        if not selected:
            show_warning_toast(self, "Please select employees first.")
            return

        filename, _ = QFileDialog.getSaveFileName(
            self, "Export Selected Employees",
            f"selected_employees_{datetime.now().strftime('%Y%m%d')}.json",
            "JSON Files (*.json)")

        if not filename: return

        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(selected, f, indent=2, ensure_ascii=False)
            show_success_toast(self, f"Exported {len(selected)} employee(s) to:\n{filename}")
            self.status_label.setText(f"✅ Exported {len(selected)} employees")
        except Exception as e:
            show_error_toast(self, f"Export failed:\n{str(e)}")


# ==================== v2.0 DIALOG CLASSES ====================

class PrintPreviewHelper(QDialog):
    """Print preview dialog"""
    def __init__(self, document_html, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Print Preview")
        self.setMinimumSize(900, 700)

        layout = QVBoxLayout(self)

        # Web view for preview
        self.web_view = QWebEngineView()
        self.web_view.setHtml(document_html)
        layout.addWidget(self.web_view)

        # Buttons
        btn_layout = QHBoxLayout()
        print_btn = ModernAnimatedButton("🖨️ Print")
        print_btn.clicked.connect(self.print_document)
        cancel_btn = ModernAnimatedButton("Cancel")
        cancel_btn.clicked.connect(self.reject)

        btn_layout.addStretch()
        btn_layout.addWidget(print_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

        self.should_print = False

    def print_document(self):
        self.should_print = True
        self.accept()


class PermissionEditorDialog(QDialog):
    """Dialog for editing user permissions"""
    def __init__(self, db, username, parent=None):
        super().__init__(parent)
        self.db = db
        self.username = username
        self.setWindowTitle(f"Edit Permissions - {username}")
        self.setMinimumSize(600, 700)

        layout = QVBoxLayout(self)

        # Info label
        info = QLabel(f"<b>Configure permissions for user: {username}</b>")
        layout.addWidget(info)

        # Scroll area for permissions
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)

        # Load current permissions
        self.perms = self.db.get_user_permissions(username)

        # Permission checkboxes
        self.perm_checks = {}

        perm_groups = {
            "Dashboard & Employees": [
                ("dashboard", "Access Dashboard"),
                ("employees", "View Employees"),
                ("add_employee", "Add New Employees"),
                ("edit_employee", "Edit Employees"),
                ("delete_employee", "Delete Employees"),
            ],
            "Features": [
                ("print_system", "Access Printing System"),
                ("bulk_operations", "Bulk Operations"),
                ("reports", "Generate Reports"),
                ("letters", "Generate Letters"),
            ],
            "Administration": [
                ("user_management", "Manage Users"),
                ("settings", "Access Settings"),
                ("audit_log", "View Audit Log"),
                ("backup_restore", "Backup & Restore"),
                ("archive", "Archive Management"),
            ]
        }

        for group_name, permissions in perm_groups.items():
            group_box = QGroupBox(group_name)
            group_layout = QVBoxLayout()

            for perm_key, perm_label in permissions:
                cb = QCheckBox(perm_label)
                cb.setChecked(self.perms.get(perm_key, False))
                self.perm_checks[perm_key] = cb
                group_layout.addWidget(cb)

            group_box.setLayout(group_layout)
            scroll_layout.addWidget(group_box)

        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        layout.addWidget(scroll)

        # Buttons
        btn_layout = QHBoxLayout()
        save_btn = PulseButton("💾 Save Permissions")
        save_btn.start_pulse()
        save_btn.clicked.connect(self.save_permissions)
        cancel_btn = ModernAnimatedButton("Cancel")
        cancel_btn.clicked.connect(self.reject)

        btn_layout.addStretch()
        btn_layout.addWidget(save_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

    def save_permissions(self):
        new_perms = {}
        for key, cb in self.perm_checks.items():
            new_perms[key] = cb.isChecked()

        if self.db.update_user_permissions(self.username, new_perms):
            show_success_toast(self, "Permissions updated successfully!")
            self.accept()
        else:
            show_error_toast(self, "Failed to update permissions")


class StoreManagementDialog(QDialog):
    """Dialog for managing stores"""
    def __init__(self, db, parent=None):
        super().__init__(parent)
        self.db = db
        self.setWindowTitle("Manage Stores/Branches")
        self.setMinimumSize(800, 600)

        layout = QVBoxLayout(self)

        # Title
        title = QLabel("<h2>🏪 Store Management</h2>")
        layout.addWidget(title)

        # Toolbar
        toolbar = QHBoxLayout()
        add_btn = ModernAnimatedButton("➕ Add Store")
        add_btn.clicked.connect(self.add_store)
        edit_btn = ModernAnimatedButton("✏️ Edit Store")
        edit_btn.clicked.connect(self.edit_store)
        toggle_btn = ModernAnimatedButton("🔄 Activate/Deactivate")
        toggle_btn.clicked.connect(self.toggle_store)
        refresh_btn = ModernAnimatedButton("🔄 Refresh")
        refresh_btn.clicked.connect(self.load_stores)

        toolbar.addWidget(add_btn)
        toolbar.addWidget(edit_btn)
        toolbar.addWidget(toggle_btn)
        toolbar.addWidget(refresh_btn)
        toolbar.addStretch()
        layout.addLayout(toolbar)

        # Table
        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["ID", "Company", "Branch", "Address", "Status"])
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        layout.addWidget(self.table)

        # Close button
        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)

        self.load_stores()

    def load_stores(self):
        stores = self.db.get_all_stores()
        self.table.setRowCount(len(stores))

        for row, store in enumerate(stores):
            self.table.setItem(row, 0, QTableWidgetItem(str(store['id'])))
            self.table.setItem(row, 1, QTableWidgetItem(store['company_name']))
            self.table.setItem(row, 2, QTableWidgetItem(store['branch_name']))
            self.table.setItem(row, 3, QTableWidgetItem(store.get('address', '')))
            status = "✅ Active" if store['active'] else "❌ Inactive"
            self.table.setItem(row, 4, QTableWidgetItem(status))

    def add_store(self):
        dialog = AddEditStoreDialog(self.db, parent=self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.load_stores()

    def edit_store(self):
        row = self.table.currentRow()
        if row < 0:
            show_warning_toast(self, "Please select a store to edit")
            return

        store_id = int(self.table.item(row, 0).text())
        dialog = AddEditStoreDialog(self.db, store_id=store_id, parent=self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.load_stores()

    def toggle_store(self):
        row = self.table.currentRow()
        if row < 0:
            show_warning_toast(self, "Please select a store")
            return

        store_id = int(self.table.item(row, 0).text())
        status_text = self.table.item(row, 4).text()
        is_active = "Active" in status_text

        action = "deactivate" if is_active else "activate"
        reply = QMessageBox.question(self, "Confirm",
                                     f"Are you sure you want to {action} this store?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            if self.db.toggle_store_active(store_id, not is_active):
                self.load_stores()
            else:
                show_error_toast(self, "Failed to update store status")


class AddEditStoreDialog(QDialog):
    """Dialog for adding/editing a store"""
    def __init__(self, db, store_id=None, parent=None):
        super().__init__(parent)
        self.db = db
        self.store_id = store_id
        self.setWindowTitle("Edit Store" if store_id else "Add New Store")
        self.setMinimumWidth(500)

        layout = QVBoxLayout(self)
        form = QVBoxLayout()

        # Company Name
        comp_label = QLabel("Company Name:")
        comp_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        form.addWidget(comp_label)
        self.company_edit = NeumorphicGradientLineEdit("Company Name")
        self.company_edit.setMinimumHeight(70)
        form.addWidget(self.company_edit)

        # Branch Name
        branch_label = QLabel("Branch Name:")
        branch_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        form.addWidget(branch_label)
        self.branch_edit = NeumorphicGradientLineEdit("Branch Name")
        self.branch_edit.setMinimumHeight(70)
        form.addWidget(self.branch_edit)

        # Address
        addr_label = QLabel("Address:")
        addr_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        form.addWidget(addr_label)
        self.address_edit = NeumorphicGradientTextEdit("Address", min_height=100)
        self.address_edit.setMinimumHeight(120)
        form.addWidget(self.address_edit)

        layout.addLayout(form)

        # Buttons
        btn_layout = QHBoxLayout()
        save_btn = PulseButton("💾 Save")
        save_btn.start_pulse()
        save_btn.clicked.connect(self.save_store)
        cancel_btn = ModernAnimatedButton("Cancel")
        cancel_btn.clicked.connect(self.reject)

        btn_layout.addStretch()
        btn_layout.addWidget(save_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

        # Load existing data if editing
        if self.store_id:
            self.load_store_data()

    def load_store_data(self):
        stores = self.db.get_all_stores()
        for store in stores:
            if store['id'] == self.store_id:
                self.company_edit.line_edit.setText(store['company_name'])
                self.branch_edit.line_edit.setText(store['branch_name'])
                self.address_edit.setPlainText(store.get('address', ''))
                break

    def save_store(self):
        company = self.company_edit.line_edit.text().strip()
        branch = self.branch_edit.line_edit.text().strip()
        address = self.address_edit.toPlainText().strip()

        if not company or not branch:
            show_warning_toast(self, "Company and Branch names are required")
            return

        if self.store_id:
            success = self.db.update_store(self.store_id, company, branch, address)
        else:
            success = self.db.add_store(company, branch, address)

        if success:
            show_success_toast(self, "Store saved successfully!")
            self.accept()
        else:
            show_error_toast(self, "Failed to save store")


class LetterGenerationDialog(QDialog):
    """Dialog for generating excuse letters"""
    def __init__(self, db, current_user, employee=None, parent=None):
        super().__init__(parent)
        self.db = db
        self.current_user = current_user
        self.setWindowTitle("Generate Excuse Letter")
        self.setMinimumSize(700, 600)

        layout = QVBoxLayout(self)

        # Title
        title = QLabel("<h2>📝 Generate Excuse Letter</h2>")
        layout.addWidget(title)

        # Form with neumorphic gradient styling
        form = QVBoxLayout()

        # Employee selection
        emp_label = QLabel("Employee:")
        emp_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        form.addWidget(emp_label)
        self.employee_combo = NeumorphicGradientComboBox("Select Employee")
        self.employee_combo.setMinimumHeight(70)
        employees = self.db.all_employees()
        for emp in employees:
            self.employee_combo.addItem(emp['name'], emp['emp_id'])
        if employee:
            idx = self.employee_combo.findData(employee['emp_id'])
            if idx >= 0:
                self.employee_combo.setCurrentIndex(idx)
        form.addWidget(self.employee_combo)

        # Store/Branch
        store_label = QLabel("Store/Branch:")
        store_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        form.addWidget(store_label)
        self.store_branch = NeumorphicGradientLineEdit("e.g., SM North Edsa, Robinson's Galleria")
        self.store_branch.setMinimumHeight(70)
        self.store_branch.line_edit.editingFinished.connect(lambda: self._format_title_case(self.store_branch.line_edit))
        form.addWidget(self.store_branch)

        # Supervisor Name
        sup_label = QLabel("Supervisor Name:")
        sup_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        form.addWidget(sup_label)
        self.supervisor_name = NeumorphicGradientLineEdit("e.g., Juan Dela Cruz")
        self.supervisor_name.setMinimumHeight(70)
        self.supervisor_name.line_edit.editingFinished.connect(lambda: self._format_title_case(self.supervisor_name.line_edit))
        form.addWidget(self.supervisor_name)

        # Supervisor Title
        title_label = QLabel("Supervisor Title:")
        title_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        form.addWidget(title_label)
        self.supervisor_title = NeumorphicGradientLineEdit("e.g., Branch Manager, HR Manager")
        self.supervisor_title.setMinimumHeight(70)
        self.supervisor_title.line_edit.editingFinished.connect(lambda: self._format_title_case(self.supervisor_title.line_edit))
        form.addWidget(self.supervisor_title)

        # Company Name
        comp_label = QLabel("Company Name (Recipient):")
        comp_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        form.addWidget(comp_label)
        self.company_name = NeumorphicGradientLineEdit("e.g., INTERNATIONAL TOYWORLD INC.")
        self.company_name.setMinimumHeight(70)
        self.company_name.line_edit.editingFinished.connect(lambda: self._format_title_case(self.company_name.line_edit))
        form.addWidget(self.company_name)

        # Date options
        date_group = QGroupBox("Letter Date")
        date_group.setStyleSheet("QGroupBox { background: transparent; color: rgba(255, 255, 255, 0.9); }")
        date_layout = QVBoxLayout()

        date_opt_label = QLabel("Date Option:")
        date_opt_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        date_layout.addWidget(date_opt_label)
        self.date_option = NeumorphicGradientComboBox("Select Date Option")
        self.date_option.setMinimumHeight(70)
        self.date_option.combo_box.setFocusPolicy(Qt.ClickFocus)
        self.date_option.addItems(["Today", "Yesterday", "Specific Date", "Date Range"])
        self.date_option.currentTextChanged.connect(self._on_date_option_changed)
        date_layout.addWidget(self.date_option)

        self.date_edit = DatePicker()
        self.date_edit.setDate(QDate.currentDate())
        self.date_edit.setVisible(False)
        date_layout.addWidget(self.date_edit)

        self.date_range_layout = QHBoxLayout()
        self.start_date = DatePicker()
        self.start_date.setDate(QDate.currentDate())
        self.end_date = DatePicker()
        self.end_date.setDate(QDate.currentDate())
        self.date_range_layout.addWidget(QLabel("From:"))
        self.date_range_layout.addWidget(self.start_date)
        self.date_range_layout.addWidget(QLabel("To:"))
        self.date_range_layout.addWidget(self.end_date)
        date_range_widget = QWidget()
        date_range_widget.setLayout(self.date_range_layout)
        date_range_widget.setVisible(False)
        self.date_range_widget = date_range_widget
        date_layout.addWidget(date_range_widget)
        
        self._wheel_guard = WheelGuard()
        self.date_option.installEventFilter(self._wheel_guard)
        self.date_edit.installEventFilter(self._wheel_guard)
        self.start_date.installEventFilter(self._wheel_guard)
        self.end_date.installEventFilter(self._wheel_guard)

        date_group.setLayout(date_layout)
        form.addWidget(date_group)

        # Reason
        reason_label = QLabel("Reason:")
        reason_label.setStyleSheet("background: transparent; color: rgba(255, 255, 255, 0.9); margin-top: 5px;")
        form.addWidget(reason_label)
        self.reason_edit = NeumorphicGradientLineEdit("e.g., personal reasons, family emergency, illness")
        self.reason_edit.setMinimumHeight(70)
        form.addWidget(self.reason_edit)

        layout.addLayout(form)

        # Buttons
        btn_layout = QHBoxLayout()
        preview_btn = ModernAnimatedButton("👁️ Preview")
        preview_btn.clicked.connect(self.preview_letter)
        generate_btn = ModernAnimatedButton("📄 Generate Letter")
        generate_btn.clicked.connect(self.generate_letter)
        cancel_btn = ModernAnimatedButton("Cancel")
        cancel_btn.clicked.connect(self.reject)

        btn_layout.addStretch()
        btn_layout.addWidget(preview_btn)
        btn_layout.addWidget(generate_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)


    def _format_title_case(self, line_edit):
        """Format text to title case (first letter of each word capitalized)"""
        cursor_pos = line_edit.cursorPosition()
        text = line_edit.text()
        formatted = titlecase(text)
        if text != formatted:
            line_edit.blockSignals(True)
            line_edit.setText(formatted)
            line_edit.setCursorPosition(min(cursor_pos, len(formatted)))
            line_edit.blockSignals(False)

    def _on_date_option_changed(self, option):
        self.date_edit.setVisible(option == "Specific Date")
        self.date_range_widget.setVisible(option == "Date Range")

    def _get_letter_date_string(self):
        option = self.date_option.currentText()
        if option == "Today":
            return datetime.now().strftime("%B %d, %Y")
        elif option == "Yesterday":
            return (datetime.now() - timedelta(days=1)).strftime("%B %d, %Y")
        elif option == "Specific Date":
            return self.date_edit.date().toString("MMMM dd, yyyy")
        elif option == "Date Range":
            start = self.start_date.date().toString("MMMM dd, yyyy")
            end = self.end_date.date().toString("MMMM dd, yyyy")
            return f"{start} to {end}"
        return ""

    def _generate_letter_content(self):
        # Get employee
        emp_id = self.employee_combo.currentData()
        employees = self.db.all_employees()
        employee = next((e for e in employees if e['emp_id'] == emp_id), None)
        if not employee:
            return None

        # Get store/branch from text field
        store_branch_text = self.store_branch.line_edit.text().strip()
        if not store_branch_text:
            store_branch_text = "Office"
        
        # Make sure an employee is selected
        if not self.employee_combo.currentData():
            show_warning_toast(self, "Please select an Employee")
            self.employee_combo.setFocus()
            return
        
        # Make sure Company Name (Recipient) is filled
        if not self.company_name.line_edit.text().strip():
            show_warning_toast(self, "Please enter Company Name (Recipient)")
            self.company_name.setFocus()
            return
        
        # v3.9: Get template or use default
        template = self.db.get_letter_template(1)
        if not template:
            # Use default template if database template is empty
            template = """[DATE]

The Store Manager
[COMPANY_NAME]
[BRANCH_NAME]

Dear Sir/Madam,

This is to inform you that [EMPLOYEE_NAME] (ID: [EMPLOYEE_ID]) was unable to report for work on [LETTER_DATE] due to [REASON].

We kindly request your understanding regarding this matter.

Thank you for your consideration.

Respectfully yours,

_________________________________
[SUPERVISOR_NAME]
[SUPERVISOR_TITLE]"""
        
        # Replace placeholders
        employee['name'] = employee['name'].split(' (Employee ID')[0]
        letter_date = self._get_letter_date_string()
        content = template.replace("[DATE]", datetime.now().strftime("%B %d, %Y"))
        company = titlecase(self.company_name.line_edit.text().strip())
        if not company:
            company = "INTERNATIONAL TOYWORLD INC."

        content = content.replace("[COMPANY_NAME]", company)
        content = content.replace("[BRANCH_NAME]", store_branch_text)
        content = content.replace("[ADDRESS]", "")
        content = content.replace("[EMPLOYEE_NAME]", employee['name'])
        content = content.replace("[EMPLOYEE_ID]", employee['emp_id'])
        content = content.replace("[LETTER_DATE]", letter_date)
        content = re.sub(r'\s*\(Employee ID:.*?\)', '', content)
        content = content.replace(
            "[REASON]",
            titlecase(self.reason_edit.line_edit.text().strip()) or "Personal Reasons"
        )
        content = content.replace("[SUPERVISOR_NAME]", self.supervisor_name.line_edit.text().strip())
        content = content.replace("[SUPERVISOR_TITLE]", self.supervisor_title.line_edit.text().strip())
        
        # Clean up
        content = re.sub(r'(?mi)^\s*INTERNATIONAL TOYWORLD INC\.\s*$', '', content).strip()
        content = re.sub(r'(?m)^(The Store Manager)\s*\n\s*\n', r'\1\n', content)
        content = re.sub(r'(?mi)^\s*re:\s*.*\r?\n?', '', content)
        content = re.sub(r'\n{3,}', '\n\n', content).strip()
        
        return content

    def preview_letter(self):
        content = self._generate_letter_content()
        if not content:
            show_warning_toast(self, "Failed to generate letter content")
            return

        # Show in message box
        # v4.4.1: Animated dialog for letter preview
        from employee_vault.ui.widgets import QuickAnimatedDialog
        preview_dialog = QuickAnimatedDialog(self, animation_style="fade")
        preview_dialog.setWindowTitle("Letter Preview")
        preview_dialog.setMinimumSize(600, 500)
        layout = QVBoxLayout(preview_dialog)

        text_edit = NeumorphicGradientTextEdit("Letter Content", min_height=400)
        text_edit.setMinimumHeight(420)
        text_edit.text_edit.setPlainText(content)
        text_edit.text_edit.setReadOnly(True)
        layout.addWidget(text_edit)

        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(preview_dialog.accept)
        layout.addWidget(close_btn)

        preview_dialog.exec()

    def generate_letter(self):
        # Validate inputs
        if not self.store_branch.line_edit.text().strip():
            show_warning_toast(self, "Please enter Store/Branch name")
            self.store_branch.setFocus()
            return

        if not self.supervisor_name.line_edit.text().strip():
            show_warning_toast(self, "Please enter Supervisor Name")
            self.supervisor_name.setFocus()
            return

        if not self.supervisor_title.line_edit.text().strip():
            show_warning_toast(self, "Please enter Supervisor Title")
            self.supervisor_title.setFocus()
            return

        if not self.reason_edit.line_edit.text().strip():
            show_warning_toast(self, "Please enter a reason for the excuse letter")
            self.reason_edit.setFocus()
            return

        content = self._generate_letter_content()
        if not content:
            show_warning_toast(self, "Failed to generate letter content")
            return

        # Get employee info
        emp_id = self.employee_combo.currentData()

        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename_docx = f"excuse_letter_{emp_id}_{timestamp}.docx"
        filepath = os.path.join(LETTERS_DIR, filename_docx)

        try:
            # Create Word document
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Set margins - more space for header and footer
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.5)  # More space for header
                section.bottom_margin = Inches(1.2)  # More space for footer
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                
                # Add header with company logo
                        # ==================== ENHANCED HEADER WITH PROPER FORMATTING ====================
                header = section.header
                header_para = header.paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Try to add header image first
                header_image_path = resource_path("cuddly_header.png")
                if not os.path.exists(header_image_path):
                    for possible_path in ["cuddly_header.png", 
                                          os.path.join(os.path.dirname(__file__), "cuddly_header.png"),
                                          os.path.join(os.getcwd(), "cuddly_header.png")]:
                        if os.path.exists(possible_path):
                            header_image_path = possible_path
                            break
                
                image_loaded_successfully = False
                if os.path.exists(header_image_path):
                    try:
                        # Best solution: User provides 'cuddly_header.png' with logo, text, and lines
                        run = header_para.add_run()
                        run.add_picture(header_image_path, width=Inches(6.2))
                        image_loaded_successfully = True # Image loaded
                    except Exception as img_error:
                        logging.warning(f"Could not add header image: {img_error}")
                        # Fallback to text header
                        pass # Will fall through to the 'if not image_loaded_successfully' block
                
                if not image_loaded_successfully:
                    # Text-based header (fallback) - REPLICATING SCREENSHOT
                    # NOTE: This is a centered approximation. A logo on the left
                    # and text on the right is not feasible with this text fallback.
                    run = header_para.add_run("Cuddly International Corp.\n")
                    run.font.bold = True
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(50, 50, 50) # Dark gray
                    
                    run = header_para.add_run(
                        "650 Jesus Ext. cor. Beata, Pandacan, Manila, PHILIPPINES\n"
                        "Tel: (632) 588-0324 to 25  Fax: (632) 588-0327\n"
                        "email: sales@cuddlyinternational.com\n"
                    )
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(80, 80, 80)
                
                # --- The old single horizontal line was here ---
                
                # --- ALWAYS ADD DOUBLE LINE ---
                # Add DOUBLE horizontal line after header (runs for both image and text fallback)
                # First line
                line1_para = header.add_paragraph()
                line1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                line1_para.paragraph_format.space_before = Pt(4) # Space before lines
                line1_para.paragraph_format.space_after = Pt(0) # 0pt gap
                run1 = line1_para.add_run("_" * 150) # Use underscore
                run1.font.size = Pt(5) # Make it very small to look solid
                run1.font.color.rgb = RGBColor(50, 50, 50) # Dark line
                
                # Second line
                line2_para = header.add_paragraph()
                line2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                line2_para.paragraph_format.space_before = Pt(0)
                line2_para.paragraph_format.space_after = Pt(0)
                run2 = line2_para.add_run("_" * 150) # Use underscore
                run2.font.size = Pt(5) # Make it very small to look solid
                run2.font.color.rgb = RGBColor(50, 50, 50) # Dark line
                
                # ==================== ENHANCED FOOTER WITH PROPER FORMATTING ====================
                footer = section.footer
                
                # Add horizontal line before footer
                footer_line_para = footer.add_paragraph()
                footer_line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = footer_line_para.add_run("_" * 120)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(128, 128, 128)
                
                footer_para = footer.add_paragraph()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Try to add footer image first
                footer_image_path = resource_path("cuddly_footer.png")
                if not os.path.exists(footer_image_path):
                    for possible_path in ["cuddly_footer.png",
                                          os.path.join(os.path.dirname(__file__), "cuddly_footer.png"),
                                          os.path.join(os.getcwd(), "cuddly_footer.png")]:
                        if os.path.exists(possible_path):
                            footer_image_path = possible_path
                            break
                
                if os.path.exists(footer_image_path):
                    try:
                        run = footer_para.add_run()
                        run.add_picture(footer_image_path, width=Inches(5.8))
                    except Exception as img_error:
                        logging.warning(f"Could not add footer image: {img_error}")
                        # Fallback to text footer
                        run = footer_para.add_run("Our Brands:\n")
                        run.font.bold = True
                        run.font.size = Pt(10)
                        
                        run = footer_para.add_run("APRUVA  |  APRUVA FOR MOM & BABY  |  POSH\n")
                        run.font.bold = True
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(192, 0, 0)
                        
                        run = footer_para.add_run(
                            "Quality Products for Your Family\n"
                            "© 2025 Cuddly International Corp. All Rights Reserved."
                        )
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(96, 96, 96)
                else:
                    # Text-based footer (fallback)
                    run = footer_para.add_run("Our Brands:\n")
                    run.font.bold = True
                    run.font.size = Pt(10)
                    
                    run = footer_para.add_run("APRUVA  |  APRUVA FOR MOM & BABY  |  POSH\n")
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(192, 0, 0)
                    
                    run = footer_para.add_run(
                        "Quality Products for Your Family\n"
                        "© 2025 Cuddly International Corp. All Rights Reserved."
                    )
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(96, 96, 96)


            # Save document
            doc.save(filepath)

            # Save to history
            letter_date = self._get_letter_date_string()
            self.db.save_letter_history(
                emp_id, "excuse", letter_date, None,
                titlecase(self.reason_edit.line_edit.text().strip()),
                self.supervisor_name.line_edit.text().strip(),
                self.supervisor_title.line_edit.text().strip(),
                filepath, self.current_user
            )

            reply = QMessageBox.question(
                self, "Success",
                f"Letter generated successfully!\n\nSaved to: {filepath}\n\nOpen the letter now?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )

            if reply == QMessageBox.StandardButton.Yes:
                QDesktopServices.openUrl(QUrl.fromLocalFile(filepath))

            self.accept()

        except Exception as e:
            show_error_toast(self, f"Failed to save letter:\n{str(e)}")


class SessionMonitorDialog(QDialog):
    """Dialog for monitoring active sessions"""
    def __init__(self, db, parent=None):
        super().__init__(parent)
        self.db = db
        self.setWindowTitle("Active Sessions")
        self.setMinimumSize(700, 400)

        layout = QVBoxLayout(self)

        # Title
        title = QLabel("<h2>👥 Active Sessions</h2>")
        layout.addWidget(title)

        # Toolbar
        toolbar = QHBoxLayout()
        refresh_btn = ModernAnimatedButton("🔄 Refresh")
        refresh_btn.clicked.connect(self.load_sessions)
        toolbar.addWidget(refresh_btn)
        toolbar.addStretch()
        layout.addLayout(toolbar)

        # Table
        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(["Username", "Computer", "Login Time", "Last Activity", "Duration"])
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        layout.addWidget(self.table)

        # Close button
        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(self.accept)
        layout.addWidget(close_btn)

        self.load_sessions()

    def load_sessions(self):
        sessions = self.db.get_active_sessions()
        self.table.setRowCount(len(sessions))

        for row, session in enumerate(sessions):
            self.table.setItem(row, 0, QTableWidgetItem(session['username']))
            self.table.setItem(row, 1, QTableWidgetItem(session.get('computer_name', 'Unknown')))
            self.table.setItem(row, 2, QTableWidgetItem(session['login_time']))
            self.table.setItem(row, 3, QTableWidgetItem(session['last_activity']))

            # Calculate duration
            try:
                login_time = datetime.fromisoformat(session['login_time'])
                duration = datetime.now() - login_time
                hours = int(duration.total_seconds() // 3600)
                minutes = int((duration.total_seconds() % 3600) // 60)
                duration_str = f"{hours}h {minutes}m"
            except (ValueError, AttributeError):
                # Invalid datetime format or calculation error
                duration_str = "Unknown"

            self.table.setItem(row, 4, QTableWidgetItem(duration_str))


# ==================== END v2.0 DIALOG CLASSES ====================


class MainWindow(QMainWindow):
    def __init__(self, db, username, user_row, icon=None):
        super().__init__(); self.setWindowTitle(f"{APP_TITLE} — {user_row['name']}")
        
        # v3.3: Modern window styling - themed title bar
        self.setAttribute(Qt.WA_TranslucentBackground, False)  # Keep solid background
        
        # Apply modern styling to the window
        current_theme = load_theme_preference()
        theme_colors = MODERN_THEMES.get(current_theme, MODERN_THEMES["dark_blue"])
        
        # Set window background to match theme
        self.setStyleSheet(f"""
            QMainWindow {{
                background-color: {theme_colors['background']};
                border: 2px solid {theme_colors['primary']};
                border-radius: 12px;
            }}
        """)
        
        if icon:
            self.setWindowIcon(icon)
        # Set fixed window size
        self.setMinimumSize(1200, 600)
        self.resize(1200, 600)
        # Center the window properly
        screen = QApplication.primaryScreen().availableGeometry()
        self.move(
            (screen.width() - self.width()) // 2,
            (screen.height() - self.height()) // 2
        )

        self.db=db; self.current_user=username; self.user_row=user_row
        self.employees=self.db.all_employees(); self._db_mtime=db_latest_mtime(DB_FILE)

        # v2.0: Load user permissions
        self.user_permissions = self.db.get_user_permissions(self.current_user)

        # v2.0: Create session tracking
        try:
            computer_name = socket.gethostname()
            self.db.create_session(self.current_user, "", computer_name)
        except Exception as e:
            logging.warning(f"Could not create session: {e}")

        # v2.0: Setup session keepalive timer
        self.session_timer = QTimer(self)
        self.session_timer.timeout.connect(lambda: self.db.update_session_activity(self.current_user))
        self.session_timer.start(60000)  # Update every minute

        # WEEK 2 FEATURE #1: Idle timeout (auto-lock after inactivity)
        self.idle_timeout_minutes = 30  # Configurable: 30 minutes default
        self.last_activity_time = datetime.now()
        self.idle_timer = QTimer(self)
        self.idle_timer.timeout.connect(self._check_idle_timeout)
        self.idle_timer.start(60000)  # Check every minute

        # Install event filter to track user activity
        self.installEventFilter(self)

        # Log login action
        self.db.log_action(username=self.current_user, action="LOGIN", details=f"User logged in")

        self.stack=QStackedWidget(); self.stack.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.dashboard = EnhancedDashboardPage()

        # Create the form first
        self.emp_form = EmployeeForm(self.db, self.current_user, self._on_form_saved)

        # Create a scroll area and put the form inside it
        self.form_scroll_area = QScrollArea()
        self.form_scroll_area.setWidgetResizable(True)
        self.form_scroll_area.setWidget(self.emp_form)

        self.employees_page=EmployeesPage(self._view_employee, self._edit_employee, self._delete_selected)
        self.employees_page.advanced_search_btn.clicked.connect(self._show_enhanced_search)
        self.dashboard.refresh(self.employees); self.employees_page.set_data(self.employees)

        # Sidebar
        # Create the sidebar frame and store a reference on the instance so
        # toggle_sidebar() can animate its width.  We set both a
        # fixed width and maximum width initially to the expanded size.
        sidebar = QFrame()
        self.sidebar = sidebar
        # Expanded and collapsed widths per reference dashboard
        # 240px expanded, 80px collapsed for modern, clean look
        self.sidebar_expanded_width = 240
        self.sidebar_collapsed_width = 80
        self.is_sidebar_collapsed = False
        self.sidebar.setFixedWidth(self.sidebar_expanded_width)
        self.sidebar.setMaximumWidth(self.sidebar_expanded_width)
        sidebar_main_layout = QVBoxLayout(sidebar)
        sidebar_main_layout.setContentsMargins(0, 0, 0, 0)
        sidebar_main_layout.setSpacing(0)

        # v3.8: COMPACT MODERN HEADER with Logo
        header_widget = QWidget()
        header_widget.setStyleSheet(f"""
            QWidget {{
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #2196F3, stop:1 #1976D2);
                padding: 8px;
            }}
        """)
        self.header_layout = QHBoxLayout(header_widget)
        self.header_layout.setContentsMargins(10, 5, 10, 5)
        self.header_layout.setSpacing(8)

        # Sidebar collapse/expand toggle button.  Move this to the far
        # right of the header and reduce its size for a more subtle
        # appearance.  The arrow glyph flips depending on collapsed
        # state.  Styling uses a semi‑transparent background and a
        # smaller font to maintain legibility without dominating the
        # header.
        self.toggle_sidebar_btn = ModernAnimatedButton("❮")
        # Make the button very small and compact to fit nicely in the header
        self.toggle_sidebar_btn.setFixedSize(12, 12)
        self.toggle_sidebar_btn.setStyleSheet(
            "QPushButton {"
            "  background-color: rgba(255, 255, 255, 0.15);"
            "  color: white;"
            "  border: none;"
            "  font-size: 8px;"
            "  border-radius: 2px;"
            "  padding: 0px;"
            "}"
            "QPushButton:hover {"
            "  background-color: rgba(255, 255, 255, 0.25);"
            "}"
        )
        # Cursor removed per user request
        # self.toggle_sidebar_btn.setCursor(Qt.PointingHandCursor)
        self.toggle_sidebar_btn.setToolTip("Collapse sidebar")
        self.toggle_sidebar_btn.clicked.connect(self._toggle_sidebar)
        # Note: We do not add the toggle button here; it will be
        # appended after the company label to position it on the right.
        
        # Try to load company logo.  Use instance attributes so they can
        # be shown/hidden when collapsing the sidebar.
        try:
            self.logo_label = QLabel()
            logo_pixmap = QPixmap("company_logo.png")
            if not logo_pixmap.isNull():
                logo_pixmap = logo_pixmap.scaled(40, 40, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                self.logo_label.setPixmap(logo_pixmap)
            else:
                self.logo_label.setText("🏢")
                self.logo_label.setStyleSheet("font-size: 30px;")
            self.header_layout.addWidget(self.logo_label)
        except Exception:
            # If loading fails simply don't add the logo
            self.logo_label = None

        # Modern solid gradient company name header
        self.company_label = QLabel("<b>Cuddly International<br>Corporation</b>")
        self.company_label.setStyleSheet("""
            QLabel {
                color: white;
                font-size: 13px;
                font-weight: bold;
                padding: 8px 16px;
                margin: 0 4px;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                                           stop:0 rgba(33, 150, 243, 0.9),
                                           stop:0.5 rgba(25, 118, 210, 0.75),
                                           stop:1 rgba(21, 101, 192, 0.6));
                border: 2px solid rgba(33, 150, 243, 0.8);
                border-top: 2px solid rgba(255, 255, 255, 0.3);
                border-left: 2px solid rgba(255, 255, 255, 0.2);
                border-radius: 22px;
            }
        """)
        self.header_layout.addWidget(self.company_label)
        self.header_layout.addStretch()
        # Append the toggle button at the end of the header layout so it
        # appears on the far right, vertically centered.
        self.header_layout.addWidget(self.toggle_sidebar_btn, 0, Qt.AlignVCenter)
        
        sidebar_main_layout.addWidget(header_widget)
        
        # Modern solid gradient USER INFO BOX with Photo
        ubox=QGroupBox()
        ubox.setStyleSheet("""
            QGroupBox {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                                           stop:0 rgba(100, 181, 246, 0.85),
                                           stop:0.5 rgba(66, 165, 245, 0.70),
                                           stop:1 rgba(33, 150, 243, 0.55));
                border: 2px solid rgba(100, 181, 246, 0.9);
                border-top: 2px solid rgba(255, 255, 255, 0.3);
                border-left: 2px solid rgba(255, 255, 255, 0.2);
                border-radius: 22px;
                padding: 12px;
                margin: 6px;
            }
            QLabel {
                background: transparent;
                border: none;
            }
        """)
        uvl=QHBoxLayout(ubox)
        uvl.setContentsMargins(8, 8, 8, 8)
        uvl.setSpacing(10)
        
        # v3.9: User photo - clickable to upload
        self.user_photo_label = QLabel()
        # Slightly larger avatar for better visibility.  A dynamic circle
        # size makes it easier to change later; we bind the size in
        # _load_user_photo() to this label's dimensions.  Increasing
        # from 45 to 56 pixels aligns it more closely with the larger
        # circular avatar used in the Add Employee panel.
        self.user_photo_label.setFixedSize(56, 56)
        # Cursor removed per user request
        # self.user_photo_label.setCursor(Qt.PointingHandCursor)  # Hand cursor
        self.user_photo_label.setStyleSheet(
            "QLabel {"
            "  border: 3px solid rgba(33, 150, 243, 0.9);"
            "  border-radius: 28px;"
            "  background-color: rgba(33, 150, 243, 0.3);"
            "}"
            "QLabel:hover {"
            "  border: 3px solid rgba(100, 181, 246, 1.0);"
            "  background-color: rgba(100, 181, 246, 0.4);"
            "}"
        )
        self.user_photo_label.setToolTip("Click to upload your photo")
        
        # Make clickable
        self.user_photo_label.mousePressEvent = lambda event: self._upload_user_photo()
        
        # Try to load user photo
        self._load_user_photo()
        
        uvl.addWidget(self.user_photo_label)
        
        # User info text - compact
        user_info_layout = QVBoxLayout()
        user_info_layout.setSpacing(0)
        name_label = QLabel(f"<b>{user_row['name']}</b>")
        name_label.setStyleSheet("color: white; font-size: 12px;")
        role_label = QLabel(f"{user_row['role']}")
        role_label.setStyleSheet("color: #b0b0b0; font-size: 11px;")
        user_info_layout.addWidget(name_label)
        user_info_layout.addWidget(role_label)
        uvl.addLayout(user_info_layout)
        uvl.addStretch()
        
        sidebar_main_layout.addWidget(ubox)
        # Keep a reference to the user info container so that we can
        # hide/show it when the sidebar collapses/expands.
        self.user_info_container = ubox

        # Create scroll area for sections
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.NoFrame)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        # Widget to hold all collapsible sections
        sections_widget = QWidget()
        s = QVBoxLayout(sections_widget)
        s.setContentsMargins(0, 0, 0, 0)
        s.setSpacing(0)

        # MAIN Section (Collapsible, starts collapsed) - BLUE
        main_section = CollapsibleSection("MAIN", start_collapsed=True, color="#4a9eff")
        main_section.add_button("  📊 Dashboard", lambda: self._show_page(self.dashboard))
        main_section.add_button("  👥 Employees", lambda: self._show_page(self.employees_page))
        main_section.add_button("  ➕ Add New", lambda: self._show_add_new())
        main_section.add_button("  📊 Reports", self._show_reports)
        s.addWidget(main_section)

        # --- Sidebar section tracking ---
        # Initialise the lists to hold the collapsible sections and their
        # corresponding icons.  These will be used by the sidebar
        # collapse/expand logic to adjust the button labels and tooltips.
        self.sidebar_sections = []
        self.section_icons = {}
        # Store the full title on the section so we can restore it
        # later.  Without this the toggle would lose the original
        # section name when switching back from collapsed mode.
        main_section.full_title = "MAIN"
        self.sidebar_sections.append(main_section)
        self.section_icons[main_section] = "🏠"

        # DATA & BACKUP Section (Collapsible, starts collapsed) - ORANGE
        backup_section = CollapsibleSection("DATA & BACKUP", start_collapsed=True, color="#ff9800")

        backup_section.add_button("  💾 Backup", self._backup_data)


        # Add Archive Manager for admins only
        if self.user_permissions.get('user_management'):
            backup_section.add_button("  ⏰ Scheduled Backup", self._show_scheduled_backup)
            backup_section.add_button("  📦 Archive Manager", self._show_archive_manager)
            backup_section.add_button("  📁 Export Data", self._export_data)
            backup_section.add_button("  📥 Import Data", self._import_data)
        s.addWidget(backup_section)
        backup_section.full_title = "DATA & BACKUP"
        self.sidebar_sections.append(backup_section)
        self.section_icons[backup_section] = "💾"

        # TOOLS Section (Collapsible, starts collapsed) - ORANGE RED
        tools_section = CollapsibleSection("TOOLS", start_collapsed=True, color="#ff5722")

        tools_section.add_button("  📦 Bulk Operations", self._show_bulk_operations)

        # v2.0: Letter generation (available to all users with permission)
        if self.user_permissions.get('letters', True):
            tools_section.add_button("  📝 Generate Letter", self._show_letter_generation)

        if self.user_permissions.get('user_management'):
            tools_section.add_button("  🆔 ID Card Generator", self._show_id_card_generator)
            tools_section.add_button("  🖨️ Print System", self._show_print_system)
            # v2.0: Session monitor (admin only)
            tools_section.add_button("  👥 Active Sessions", self._show_session_monitor)
        s.addWidget(tools_section)
        tools_section.full_title = "TOOLS"
        self.sidebar_sections.append(tools_section)
        self.section_icons[tools_section] = "🛠"

        # SYSTEM SETTINGS Section (Collapsible, starts collapsed) - PURPLE
        settings_section = CollapsibleSection("SYSTEM SETTINGS", start_collapsed=True, color="#9c27b0")

        settings_section.add_button("  🎨 Change Theme", self._show_theme_selector)
        settings_section.add_button("  ℹ️ About", self._show_about)

        # Add admin-only features
        if self.user_permissions.get('user_management'):
            settings_section.add_button("  👥 User Management", self._show_user_management)
            settings_section.add_button("  🌐 Network Config", self._show_network_config)
            settings_section.add_button("  📧 Email Notifications", self._show_email_settings)
            settings_section.add_button("  📜 Audit Log", self._show_audit_log)
            settings_section.add_button("  🔒 Security Audit", self._show_security_audit)

        s.addWidget(settings_section)
        settings_section.full_title = "SYSTEM SETTINGS"
        self.sidebar_sections.append(settings_section)
        self.section_icons[settings_section] = "⚙"

        # Add stretch to push content up but allow scrolling
        s.addStretch(1)

        # Set the sections widget to scroll area
        scroll_area.setWidget(sections_widget)
        sidebar_main_layout.addWidget(scroll_area)

        # Logout button (not scrollable, always visible at bottom)
        self.logout_btn=ModernAnimatedButton("🔓 Logout")
        self.logout_btn.clicked.connect(self._logout)
        sidebar_main_layout.addWidget(self.logout_btn)

        container=QWidget()
        layout=QHBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)  # Fix margins when maximized
        layout.setSpacing(0)
        layout.addWidget(sidebar)
        layout.addWidget(self.stack, 1)
        self.setCentralWidget(container)
        self.stack.addWidget(self.dashboard)
        self.stack.addWidget(self.employees_page)
        self.stack.addWidget(self.form_scroll_area) # <-- Add the SCROLL AREA here
        self._show_page(self.dashboard)

        # Tray icon
        pm=QPixmap(64,64); pm.fill(Qt.transparent); p=QPainter(pm); p.setRenderHint(QPainter.Antialiasing,True); p.setBrush(QColor("#3d6cff")); p.setPen(Qt.NoPen); p.drawEllipse(0,0,64,64); p.end()
        self.tray=QSystemTrayIcon(QIcon(pm), self); self.tray.setToolTip("Cuddly Employees Information"); self.tray.setVisible(True)

        # Timers - v2.0 FIXED: 3 seconds for better multi-user sync
        self.timer=QTimer(self); self.timer.setInterval(3000); self.timer.timeout.connect(self._auto_refresh); self.timer.start()
        self.contract_timer=QTimer(self); self.contract_timer.setInterval(30*60*1000); self.contract_timer.timeout.connect(self._notify_contracts); self.contract_timer.start()

        # v2.0 FIXED: Apply user permissions after UI is built
        self._apply_user_permissions()
        
        # Setup keyboard shortcuts
        self._setup_keyboard_shortcuts()

    def _setup_keyboard_shortcuts(self):
        """Setup global keyboard shortcuts for the application"""
        from PySide6.QtGui import QShortcut, QKeySequence
        
        # Ctrl+F - Focus search
        self.shortcut_search = QShortcut(QKeySequence("Ctrl+F"), self)
        self.shortcut_search.activated.connect(self._focus_search)
        
        # Ctrl+N - New employee
        self.shortcut_new = QShortcut(QKeySequence("Ctrl+N"), self)
        self.shortcut_new.activated.connect(self._show_add_new)
        
        # Ctrl+R or F5 - Refresh
        self.shortcut_refresh = QShortcut(QKeySequence("F5"), self)
        self.shortcut_refresh.activated.connect(self._force_refresh)
        self.shortcut_refresh2 = QShortcut(QKeySequence("Ctrl+R"), self)
        self.shortcut_refresh2.activated.connect(self._force_refresh)
        
        # Ctrl+E - Edit selected employee
        self.shortcut_edit = QShortcut(QKeySequence("Ctrl+E"), self)
        self.shortcut_edit.activated.connect(self._edit_selected)
        
        # Ctrl+S - Save (context-aware)
        self.shortcut_save = QShortcut(QKeySequence("Ctrl+S"), self)
        self.shortcut_save.activated.connect(self._shortcut_save)
        
        # Ctrl+B - Backup
        self.shortcut_backup = QShortcut(QKeySequence("Ctrl+B"), self)
        self.shortcut_backup.activated.connect(self._quick_backup)
        
        # Ctrl+P - Print/Export ID card
        self.shortcut_print = QShortcut(QKeySequence("Ctrl+P"), self)
        self.shortcut_print.activated.connect(self._print_selected)
        
        # Ctrl+K - Command palette
        self.shortcut_command = QShortcut(QKeySequence("Ctrl+K"), self)
        self.shortcut_command.activated.connect(self._show_command_palette)
        
        # Delete - Delete selected
        self.shortcut_delete = QShortcut(QKeySequence("Delete"), self)
        self.shortcut_delete.activated.connect(self._delete_selected)
        
        # Escape - Go back / cancel
        self.shortcut_escape = QShortcut(QKeySequence("Escape"), self)
        self.shortcut_escape.activated.connect(self._shortcut_escape)
        
        # Ctrl+1 - Dashboard
        self.shortcut_dashboard = QShortcut(QKeySequence("Ctrl+1"), self)
        self.shortcut_dashboard.activated.connect(lambda: self._show_page(self.dashboard))
        
        # Ctrl+2 - Employees list
        self.shortcut_employees = QShortcut(QKeySequence("Ctrl+2"), self)
        self.shortcut_employees.activated.connect(lambda: self._show_page(self.employees_page))
        
        # Ctrl+H - Show help/keyboard shortcuts
        self.shortcut_help = QShortcut(QKeySequence("Ctrl+H"), self)
        self.shortcut_help.activated.connect(self._show_keyboard_shortcuts_help)
        
    def _focus_search(self):
        """Focus the search box in employees page"""
        self._show_page(self.employees_page)
        if hasattr(self.employees_page, 'search_edit'):
            self.employees_page.search_edit.setFocus()
            self.employees_page.search_edit.selectAll()
    
    def _force_refresh(self):
        """Force refresh employee data"""
        self.employees = self.db.all_employees()
        self.dashboard.refresh(self.employees)
        self.employees_page.set_data(self.employees)
        show_info_toast(self, "Data refreshed")
    
    def _edit_selected(self):
        """Edit currently selected employee"""
        if self.stack.currentWidget() == self.employees_page:
            selected = self.employees_page.table.selectionModel().selectedRows()
            if len(selected) == 1:
                row = self.employees_page.proxy.mapToSource(selected[0]).row()
                emp = self.employees_page.model.data_list[row]
                self._edit_employee(emp)
    
    def _shortcut_save(self):
        """Context-aware save shortcut"""
        if self.stack.currentWidget() == self.form_scroll_area:
            # In form view, trigger save
            if hasattr(self.emp_form, 'save_btn'):
                self.emp_form.save_btn.click()
    
    def _quick_backup(self):
        """Create quick backup"""
        try:
            backup_path = self.db.backup_database()
            if backup_path:
                show_success_toast(self, f"Backup created: {os.path.basename(backup_path)}")
            else:
                show_error_toast(self, "Backup failed")
        except Exception as e:
            show_error_toast(self, f"Backup error: {e}")
    
    def _print_selected(self):
        """Print/export ID card for selected employee"""
        if self.stack.currentWidget() == self.employees_page:
            selected = self.employees_page.table.selectionModel().selectedRows()
            if len(selected) == 1:
                row = self.employees_page.proxy.mapToSource(selected[0]).row()
                emp = self.employees_page.model.data_list[row]
                self._generate_id_card(emp)
    
    def _shortcut_escape(self):
        """Handle escape key - go back or cancel"""
        if self.stack.currentWidget() == self.form_scroll_area:
            # In form, go back to employees list
            self._show_page(self.employees_page)
        elif self.stack.currentWidget() == self.employees_page:
            # In employees list, clear selection
            self.employees_page.table.clearSelection()
    
    def _show_keyboard_shortcuts_help(self):
        """Show keyboard shortcuts help dialog"""
        from employee_vault.ui.widgets import SmoothAnimatedDialog
        dlg = SmoothAnimatedDialog(self, animation_style="fade")
        dlg.setWindowTitle("⌨️ Keyboard Shortcuts")
        dlg.resize(450, 500)
        
        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel("<h2>⌨️ Keyboard Shortcuts</h2>"))
        
        shortcuts_html = """
        <style>
            table { width: 100%; border-collapse: collapse; }
            th, td { padding: 8px; text-align: left; border-bottom: 1px solid #444; }
            th { background: #333; color: #4CAF50; }
            .key { background: #2196F3; color: white; padding: 2px 8px; border-radius: 4px; font-family: monospace; }
        </style>
        <table>
            <tr><th>Shortcut</th><th>Action</th></tr>
            <tr><td><span class="key">Ctrl+F</span></td><td>Focus search box</td></tr>
            <tr><td><span class="key">Ctrl+N</span></td><td>Add new employee</td></tr>
            <tr><td><span class="key">Ctrl+E</span></td><td>Edit selected employee</td></tr>
            <tr><td><span class="key">Ctrl+S</span></td><td>Save (in form)</td></tr>
            <tr><td><span class="key">Ctrl+R</span> / <span class="key">F5</span></td><td>Refresh data</td></tr>
            <tr><td><span class="key">Ctrl+B</span></td><td>Create backup</td></tr>
            <tr><td><span class="key">Ctrl+P</span></td><td>Print/Export ID card</td></tr>
            <tr><td><span class="key">Ctrl+K</span></td><td>Command palette</td></tr>
            <tr><td><span class="key">Ctrl+1</span></td><td>Go to Dashboard</td></tr>
            <tr><td><span class="key">Ctrl+2</span></td><td>Go to Employees list</td></tr>
            <tr><td><span class="key">Ctrl+H</span></td><td>Show this help</td></tr>
            <tr><td><span class="key">Delete</span></td><td>Delete selected employee(s)</td></tr>
            <tr><td><span class="key">Escape</span></td><td>Go back / Cancel</td></tr>
        </table>
        """
        
        shortcuts_label = QLabel(shortcuts_html)
        shortcuts_label.setWordWrap(True)
        layout.addWidget(shortcuts_label)
        
        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(dlg.close)
        layout.addWidget(close_btn)
        
        dlg.exec()

    def _show_command_palette(self):
        """Show command palette for quick actions (Ctrl+K)"""
        from employee_vault.ui.widgets import SmoothAnimatedDialog
        
        dlg = SmoothAnimatedDialog(self, animation_style="fade")
        dlg.setWindowTitle("Command Palette")
        dlg.setWindowFlags(dlg.windowFlags() | Qt.FramelessWindowHint)
        dlg.resize(500, 400)
        
        # Center on screen
        screen = QApplication.primaryScreen().availableGeometry()
        dlg.move(
            (screen.width() - dlg.width()) // 2,
            (screen.height() - dlg.height()) // 3  # Upper third
        )
        
        layout = QVBoxLayout(dlg)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Search input
        search = NeumorphicGradientLineEdit("🔍 Type a command... (Esc to close)")
        search.setMinimumHeight(65)
        layout.addWidget(search)
        
        # Commands list
        commands = [
            ("➕ Add New Employee", "Ctrl+N", self._show_add_new),
            ("🔍 Search Employees", "Ctrl+F", self._focus_search),
            ("🔄 Refresh Data", "F5", self._force_refresh),
            ("💾 Create Backup", "Ctrl+B", self._quick_backup),
            ("📊 Dashboard", "Ctrl+1", lambda: self._show_page(self.dashboard)),
            ("👥 Employee List", "Ctrl+2", lambda: self._show_page(self.employees_page)),
            ("🎨 Change Theme", "", self._show_theme_selector),
            ("📜 View Audit Log", "", self._show_audit_log),
            ("🔒 Security Audit", "", self._show_security_audit),
            ("📈 Generate Reports", "", self._show_reports),
            ("⌨️ Keyboard Shortcuts", "Ctrl+H", self._show_keyboard_shortcuts_help),
            ("👤 My Account", "", self._show_profile),
            ("ℹ️ About", "", self._show_about),
            ("🚪 Logout", "", self._logout),
        ]
        
        # Add admin commands if user has permission
        if self.user_permissions.get('user_management'):
            commands.insert(-1, ("👥 User Management", "", self._show_user_management))
            commands.insert(-1, ("🌐 Network Config", "", self._show_network_config))
        
        list_widget = QListWidget()
        list_widget.setStyleSheet("""
            QListWidget {
                border: none;
                background: #16213e;
                color: white;
                font-size: 14px;
            }
            QListWidget::item {
                padding: 12px 20px;
                border-bottom: 1px solid #1a1a2e;
            }
            QListWidget::item:hover {
                background: #1a1a2e;
            }
            QListWidget::item:selected {
                background: #2196F3;
                color: white;
            }
        """)
        
        # Populate list
        all_items = []
        for name, shortcut, action in commands:
            item = QListWidgetItem(f"{name}    {shortcut}" if shortcut else name)
            item.setData(Qt.UserRole, action)
            list_widget.addItem(item)
            all_items.append((name.lower(), item))
        
        layout.addWidget(list_widget)
        
        def filter_commands():
            text = search.line_edit.text().lower()
            for name, item in all_items:
                item.setHidden(text not in name)
            # Select first visible item
            for i in range(list_widget.count()):
                if not list_widget.item(i).isHidden():
                    list_widget.setCurrentRow(i)
                    break
        
        def execute_command(item):
            action = item.data(Qt.UserRole)
            dlg.close()
            if callable(action):
                action()
        
        def on_key_press(event):
            if event.key() == Qt.Key_Escape:
                dlg.close()
            elif event.key() == Qt.Key_Return:
                current = list_widget.currentItem()
                if current and not current.isHidden():
                    execute_command(current)
            elif event.key() == Qt.Key_Down:
                # Move selection down
                for i in range(list_widget.currentRow() + 1, list_widget.count()):
                    if not list_widget.item(i).isHidden():
                        list_widget.setCurrentRow(i)
                        break
            elif event.key() == Qt.Key_Up:
                # Move selection up
                for i in range(list_widget.currentRow() - 1, -1, -1):
                    if not list_widget.item(i).isHidden():
                        list_widget.setCurrentRow(i)
                        break
            else:
                QLineEdit.keyPressEvent(search, event)
        
        search.keyPressEvent = on_key_press
        search.line_edit.textChanged.connect(filter_commands)
        list_widget.itemDoubleClicked.connect(execute_command)
        
        # Focus search and select first item
        search.setFocus()
        list_widget.setCurrentRow(0)
        
        dlg.exec()

    def _show_page(self, w: QWidget): self.stack.setCurrentWidget(w)
    def _show_add_new(self): self.emp_form.edit_employee(None); self._show_page(self.form_scroll_area)

    def _view_employee(self, emp):
        # v4.4.1: Animated dialog for employee details
        from employee_vault.ui.widgets import SmoothAnimatedDialog
        dlg=SmoothAnimatedDialog(self, animation_style="slide"); dlg.setWindowTitle(f"Employee Details — {emp.get('name','?')}"); dlg.resize(720,600)

        # Main dialog layout
        main_layout = QVBoxLayout(dlg)

        # Create scroll area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)

        # Content widget inside scroll area
        content_widget = QWidget()
        v = QVBoxLayout(content_widget)

        # Header with photo
        header=QHBoxLayout(); photo=os.path.join(PHOTOS_DIR,f"{emp.get('emp_id','')}.png")
        img=QLabel(); img.setText("📷") if not os.path.exists(photo) else img.setPixmap(QPixmap(photo).scaled(60,60,Qt.KeepAspectRatio,Qt.SmoothTransformation))
        header.addWidget(img); info=QVBoxLayout(); status="🟢 Active" if not emp.get("resign_date") else "🔴 Resigned"; info.addWidget(QLabel(f"<b>{emp.get('name','?')}</b> — {status}")); info.addWidget(QLabel(f"{emp.get('emp_id','?')} • {emp.get('position','N/A')}")); header.addLayout(info,1); v.addLayout(header)

        # Employee fields
        fields=[("👤 Full Name", emp.get('name','?')),
                ("🆔 Employee ID", emp.get('emp_id','?')),
                ("🔢 SSS #", emp.get('sss_number','—') or "—"),
                ("💳 TIN", emp.get('tin_number','—') or "—"),
                ("🏦 Pag-IBIG #", emp.get('pagibig_number','—') or "—"),
                ("🏥 PhilHealth #", emp.get('philhealth_number','—') or "—"),
                ("📧 Email", emp.get('email','N/A')),
                ("📱 Phone", emp.get('phone','N/A')),
                ("🏢 Department", emp.get('department','N/A')),
                ("🏢 Agency", emp.get('agency','—') or "—"),
                ("💼 Position", emp.get('position','N/A')),
                ("📅 Hire Date", emp.get('hire_date','N/A')),
                ("📅 Resign Date", emp.get('resign_date','—') or "—"),
                ("💰 Salary/Day", f"{emp.get('salary',0):,.2f}")]

        # Add contract information if available
        if emp.get('contract_start_date') or emp.get('contract_expiry'):
            fields.append(("📄 Contract Start", emp.get('contract_start_date','—') or "—"))
            if emp.get('contract_months'):
                months = emp.get('contract_months')
                if months >= 12:
                    years = months // 12
                    remaining = months % 12
                    duration = f"{years}y {remaining}m" if remaining else f"{years} year(s)"
                else:
                    duration = f"{months} month(s)"
                fields.append(("⏱️ Contract Duration", duration))
            fields.append(("📄 Contract Expiry", emp.get('contract_expiry','—') or "—"))

        for lab,val in fields:
            r=QHBoxLayout(); l=QLabel(f"<b>{lab}</b>"); l.setFixedWidth(170); r.addWidget(l); r.addWidget(QLabel(str(val)),1); v.addLayout(r)

        # Emergency contact section
        if emp.get('emergency_contact_name') or emp.get('emergency_contact_phone'):
            v.addWidget(QLabel("<b style='color:#ff9966;'>🚨 Emergency Contact</b>"))
            emergency_fields = []
            if emp.get('emergency_contact_name'):
                emergency_fields.append(("👤 Name", emp.get('emergency_contact_name')))
            if emp.get('emergency_contact_phone'):
                emergency_fields.append(("📱 Phone", emp.get('emergency_contact_phone')))
            for lab, val in emergency_fields:
                r=QHBoxLayout(); l=QLabel(f"<b>{lab}</b>"); l.setFixedWidth(170); r.addWidget(l); r.addWidget(QLabel(str(val)),1); v.addLayout(r)

        # Contract status
        d=contract_days_left(emp)
        if d is not None:
            if d<0: v.addWidget(QLabel(f"<b style='color:#ff6b6b'>Contract expired {-d} day(s) ago</b>"))
            elif d==0: v.addWidget(QLabel("<b style='color:#ff6b6b'>Contract expires today</b>"))
            elif d<=ALERT_DAYS: v.addWidget(QLabel(f"<b style='color:#ffcc66'>Contract expires in {d} day(s)</b>"))
            else: v.addWidget(QLabel(f"<b style='color:#9ad17a'>Contract valid ({d} day(s) left)</b>"))

        # Notes section
        if emp.get('notes'):
            v.addWidget(QLabel("<b>📝 Notes</b>"))
            notes_label = QLabel(emp.get('notes', ''))
            notes_label.setWordWrap(True)
            notes_label.setStyleSheet("padding: 10px; background: #1c1c1c; border-radius: 8px;")
            v.addWidget(notes_label)

        # Attached files - IMPROVED
        v.addWidget(QLabel("<b>📎 Attached Files</b>"))
        folder=os.path.join(FILES_DIR, emp.get('emp_id',''))
        if os.path.exists(folder):
            files=os.listdir(folder)
            if files:
                files_container = QWidget()
                files_container.setStyleSheet("background: #1c1c1c; border-radius: 8px; padding: 10px;")
                files_layout = QVBoxLayout(files_container)
                files_layout.setContentsMargins(5, 5, 5, 5)
                files_layout.setSpacing(5)

                for f in files:
                    pth=os.path.join(folder,f)
                    btn=ModernAnimatedButton(f"📄 {f}")
                    btn.setFlat(True)
                    btn.setStyleSheet("""
                        QPushButton {
                            text-align: left;
                            padding: 8px;
                            background: #2a2a2a;
                            border-radius: 4px;
                        }
                        QPushButton:hover {
                            background: #3a3a3a;
                        }
                    """)
                    from PySide6.QtCore import QUrl
                    btn.clicked.connect(lambda _=False, pp=pth: QDesktopServices.openUrl(QUrl.fromLocalFile(pp)))
                    files_layout.addWidget(btn)

                v.addWidget(files_container)
            else:
                v.addWidget(QLabel("<i style='color:#888;'>No files attached</i>"))
        else:
            v.addWidget(QLabel("<i style='color:#888;'>No files attached</i>"))

        # Add stretch to push content to top
        v.addStretch()

        # Set content widget to scroll area
        scroll.setWidget(content_widget)

        # Add scroll area to main layout
        main_layout.addWidget(scroll)

        # Close button outside scroll area (always visible)
        button_layout = QHBoxLayout()

        # Add View History button for admins
        if self.user_row['role'] == 'admin':
            history_btn = ModernAnimatedButton("📜 View History")
            history_btn.clicked.connect(lambda: self._show_employee_history(emp))
            button_layout.addWidget(history_btn)

        close_btn = ModernAnimatedButton("Close (ESC)")
        close_btn.clicked.connect(dlg.close)
        button_layout.addWidget(close_btn)

        main_layout.addLayout(button_layout)

        dlg.exec()

    def _show_employee_history(self, emp):
        """Show audit history for a specific employee"""
        # v4.4.1: Animated dialog for employee history
        from employee_vault.ui.widgets import AnimatedDialogBase
        dlg = AnimatedDialogBase(self, animation_style="fade")
        dlg.setWindowTitle(f"📜 History - {emp.get('name', '?')}")
        dlg.resize(900, 500)

        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel(f"<h2>📜 Activity History for {emp.get('name', '?')}</h2>"))
        layout.addWidget(QLabel(f"Employee ID: <b>{emp.get('emp_id', '?')}</b>"))

        # Table for history
        from PySide6.QtWidgets import QTableWidget, QTableWidgetItem, QHeaderView
        table = QTableWidget()
        table.setColumnCount(6)
        table.setHorizontalHeaderLabels(["Timestamp", "User", "Action", "Old Value", "New Value", "Details"])
        table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        table.horizontalHeader().setStretchLastSection(True)
        table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        table.setSelectionBehavior(QAbstractItemView.SelectRows)

        # Load history
        history = self.db.get_employee_history(emp.get('emp_id'))

        table.setRowCount(len(history))
        for i, log in enumerate(history):
            table.setItem(i, 0, QTableWidgetItem(log.get('timestamp', '')))
            table.setItem(i, 1, QTableWidgetItem(log.get('username', '')))

            # Color code actions
            action_item = QTableWidgetItem(log.get('action', ''))
            action_text = log.get('action', '')
            if action_text == 'INSERT':
                action_item.setForeground(QColor('#2196F3'))  # Blue
            elif action_text == 'UPDATE':
                action_item.setForeground(QColor('#FF9800'))  # Orange
            elif action_text == 'DELETE':
                action_item.setForeground(QColor('#F44336'))  # Red
            table.setItem(i, 2, action_item)

            table.setItem(i, 3, QTableWidgetItem(log.get('old_value', '') or ''))
            table.setItem(i, 4, QTableWidgetItem(log.get('new_value', '') or ''))

            details_text = log.get('details', '') or ''
            details_item = QTableWidgetItem(details_text[:100] + ('...' if len(details_text) > 100 else ''))
            details_item.setToolTip(details_text)
            table.setItem(i, 5, details_item)

        table.resizeColumnsToContents()
        layout.addWidget(table)

        # Stats
        layout.addWidget(QLabel(f"Total changes: <b>{len(history)}</b>"))

        # Close button
        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(dlg.close)
        layout.addWidget(close_btn)

        dlg.exec()

    def _edit_employee(self, emp): self.emp_form.edit_employee(emp); self._show_page(self.form_scroll_area)

    def _delete_selected(self, rows):
        # Only admins can delete employees
        if not self.user_permissions.get('delete_employee'):
            show_warning_toast(
                self, "Only administrators can delete employees.\n\n"
                "Please contact your administrator if you need to delete employee records."
            )
            return

        if not rows:
            return

        preview = "<br>".join(f"{r['emp_id']} — {r['name']}" for r in rows[:10])
        if len(rows) > 10:
            preview += "<br>…"

        # Ask for archive reason
        # Build the prompt outside of the f-string to avoid backslash
        # confusion within the f-string expression.  Newlines are
        # escaped here so that Python does not interpret them inside an
        # f-string expression.  Using join() would also work but this
        # approach is clearer.
        prompt_lines = [
            f"Archive {len(rows)} employee(s)?",
            "",
            preview.replace('<br>', '\n'),
            "",
            "These employees will be moved to the archive and can be restored later.",
            "",
            "Reason for archiving (optional):"
        ]
        prompt_text = "\n".join(prompt_lines)
        reason, ok = QInputDialog.getText(self, "Archive Reason", prompt_text)

        if not ok:
            return

        # Archive employees instead of deleting
        for row in rows:
            emp_id = row["emp_id"]
            self.db.archive_employee(emp_id, self.current_user, reason or "No reason specified")

            # Move file folders to archive
            src_folder = os.path.join(FILES_DIR, emp_id)
            if os.path.isdir(src_folder):
                archive_folder = os.path.join(FILES_DIR, "_archived", emp_id)
                os.makedirs(os.path.dirname(archive_folder), exist_ok=True)
                try:
                    shutil.move(src_folder, archive_folder)
                except (OSError, shutil.Error):
                    # File may already be moved or locked
                    pass

            # Move photos to archive
            src_photo = os.path.join(PHOTOS_DIR, f"{emp_id}.png")
            if os.path.exists(src_photo):
                archive_photo = os.path.join(PHOTOS_DIR, "_archived", f"{emp_id}.png")
                os.makedirs(os.path.dirname(archive_photo), exist_ok=True)
                try:
                    shutil.move(src_photo, archive_photo)
                except (OSError, shutil.Error):
                    # Photo may already be moved or locked
                    pass

        show_success_toast(
            self, f"{len(rows)} employee(s) have been archived.\n\n"
            "You can restore them from the Archive Manager."
        )

        self._refresh_all()

    def _on_form_saved(self, cancel_only=False, switch_page=True):
        if not cancel_only: self._refresh_all()
        if switch_page: self._show_page(self.employees_page)

    def _refresh_all(self):
        self.employees=self.db.all_employees()
        self.dashboard.refresh(self.employees)
        self.employees_page.set_data(self.employees)

    def _show_enhanced_search(self):
        """Show enhanced search dialog"""
        # v4.4.1: Animated dialog for enhanced search
        from employee_vault.ui.widgets import AnimatedDialogBase
        dialog = AnimatedDialogBase(self, animation_style="fade")
        dialog.setWindowTitle("🔍 Enhanced Search")
        dialog.setMinimumWidth(500)
        layout = QFormLayout(dialog)

        # Search criteria
        name_input = NeumorphicGradientLineEdit("Employee name")
        name_input.setMinimumHeight(70)
        dept_combo = NeumorphicGradientComboBox("Select Department")
        dept_combo.setMinimumHeight(70)
        dept_combo.addItems(["", "Office", "Store - Sm Novaliches", "Store - Sm San Fernando"])
        position_input = NeumorphicGradientLineEdit("Position/Role")
        position_input.setMinimumHeight(70)
        status_combo = NeumorphicGradientComboBox("Select Status")
        status_combo.setMinimumHeight(70)
        status_combo.addItems(["All", "Active", "Resigned"])

        layout.addRow("Name:", name_input)
        layout.addRow("Department:", dept_combo)
        layout.addRow("Position:", position_input)
        layout.addRow("Status:", status_combo)

        # Buttons
        btn_layout = QHBoxLayout()
        search_btn = ModernAnimatedButton("Search")
        clear_btn = ModernAnimatedButton("Clear")
        close_btn = ModernAnimatedButton("Close")

        def do_search():
            # Filter employees based on criteria
            results = self.employees.copy()

            if name_input.line_edit.text():
                results = [e for e in results if name_input.line_edit.text().lower() in e['name'].lower()]
            if dept_combo.combo_box.currentText():
                results = [e for e in results if e.get('department') == dept_combo.combo_box.currentText()]
            if position_input.line_edit.text():
                results = [e for e in results if position_input.line_edit.text().lower() in e.get('position', '').lower()]
            if status_combo.combo_box.currentText() == "Active":
                results = [e for e in results if not e.get('resign_date')]
            elif status_combo.combo_box.currentText() == "Resigned":
                results = [e for e in results if e.get('resign_date')]

            # Update table
            self.employees_page.refresh(results)
            dialog.accept()
            self._show_page(self.employees_page)

        search_btn.clicked.connect(do_search)
        clear_btn.clicked.connect(lambda: [name_input.line_edit.clear(), position_input.line_edit.clear()])
        close_btn.clicked.connect(dialog.reject)

        btn_layout.addWidget(search_btn)
        btn_layout.addWidget(clear_btn)
        btn_layout.addWidget(close_btn)
        layout.addRow(btn_layout)

        dialog.exec()


    def _show_id_card_generator(self):
        """Open ID card generator"""
        dlg = IDCardGeneratorBackenderator(self, self.db)
        dlg.exec()

    def _show_print_system(self):
        """Open print system dialog"""
        dlg = PrintSystemDialog(self, self.db, self.employees)
        dlg.exec()

    def _show_bulk_operations(self):
        """Open bulk operations dialog"""
        dlg = BulkOperationsDialog(self, self.db, self.employees)
        if dlg.exec() == QDialog.Accepted:
            self._refresh_all()


    def _export_data(self):
        """Export employee data to JSON or Excel"""
        formats = "JSON (*.json);;Excel (*.xlsx)"
        path, selected_filter = QFileDialog.getSaveFileName(
            self, "Export Employee Data", "employees_export.json", formats
        )
        if not path:
            return

        try:
            if path.endswith('.xlsx'):
                self._export_to_excel(path)
            else:
                # JSON export
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(self.employees, f, indent=2)
            show_success_toast(self, f"Data exported successfully to:\n{path}")
        except Exception as e:
            logging.error(f"Export error: {e}")
            show_error_toast(self, f"Failed to export data:\n{str(e)}")

    def _export_to_excel(self, filepath):
        """Export employee data to Excel with formatting"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            show_warning_toast(
                self, "Excel export requires 'openpyxl' library.\n\n"
                "Install it with:\npip install openpyxl"
            )
            return

        wb = Workbook()
        ws = wb.active
        ws.title = "Employees"

        # Headers
        headers = ["Employee ID", "Name", "SSS #", "Email", "Phone", "Department",
                  "Position", "Agency", "Hire Date", "Resign Date", "Salary/Day",
                  "Contract Start", "Contract Duration", "Contract Expiry", "Status",
                  "Emergency Contact", "Emergency Phone"]

        # Style headers
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_border

        # Data rows
        for row_num, emp in enumerate(self.employees, 2):
            status = "Active" if not emp.get('resign_date') else "Resigned"

            # Calculate contract duration
            contract_duration = ""
            if emp.get('contract_months'):
                months = emp.get('contract_months')
                if months >= 12:
                    years = months // 12
                    remaining = months % 12
                    contract_duration = f"{years}y {remaining}m" if remaining else f"{years} year(s)"
                else:
                    contract_duration = f"{months} month(s)"

            row_data = [
                emp.get('emp_id', ''),
                emp.get('name', ''),
                emp.get('sss_number', ''),
                emp.get('email', ''),
                emp.get('phone', ''),
                emp.get('department', ''),
                emp.get('position', ''),
                emp.get('agency', ''),
                emp.get('hire_date', ''),
                emp.get('resign_date', ''),
                emp.get('salary', 0),
                emp.get('contract_start_date', ''),
                contract_duration,
                emp.get('contract_expiry', ''),
                status,
                emp.get('emergency_contact_name', ''),
                emp.get('emergency_contact_phone', '')
            ]

            for col_num, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_num, column=col_num, value=value)
                cell.border = thin_border
                if col_num == 11:  # Salary column
                    cell.number_format = '#,##0.00'
                # Color-code status
                if col_num == 15:  # Status column
                    if status == "Active":
                        cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
                    else:
                        cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")

        # Adjust column widths
        for col_num in range(1, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col_num)].width = 15

        wb.save(filepath)
        logging.info(f"Exported {len(self.employees)} employees to Excel: {filepath}")

    def _import_data(self):
        """Import employees from Excel or CSV file"""
        from employee_vault.ui.widgets import SmoothAnimatedDialog
        
        filepath, _ = QFileDialog.getOpenFileName(
            self, 
            "Import Employee Data", 
            "",
            "Spreadsheet Files (*.xlsx *.csv);;Excel Files (*.xlsx);;CSV Files (*.csv)"
        )
        
        if not filepath:
            return
        
        try:
            # Parse file based on extension
            if filepath.endswith('.csv'):
                employees = self._parse_csv_import(filepath)
            else:
                employees = self._parse_excel_import(filepath)
            
            if not employees:
                show_warning_toast(self, "No valid employee data found in file")
                return
            
            # Show preview dialog
            self._show_import_preview(employees, filepath)
            
        except Exception as e:
            logging.error(f"Import error: {e}")
            show_error_toast(self, f"Failed to read file:\n{str(e)}")

    def _parse_csv_import(self, filepath):
        """Parse CSV file for import"""
        import csv
        employees = []
        
        with open(filepath, 'r', encoding='utf-8-sig') as f:
            reader = csv.DictReader(f)
            for row in reader:
                emp = self._normalize_import_row(row)
                if emp:
                    employees.append(emp)
        
        return employees

    def _parse_excel_import(self, filepath):
        """Parse Excel file for import"""
        try:
            from openpyxl import load_workbook
        except ImportError:
            show_error_toast(self, "openpyxl required for Excel import. Install with: pip install openpyxl")
            return []
        
        employees = []
        wb = load_workbook(filepath, read_only=True)
        ws = wb.active
        
        # Get headers from first row
        headers = [cell.value for cell in ws[1]]
        
        for row in ws.iter_rows(min_row=2, values_only=True):
            row_dict = dict(zip(headers, row))
            emp = self._normalize_import_row(row_dict)
            if emp:
                employees.append(emp)
        
        wb.close()
        return employees

    def _normalize_import_row(self, row):
        """Normalize imported row to employee dict"""
        # Map common column names to our field names
        field_map = {
            'employee id': 'emp_id', 'emp id': 'emp_id', 'id': 'emp_id', 'emp_id': 'emp_id',
            'name': 'name', 'full name': 'name', 'employee name': 'name',
            'email': 'email', 'e-mail': 'email',
            'phone': 'phone', 'phone number': 'phone', 'mobile': 'phone',
            'department': 'department', 'dept': 'department',
            'position': 'position', 'job title': 'position', 'title': 'position',
            'hire date': 'hire_date', 'start date': 'hire_date', 'date hired': 'hire_date',
            'resign date': 'resign_date', 'end date': 'resign_date', 'termination date': 'resign_date',
            'salary': 'salary', 'basic salary': 'salary',
            'agency': 'agency',
            'sss': 'sss_number', 'sss number': 'sss_number', 'sss_number': 'sss_number',
            'contract expiry': 'contract_expiry', 'contract_expiry': 'contract_expiry',
            'notes': 'notes', 'remarks': 'notes',
        }
        
        emp = {}
        for key, value in row.items():
            if key and value:
                normalized_key = str(key).lower().strip()
                if normalized_key in field_map:
                    emp[field_map[normalized_key]] = str(value).strip()
        
        # Must have at least name to be valid
        if not emp.get('name'):
            return None
        
        return emp

    def _show_import_preview(self, employees, filepath):
        """Show preview of data to import with validation"""
        from employee_vault.ui.widgets import SmoothAnimatedDialog
        
        dlg = SmoothAnimatedDialog(self, animation_style="slide")
        dlg.setWindowTitle("📥 Import Preview")
        dlg.resize(900, 600)
        
        layout = QVBoxLayout(dlg)
        
        # Header
        layout.addWidget(QLabel(f"<h2>📥 Import Preview</h2>"))
        layout.addWidget(QLabel(f"File: {os.path.basename(filepath)} | Found: {len(employees)} employees"))
        
        # Validation summary
        existing_ids = {e.get('emp_id') for e in self.employees}
        new_count = sum(1 for e in employees if e.get('emp_id') not in existing_ids)
        update_count = len(employees) - new_count
        
        summary = QLabel(f"""
            <b>Import Summary:</b><br>
            🆕 New employees: <span style="color: #4CAF50;">{new_count}</span><br>
            🔄 Updates (matching ID): <span style="color: #FF9800;">{update_count}</span>
        """)
        layout.addWidget(summary)
        
        # Preview table
        table = QTableWidget()
        table.setColumnCount(7)
        table.setHorizontalHeaderLabels(["Status", "Emp ID", "Name", "Department", "Position", "Hire Date", "Email"])
        table.setRowCount(len(employees))
        table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        table.setAlternatingRowColors(True)
        
        for i, emp in enumerate(employees):
            is_update = emp.get('emp_id') in existing_ids
            status = "🔄 Update" if is_update else "🆕 New"
            
            status_item = QTableWidgetItem(status)
            status_item.setForeground(QColor('#FF9800' if is_update else '#4CAF50'))
            table.setItem(i, 0, status_item)
            
            table.setItem(i, 1, QTableWidgetItem(emp.get('emp_id', '—')))
            table.setItem(i, 2, QTableWidgetItem(emp.get('name', '—')))
            table.setItem(i, 3, QTableWidgetItem(emp.get('department', '—')))
            table.setItem(i, 4, QTableWidgetItem(emp.get('position', '—')))
            table.setItem(i, 5, QTableWidgetItem(emp.get('hire_date', '—')))
            table.setItem(i, 6, QTableWidgetItem(emp.get('email', '—')))
        
        table.resizeColumnsToContents()
        layout.addWidget(table)
        
        # Options
        options_layout = QHBoxLayout()
        
        skip_existing = QCheckBox("Skip existing employees (don't update)")
        skip_existing.setChecked(False)
        options_layout.addWidget(skip_existing)
        
        generate_ids = QCheckBox("Auto-generate missing Employee IDs")
        generate_ids.setChecked(True)
        options_layout.addWidget(generate_ids)
        
        layout.addLayout(options_layout)
        
        # Buttons
        btn_layout = QHBoxLayout()
        
        import_btn = ModernAnimatedButton("✅ Import All")
        import_btn.setStyleSheet("background: #4CAF50;")
        
        cancel_btn = ModernAnimatedButton("Cancel")
        
        def do_import():
            try:
                imported = 0
                updated = 0
                skipped = 0
                
                for emp in employees:
                    emp_id = emp.get('emp_id')
                    
                    # Generate ID if missing
                    if not emp_id and generate_ids.isChecked():
                        dept = emp.get('department', 'GEN')
                        prefix = self._prefix_from_department(dept) if hasattr(self, '_prefix_from_department') else 'EMP'
                        seq = self.db.next_sequence()
                        emp_id = f"{prefix}-{seq:03d}"
                        emp['emp_id'] = emp_id
                    
                    if not emp_id:
                        skipped += 1
                        continue
                    
                    # Check if exists
                    if emp_id in existing_ids:
                        if skip_existing.isChecked():
                            skipped += 1
                            continue
                        # Update existing
                        self.db.update_employee(emp_id, emp, self.current_user)
                        updated += 1
                    else:
                        # Add new with required fields
                        emp.setdefault('hire_date', datetime.now().strftime('%m-%d-%Y'))
                        self.db.add_employee(emp, self.current_user)
                        imported += 1
                
                dlg.close()
                
                # Refresh data
                self.employees = self.db.all_employees()
                self.dashboard.refresh(self.employees)
                self.employees_page.set_data(self.employees)
                
                show_success_toast(
                    self, 
                    f"Import complete!\n\n🆕 Added: {imported}\n🔄 Updated: {updated}\n⏭️ Skipped: {skipped}"
                )
                
                # Log security event
                self.db.log_security_event(
                    event_type="DATA_IMPORT",
                    username=self.current_user,
                    details=f"Imported {imported} new, updated {updated} from {os.path.basename(filepath)}",
                    severity="INFO"
                )
                
            except Exception as e:
                logging.error(f"Import failed: {e}")
                show_error_toast(self, f"Import failed:\n{str(e)}")
        
        import_btn.clicked.connect(do_import)
        cancel_btn.clicked.connect(dlg.close)
        
        btn_layout.addWidget(import_btn)
        btn_layout.addStretch()
        btn_layout.addWidget(cancel_btn)
        
        layout.addLayout(btn_layout)
        
        dlg.exec()

    def _backup_data(self):
        """Manual backup with timestamp"""
        folder = QFileDialog.getExistingDirectory(self, "Choose Backup Folder")
        if not folder:
            return

        try:
            # Create timestamped backup folder
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_folder = os.path.join(folder, f"employee_vault_backup_{timestamp}")
            os.makedirs(backup_folder, exist_ok=True)

            # Copy database files
            files_copied = 0
            for fn in [DB_FILE, DB_FILE + "-wal", DB_FILE + "-shm"]:
                if os.path.exists(fn):
                    dest = os.path.join(backup_folder, os.path.basename(fn))
                    shutil.copy2(fn, dest)
                    files_copied += 1

            # Also backup the employee files and photos
            if os.path.exists(FILES_DIR):
                shutil.copytree(FILES_DIR, os.path.join(backup_folder, FILES_DIR))
            if os.path.exists(PHOTOS_DIR):
                shutil.copytree(PHOTOS_DIR, os.path.join(backup_folder, PHOTOS_DIR))

            # Create backup info file
            info = {
                "backup_date": datetime.now().isoformat(),
                "total_employees": len(self.employees),
                "database_files": files_copied,
                "backed_up_by": self.current_user
            }
            with open(os.path.join(backup_folder, "backup_info.json"), "w") as f:
                json.dump(info, f, indent=2)

            logging.info(f"Backup created: {backup_folder}")
            show_success_toast(
                self, f"Backup created successfully!\n\n"
                f"Location: {backup_folder}\n"
                f"Files backed up: {files_copied} database files + attachments"
            )
        except Exception as e:
            logging.error(f"Backup error: {e}")
            show_error_toast(self, f"Failed to create backup:\n{str(e)}")

    def _setup_auto_backup(self):
        """Setup automated daily backup"""
        # This would be called from settings in a full implementation
        pass

    def _show_network_config(self):
        """Show network configuration"""
        # v4.4.1: Animated dialog for network config
        from employee_vault.ui.widgets import AnimatedDialogBase
        dialog = AnimatedDialogBase(self, animation_style="fade")
        dialog.setWindowTitle("🌐 Network Configuration")
        layout = QVBoxLayout(dialog)

        # Share path setting
        layout.addWidget(QLabel("<b>Network Share Path:</b>"))
        share_input = NeumorphicGradientLineEdit("Network UNC Path")
        share_input.setMinimumHeight(70)
        share_input.line_edit.setText(EXPECTED_UNC_PREFIX)
        share_input.line_edit.setReadOnly(True)
        layout.addWidget(share_input)

        # Database location
        layout.addWidget(QLabel("<b>Database Location:</b>"))
        db_input = NeumorphicGradientLineEdit("Database File Path")
        db_input.setMinimumHeight(70)
        db_input.line_edit.setText(str(Path(DB_FILE).resolve()))
        db_input.line_edit.setReadOnly(True)
        layout.addWidget(db_input)

        # Info
        layout.addWidget(QLabel("<i>Network settings are configured in code.</i>"))

        # Close button
        btn = ModernAnimatedButton("Close")
        btn.clicked.connect(dialog.accept)
        layout.addWidget(btn)

        dialog.exec()

    def _test_network_paths(self, db_path, files_path, photos_path):
        """Test if network paths are accessible"""
        results = []

        # Test database
        if os.path.exists(db_path) and os.access(db_path, os.R_OK | os.W_OK):
            results.append("✅ Database path is accessible")
        else:
            results.append("❌ Database path is not accessible")

        # Test files directory
        if os.path.exists(files_path) and os.access(files_path, os.R_OK | os.W_OK):
            results.append("✅ Files directory is accessible")
        else:
            results.append("❌ Files directory is not accessible")

        # Test photos directory
        if os.path.exists(photos_path) and os.access(photos_path, os.R_OK | os.W_OK):
            results.append("✅ Photos directory is accessible")
        else:
            results.append("❌ Photos directory is not accessible")

        show_success_toast(self, "\n".join(results))

    def _show_email_settings(self):
        """Show email notification configuration dialog"""
        from employee_vault.ui.widgets import SmoothAnimatedDialog
        from employee_vault.email_notifications import get_email_config_from_db, save_email_config_to_db, EmailNotifier
        
        dlg = SmoothAnimatedDialog(self, animation_style="fade")
        dlg.setWindowTitle("📧 Email Notifications")
        dlg.resize(500, 550)
        
        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel("<h2>📧 Email Notification Settings</h2>"))
        layout.addWidget(QLabel("Configure email alerts for contract expiry and security events"))
        
        # Load current config
        config = get_email_config_from_db(self.db)
        
        # Enable toggle
        enable_check = QCheckBox("Enable Email Notifications")
        enable_check.setChecked(config.get('enabled', False))
        layout.addWidget(enable_check)
        
        # SMTP settings group
        smtp_group = QGroupBox("SMTP Server Settings")
        smtp_layout = QFormLayout(smtp_group)

        server_input = NeumorphicGradientLineEdit("smtp.gmail.com")
        server_input.setMinimumHeight(70)
        server_input.line_edit.setText(config.get('smtp_server', 'smtp.gmail.com'))
        smtp_layout.addRow("SMTP Server:", server_input)

        port_input = NeumorphicGradientSpinBox("Port Number")
        port_input.setMinimumHeight(70)
        port_input.spin_box.setRange(1, 65535)
        port_input.spin_box.setValue(config.get('smtp_port', 587))
        smtp_layout.addRow("SMTP Port:", port_input)

        username_input = NeumorphicGradientLineEdit("your-email@gmail.com")
        username_input.setMinimumHeight(70)
        username_input.line_edit.setText(config.get('username', ''))
        smtp_layout.addRow("Username:", username_input)

        password_input = NeumorphicGradientPasswordInput("App password (not your regular password)")
        password_input.setMinimumHeight(70)
        password_input.line_edit.setText(config.get('password', ''))
        smtp_layout.addRow("Password:", password_input)
        
        layout.addWidget(smtp_group)
        
        # Notification settings group
        notif_group = QGroupBox("Notification Recipients")
        notif_layout = QFormLayout(notif_group)

        from_input = NeumorphicGradientLineEdit("noreply@yourcompany.com")
        from_input.setMinimumHeight(70)
        from_input.line_edit.setText(config.get('from_email', ''))
        notif_layout.addRow("From Email:", from_input)

        admin_input = NeumorphicGradientLineEdit("admin@yourcompany.com")
        admin_input.setMinimumHeight(70)
        admin_input.line_edit.setText(config.get('admin_email', ''))
        notif_layout.addRow("Admin Email:", admin_input)
        
        layout.addWidget(notif_group)
        
        # Info
        info = QLabel("""
            <b>Gmail Users:</b> Use an App Password instead of your regular password.<br>
            Go to <a href="https://myaccount.google.com/apppasswords">Google App Passwords</a> to create one.<br><br>
            <b>Notifications will be sent for:</b><br>
            • Contract expiry alerts (daily summary)<br>
            • Critical security events<br>
        """)
        info.setWordWrap(True)
        info.setOpenExternalLinks(True)
        info.setStyleSheet("color: #888; font-size: 11px;")
        layout.addWidget(info)
        
        # Buttons
        btn_layout = QHBoxLayout()
        
        test_btn = ModernAnimatedButton("🔧 Test Connection")
        save_btn = ModernAnimatedButton("💾 Save Settings")
        cancel_btn = ModernAnimatedButton("Cancel")
        
        def test_connection():
            notifier = EmailNotifier({
                'smtp_server': server_input.line_edit.text(),
                'smtp_port': port_input.spin_box.value(),
                'username': username_input.line_edit.text(),
                'password': password_input.line_edit.text(),
                'from_email': from_input.line_edit.text()
            })

            if not notifier.is_configured():
                show_warning_toast(dlg, "Please fill in all SMTP settings")
                return

            test_email = admin_input.line_edit.text() or from_input.line_edit.text()
            if not test_email:
                show_warning_toast(dlg, "Please enter an admin email to receive the test")
                return
            
            try:
                success = notifier.send_email(
                    test_email,
                    "🧪 Employee Vault - Test Email",
                    "<h1>✅ Email Configuration Working!</h1><p>Your email notifications are properly configured.</p>"
                )
                if success:
                    show_success_toast(dlg, f"Test email sent to {test_email}!")
                else:
                    show_error_toast(dlg, "Failed to send test email. Check your settings.")
            except Exception as e:
                show_error_toast(dlg, f"Connection failed:\n{str(e)}")
        
        def save_settings():
            new_config = {
                'smtp_server': server_input.line_edit.text(),
                'smtp_port': port_input.spin_box.value(),
                'username': username_input.line_edit.text(),
                'password': password_input.line_edit.text(),
                'from_email': from_input.line_edit.text(),
                'admin_email': admin_input.line_edit.text(),
                'enabled': enable_check.isChecked()
            }
            
            if save_email_config_to_db(self.db, new_config):
                show_success_toast(dlg, "Email settings saved successfully!")
                dlg.close()
            else:
                show_error_toast(dlg, "Failed to save settings")
        
        test_btn.clicked.connect(test_connection)
        save_btn.clicked.connect(save_settings)
        cancel_btn.clicked.connect(dlg.close)
        
        btn_layout.addWidget(test_btn)
        btn_layout.addWidget(save_btn)
        btn_layout.addStretch()
        btn_layout.addWidget(cancel_btn)
        
        layout.addLayout(btn_layout)
        
        dlg.exec()

    def _show_scheduled_backup(self):
        """Show scheduled backup configuration dialog"""
        # v4.4.1: Animated dialog for scheduled backup
        from employee_vault.ui.widgets import AnimatedDialogBase
        dlg = AnimatedDialogBase(self, animation_style="fade")
        dlg.setWindowTitle("⏰ Scheduled Backup")
        dlg.resize(650, 600)

        wheel_guard = WheelGuard(dlg)

        try:
            config_str = self.db.get_setting('backup_config', '{}')
            config = json.loads(config_str)
        except Exception:
            config = {}  # Use defaults if loading fails

        # This is the MAIN layout for the dialog (non-scrolling)
        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel("<h2>⏰ Scheduled Backup Configuration</h2>"))
        layout.addWidget(QLabel("Automatically backup your data on a schedule"))

        # --- Scroll Area Setup (as before) ---
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_content_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_content_widget)


        # --- Enable/Disable ---
        enable_box = QGroupBox("Status")
        enable_layout = QVBoxLayout(enable_box)

        enable_check = QCheckBox("Enable Scheduled Backups")
        # Load saved value
        enable_check.setChecked(config.get('enabled', False))
        enable_layout.addWidget(enable_check)

        status_label = QLabel("Status: ❌ Disabled")
        status_label.setStyleSheet("color: #ff6b6b; padding: 5px;")
        enable_layout.addWidget(status_label)

        def toggle_status():
            if enable_check.isChecked():
                status_label.setText("Status: ✅ Enabled")
                status_label.setStyleSheet("color: #9ad17a; padding: 5px;")
            else:
                status_label.setText("Status: ❌ Disabled")
                status_label.setStyleSheet("color: #ff6b6b; padding: 5px;")

        enable_check.stateChanged.connect(toggle_status)
        toggle_status() # Set initial text
        scroll_layout.addWidget(enable_box)

        # --- Schedule settings ---
        schedule_box = QGroupBox("Schedule")
        schedule_layout = QFormLayout(schedule_box)

        frequency_combo = NeumorphicGradientComboBox("Select Frequency")
        frequency_combo.setMinimumHeight(70)
        frequency_combo.combo_box.setFocusPolicy(Qt.ClickFocus)
        frequency_combo.combo_box.installEventFilter(wheel_guard)
        frequency_combo.addItems(["Daily", "Weekly", "Monthly"])
        # Load saved value
        frequency_combo.combo_box.setCurrentText(config.get('frequency', 'Daily'))
        schedule_layout.addRow("Frequency:", frequency_combo)

        # Time selection with dropdowns
        time_layout = QHBoxLayout()
        hour_combo = NeumorphicGradientComboBox("Hour")
        hour_combo.setMinimumHeight(70)
        hour_combo.combo_box.setFocusPolicy(Qt.ClickFocus)
        hour_combo.combo_box.installEventFilter(wheel_guard)
        hour_combo.addItems([f"{h:02d}" for h in range(24)])
        # Load saved value
        hour_combo.combo_box.setCurrentText(config.get('hour', '18'))

        minute_combo = NeumorphicGradientComboBox("Minute")
        minute_combo.setMinimumHeight(70)
        minute_combo.combo_box.setFocusPolicy(Qt.ClickFocus)
        minute_combo.combo_box.installEventFilter(wheel_guard)
        minute_combo.addItems([f"{m:02d}" for m in range(0, 60, 5)])
        # Load saved value
        minute_combo.combo_box.setCurrentText(config.get('minute', '00'))

        time_layout.addWidget(QLabel("Hour:"))
        time_layout.addWidget(hour_combo)
        time_layout.addWidget(QLabel("Minute:"))
        time_layout.addWidget(minute_combo)
        time_layout.addStretch()
        schedule_layout.addRow("Backup Time:", time_layout)

        day_combo = NeumorphicGradientComboBox("Select Day")
        day_combo.setMinimumHeight(70)
        day_combo.combo_box.setFocusPolicy(Qt.ClickFocus)
        day_combo.combo_box.installEventFilter(wheel_guard)
        day_combo.addItems(["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"])
        # Load saved value
        day_combo.combo_box.setCurrentText(config.get('day_of_week', 'Monday'))

        schedule_layout.addRow("Day of Week:", day_combo)

        def update_day_combo():
            day_combo.setEnabled(frequency_combo.combo_box.currentText() == "Weekly")

        frequency_combo.combo_box.currentTextChanged.connect(update_day_combo)
        update_day_combo() # Set initial state

        scroll_layout.addWidget(schedule_box)

        # --- Backup location ---
        location_box = QGroupBox("Backup Location")
        location_layout = QVBoxLayout(location_box)

        location_path = NeumorphicGradientLineEdit("\\\\SERVER\\Backups\\EmployeeVault")
        location_path.setMinimumHeight(70)
        # Load saved value
        location_path.line_edit.setText(config.get('location', ''))
        location_browse = ModernAnimatedButton("Browse...")

        def browse_backup_location():
            folder = QFileDialog.getExistingDirectory(dlg, "Select Backup Location")
            if folder:
                # Normalize path to use forward slashes for consistency
                location_path.line_edit.setText(folder.replace("\\", "/"))

        location_browse.clicked.connect(browse_backup_location)

        loc_layout = QHBoxLayout()
        loc_layout.addWidget(location_path)
        loc_layout.addWidget(location_browse)
        location_layout.addLayout(loc_layout)

        scroll_layout.addWidget(location_box)

        # --- Retention policy ---
        retention_box = QGroupBox("Retention Policy")
        retention_layout = QFormLayout(retention_box)

        keep_backups = NeumorphicGradientLineEdit("Number of backups to keep")
        keep_backups.setMinimumHeight(70)
        # Load saved value
        keep_backups.line_edit.setText(config.get('keep_n', '30'))
        retention_layout.addRow("Keep Last N Backups:", keep_backups)

        auto_delete = QCheckBox("Auto-delete old backups")
        # Load saved value
        auto_delete.setChecked(config.get('auto_delete', True))
        retention_layout.addRow("", auto_delete)

        scroll_layout.addWidget(retention_box)

        # --- Notifications ---
        notif_box = QGroupBox("Notifications")
        notif_layout = QVBoxLayout(notif_box)

        notify_success = QCheckBox("Notify on successful backup")
        # Load saved value
        notify_success.setChecked(config.get('notify_success', True))
        notif_layout.addWidget(notify_success)

        notify_fail = QCheckBox("Notify on backup failure")
        # Load saved value
        notify_fail.setChecked(config.get('notify_fail', True))
        notif_layout.addWidget(notify_fail)

        scroll_layout.addWidget(notif_box)

        # --- End of scrolling content ---
        scroll_layout.addStretch(1)
        scroll.setWidget(scroll_content_widget)
        layout.addWidget(scroll)


        # --- NEW: Real Save Function ---
        def _save_settings():
            # Create a dictionary with all the current values
            new_config = {
                'enabled': enable_check.isChecked(),
                'frequency': frequency_combo.combo_box.currentText(),
                'hour': hour_combo.combo_box.currentText(),
                'minute': minute_combo.combo_box.currentText(),
                'day_of_week': day_combo.combo_box.currentText(),
                'location': location_path.line_edit.text(),
                'keep_n': keep_backups.line_edit.text(),
                'auto_delete': auto_delete.isChecked(),
                'notify_success': notify_success.isChecked(),
                'notify_fail': notify_fail.isChecked(),
            }
            try:
                # Save the dictionary to the database as a JSON string
                self.db.set_setting('backup_config', json.dumps(new_config))
                # Log this action
                self.db.log_action(self.current_user, "BACKUP_CONFIG_CHANGED", details="Scheduled backup settings updated.")
                show_success_toast(dlg, "Scheduled backup configuration saved successfully!")
                dlg.close()
            except Exception as e:
                show_error_toast(dlg, f"Failed to save settings to database: {e}")

        # --- NEW: Real Test Backup Function ---
        def _run_test_backup():
            location = location_path.line_edit.text()
            if not location:
                show_warning_toast(dlg, "Please set a Backup Location first.")
                return

            if not os.path.exists(location):
                try:
                    os.makedirs(location, exist_ok=True)
                except Exception as e:
                    show_error_toast(dlg, f"Invalid backup location. Could not create directory:\n{e}")
                    return

            # Create a timestamped folder for this specific test
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_folder = os.path.join(location, f"TEST_BACKUP_{timestamp}")

            try:
                os.makedirs(backup_folder, exist_ok=True)

                # Fix: Use SQLite backup API instead of file copying to avoid WinError 33 (locked files)
                # This properly handles the database even when it's in use
                backup_path = self.db.backup_database(backup_folder)

                if backup_path:
                    show_success_toast(dlg, f"Test backup created successfully!\n\n"
                        f"Backup file:\n{backup_path}\n\n"
                        f"The database was safely backed up using SQLite's backup API,\n"
                        f"which avoids file locking issues.")

                    # Log this action
                    self.db.log_action(self.current_user, "BACKUP_TEST_RUN", details=f"Test backup created at {backup_path}")
                else:
                    show_warning_toast(dlg, "Backup operation completed but no backup file was created.\n"
                        "Check the logs for details.")

            except Exception as e:
                show_error_toast(dlg, f"An error occurred while creating backup folder:\n{e}")

        # --- Test backup button (Fixed at bottom) ---
        test_btn = ModernAnimatedButton("🧪 Test Backup Now")
        # --- UPDATED: Connect to the new function ---
        test_btn.clicked.connect(_run_test_backup)
        layout.addWidget(test_btn)

        # --- Buttons (Fixed at bottom) ---
        button_box = QHBoxLayout()
        save_btn = PulseButton("💾 Save Autobackup Settings")
        save_btn.start_pulse()
        # --- UPDATED: Connect to the new function ---
        save_btn.clicked.connect(_save_settings)

        cancel_btn = ModernAnimatedButton("Cancel")
        cancel_btn.clicked.connect(dlg.close)

        button_box.addWidget(save_btn)
        button_box.addWidget(cancel_btn)
        layout.addLayout(button_box)

        dlg.exec()

    def _show_archive_manager(self):
        """Show Archive Manager dialog (Priority #2 - Delete Protection)"""
        # v4.4.1: Animated dialog for archive manager
        from employee_vault.ui.widgets import SmoothAnimatedDialog
        dlg = SmoothAnimatedDialog(self, animation_style="slide")
        dlg.setWindowTitle("📦 Archive Manager")
        dlg.resize(1000, 600)

        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel("<h2>📦 Archive Manager - Deleted Employees</h2>"))
        layout.addWidget(QLabel("Employees that have been deleted are stored here and can be restored."))

        # Toolbar
        toolbar = QHBoxLayout()

        restore_btn = ModernAnimatedButton("↩️ Restore")
        restore_btn.clicked.connect(lambda: (
            self._restore_archived(table),
            self._load_archived(table, info_label, restore_btn, delete_btn)
        ))
        restore_btn.setEnabled(False)
        toolbar.addWidget(restore_btn)

        delete_btn = ModernAnimatedButton("🗑️ Permanently Delete")
        delete_btn.clicked.connect(lambda: (
            self._permanently_delete_archived(table),
            self._load_archived(table, info_label, restore_btn, delete_btn)
        ))
        delete_btn.setEnabled(False)
        delete_btn.setStyleSheet("background: #d32f2f;")
        toolbar.addWidget(delete_btn)

        toolbar.addStretch()



        layout.addLayout(toolbar)

        # Table
        from PySide6.QtWidgets import QTableWidget, QTableWidgetItem, QHeaderView
        table = QTableWidget()
        table.setColumnCount(9)
        table.setHorizontalHeaderLabels([
            "Employee ID", "Name", "Department", "Position",
            "Hire Date", "Archived Date", "Archived By", "Reason", "Status"
        ])
        table.horizontalHeader().setStretchLastSection(True)
        table.verticalHeader().setVisible(False)
        table.setSelectionBehavior(QTableView.SelectRows)
        table.setSelectionMode(QTableView.SingleSelection)
        table.setAlternatingRowColors(True)
        table.setEditTriggers(QTableView.NoEditTriggers)

        def on_selection():
            has_selection = len(table.selectedItems()) > 0
            restore_btn.setEnabled(has_selection)
            delete_btn.setEnabled(has_selection)

        table.itemSelectionChanged.connect(on_selection)

        layout.addWidget(table)

        # Info label
        info_label = QLabel()
        layout.addWidget(info_label)

        # Close button
        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(dlg.close)
        bottom_layout = QHBoxLayout()
        bottom_layout.addStretch()
        bottom_layout.addWidget(close_btn)
        layout.addLayout(bottom_layout)

        # Load archived employees
        self._load_archived(table, info_label, restore_btn, delete_btn)

        # Create an auto-refresh timer that is a child of the dialog
        # This ensures the timer is destroyed when the dialog is closed.
        refresh_timer = QTimer(dlg)
        refresh_timer.setInterval(5000) # 15 seconds (or your preference)

        # Connect the timer's timeout signal to the _load_archived function
        # We use a lambda to pass all the required arguments.
        refresh_timer.timeout.connect(lambda: self._load_archived(
            table, info_label, restore_btn, delete_btn
        ))

        # Start the timer
        refresh_timer.start()

        # --- END OF ADDED CODE ---

        # This line should already be here:
        dlg.exec()

    def _load_archived(self, table, info_label, restore_btn, delete_btn):
        """Load archived employees into table"""
        archived = self.db.get_archived_employees()
        table.setRowCount(len(archived))

        for i, emp in enumerate(archived):
            table.setItem(i, 0, QTableWidgetItem(emp.get('emp_id', '')))
            table.setItem(i, 1, QTableWidgetItem(emp.get('name', '')))
            table.setItem(i, 2, QTableWidgetItem(emp.get('department', '') or '—'))
            table.setItem(i, 3, QTableWidgetItem(emp.get('position', '') or '—'))
            table.setItem(i, 4, QTableWidgetItem(emp.get('hire_date', '') or '—'))
            table.setItem(i, 5, QTableWidgetItem(emp.get('archived_date', '')))
            table.setItem(i, 6, QTableWidgetItem(emp.get('archived_by', '')))
            table.setItem(i, 7, QTableWidgetItem(emp.get('archive_reason', '') or 'No reason'))

            # Status
            if emp.get('resign_date'):
                table.setItem(i, 8, QTableWidgetItem(f"Resigned - {emp.get('resign_date')}"))
            else:
                table.setItem(i, 8, QTableWidgetItem("Was Active"))

        table.resizeColumnsToContents()
        info_label.setText(f"📊 Total archived employees: <b>{len(archived)}</b>")

        restore_btn.setEnabled(False)
        delete_btn.setEnabled(False)

    def _restore_archived(self, table):
        """Restore selected archived employee"""
        selected = table.selectedItems()
        if not selected:
            return

        row = table.currentRow()
        emp_id = table.item(row, 0).text()
        name = table.item(row, 1).text()

        reply = QMessageBox.question(
            self,
            "Confirm Restore",
            f"Restore employee:\n\n"
            f"👤 {name} ({emp_id})\n\n"
            f"This will move the employee back to the active employees list.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            if self.db.restore_employee(emp_id, self.current_user):
                # Restore files
                archive_folder = os.path.join(FILES_DIR, "_archived", emp_id)
                if os.path.isdir(archive_folder):
                    dest_folder = os.path.join(FILES_DIR, emp_id)
                    try:
                        shutil.move(archive_folder, dest_folder)
                    except (OSError, shutil.Error):
                        # Folder may already exist or be locked
                        pass

                # Restore photo
                archive_photo = os.path.join(PHOTOS_DIR, "_archived", f"{emp_id}.png")
                if os.path.exists(archive_photo):
                    dest_photo = os.path.join(PHOTOS_DIR, f"{emp_id}.png")
                    try:
                        shutil.move(archive_photo, dest_photo)
                    except (OSError, shutil.Error):
                        # Photo may already exist or be locked
                        pass

                show_success_toast(self, f"Employee {name} has been restored!")
                self._refresh_all()
            else:
                show_warning_toast(self, "Failed to restore employee.")

    def _permanently_delete_archived(self, table):
        """Permanently delete selected archived employee"""
        selected = table.selectedItems()
        if not selected:
            return

        row = table.currentRow()
        emp_id = table.item(row, 0).text()
        name = table.item(row, 1).text()

        reply = show_warning_toast(
            self, f"PERMANENTLY DELETE employee:\n\n"
            f"👤 {name} ({emp_id})\n\n"
            f"⚠️ WARNING: This action CANNOT be undone!\n"
            f"The employee and all associated files will be permanently deleted.\n\n"
            f"Are you absolutely sure?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            # Second confirmation
            confirm = QInputDialog.getText(
                self,
                "Final Confirmation",
                f"Type the employee name to confirm permanent deletion:\n\n{name}"
            )

            if confirm[1] and confirm[0] == name:
                if self.db.permanently_delete_archived(emp_id, self.current_user):
                    # Delete archived files permanently
                    archive_folder = os.path.join(FILES_DIR, "_archived", emp_id)
                    if os.path.isdir(archive_folder):
                        shutil.rmtree(archive_folder, ignore_errors=True)

                    # Delete archived photo
                    archive_photo = os.path.join(PHOTOS_DIR, "_archived", f"{emp_id}.png")
                    if os.path.exists(archive_photo):
                        os.remove(archive_photo)

                    show_success_toast(self, f"Employee {name} has been permanently deleted.")
                else:
                    show_warning_toast(self, "Failed to delete employee.")
            else:
                show_success_toast(self, "Deletion cancelled - name did not match.")

    def _show_user_management(self):
        """Show user management dialog"""
        dlg = UserManagementDialog(self.db, self.current_user, self)
        dlg.exec()

    def _show_audit_log(self):
        """Show audit log viewer dialog"""
        # v4.4.1: Animated dialog for audit log
        from employee_vault.ui.widgets import SmoothAnimatedDialog
        dlg = SmoothAnimatedDialog(self, animation_style="slide")
        dlg.setWindowTitle("📜 Audit Log")
        dlg.resize(1000, 600)

        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel("<h2>📜 Audit Log - Activity History</h2>"))

        # Filters
        filter_box = QGroupBox("Filters")
        filter_layout = QHBoxLayout(filter_box)

        # Username filter
        filter_layout.addWidget(QLabel("Username:"))
        username_filter = NeumorphicGradientComboBox("All Users")
        username_filter.setMinimumHeight(70)
        username_filter.addItem("All Users")
        # Get unique usernames from audit log
        users = set()
        for log in self.db.get_audit_log(limit=1000):
            users.add(log['username'])
        for user in sorted(users):
            username_filter.addItem(user)
        filter_layout.addWidget(username_filter)

        # Action filter
        filter_layout.addWidget(QLabel("Action:"))
        action_filter = NeumorphicGradientComboBox("All Actions")
        action_filter.setMinimumHeight(70)
        action_filter.addItems(["All Actions", "LOGIN", "LOGOUT", "INSERT", "UPDATE", "DELETE"])
        filter_layout.addWidget(action_filter)

        # Record ID filter
        filter_layout.addWidget(QLabel("Employee ID:"))
        record_filter = NeumorphicGradientLineEdit("Filter by Employee ID")
        record_filter.setMinimumHeight(70)
        filter_layout.addWidget(record_filter)

        # Refresh button
        refresh_btn = ModernAnimatedButton("🔄 Refresh")
        filter_layout.addWidget(refresh_btn)

        layout.addWidget(filter_box)

        # Table for audit log
        from PySide6.QtWidgets import QTableWidget, QTableWidgetItem, QHeaderView
        table = QTableWidget()
        table.setColumnCount(8)
        table.setHorizontalHeaderLabels(["ID", "Timestamp", "Username", "Action", "Table", "Record ID", "Old Value", "New Value", "Details"])
        table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        table.horizontalHeader().setStretchLastSection(True)
        table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        table.setSelectionBehavior(QAbstractItemView.SelectRows)

        def load_logs():
            """Load logs with current filters"""
            username = None if username_filter.combo_box.currentText() == "All Users" else username_filter.combo_box.currentText()
            action = None if action_filter.combo_box.currentText() == "All Actions" else action_filter.combo_box.currentText()
            record_id = record_filter.line_edit.text().strip() or None

            logs = self.db.get_audit_log(limit=500, username=username, action=action, record_id=record_id)

            table.setRowCount(len(logs))
            for i, log in enumerate(logs):
                table.setItem(i, 0, QTableWidgetItem(str(log.get('id', ''))))
                table.setItem(i, 1, QTableWidgetItem(log.get('timestamp', '')))
                table.setItem(i, 2, QTableWidgetItem(log.get('username', '')))

                # Color code actions
                action_item = QTableWidgetItem(log.get('action', ''))
                action_text = log.get('action', '')
                if action_text == 'LOGIN':
                    action_item.setForeground(QColor('#4CAF50'))  # Green
                elif action_text == 'LOGOUT':
                    action_item.setForeground(QColor('#9E9E9E'))  # Gray
                elif action_text == 'INSERT':
                    action_item.setForeground(QColor('#2196F3'))  # Blue
                elif action_text == 'UPDATE':
                    action_item.setForeground(QColor('#FF9800'))  # Orange
                elif action_text == 'DELETE':
                    action_item.setForeground(QColor('#F44336'))  # Red
                table.setItem(i, 3, action_item)

                table.setItem(i, 4, QTableWidgetItem(log.get('table_name', '') or ''))
                table.setItem(i, 5, QTableWidgetItem(log.get('record_id', '') or ''))
                table.setItem(i, 6, QTableWidgetItem(log.get('old_value', '') or ''))
                table.setItem(i, 7, QTableWidgetItem(log.get('new_value', '') or ''))

                # Show details in a separate column
                details_text = log.get('details', '') or ''
                details_item = QTableWidgetItem(details_text[:100] + ('...' if len(details_text) > 100 else ''))
                details_item.setToolTip(details_text)  # Full text on hover
                table.setItem(i, 8, details_item)

            # Adjust column widths
            table.resizeColumnsToContents()
            table.setColumnWidth(8, 200)  # Details column wider

        # Initial load
        load_logs()

        # Connect refresh button
        refresh_btn.clicked.connect(load_logs)
        username_filter.currentTextChanged.connect(load_logs)
        action_filter.currentTextChanged.connect(load_logs)

        layout.addWidget(table)

        # Stats
        stats_label = QLabel()
        def update_stats():
            stats_label.setText(f"Showing {table.rowCount()} entries")
        update_stats()
        layout.addWidget(stats_label)

        # Export button
        export_btn = ModernAnimatedButton("📥 Export to JSON")
        def export_logs():
            filename, _ = QFileDialog.getSaveFileName(dlg, "Save Audit Log", "audit_log.json", "JSON Files (*.json)")
            if filename:
                try:
                    logs = self.db.get_audit_log(limit=5000)
                    with open(filename, 'w') as f:
                        json.dump(logs, f, indent=2)
                    show_success_toast(dlg, f"Audit log exported to {filename}")
                except Exception as e:
                    show_error_toast(dlg, f"Failed to export: {str(e)}")
        export_btn.clicked.connect(export_logs)
        layout.addWidget(export_btn)

        # Close button
        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(dlg.close)
        layout.addWidget(close_btn)

        dlg.exec()

    def _show_security_audit(self):
        """Show security audit log with tamper detection verification"""
        from employee_vault.ui.widgets import SmoothAnimatedDialog
        dlg = SmoothAnimatedDialog(self, animation_style="slide")
        dlg.setWindowTitle("🔒 Security Audit Log")
        dlg.resize(1100, 700)

        layout = QVBoxLayout(dlg)
        
        # Header with integrity status
        header_layout = QHBoxLayout()
        header_layout.addWidget(QLabel("<h2>🔒 Security Audit - Tamper-Proof Log</h2>"))
        
        # Integrity verification button and status
        integrity_status = QLabel("⏳ Checking integrity...")
        integrity_status.setStyleSheet("color: #888; font-size: 12px;")
        header_layout.addStretch()
        header_layout.addWidget(integrity_status)
        
        verify_btn = ModernAnimatedButton("🔍 Verify Integrity")
        header_layout.addWidget(verify_btn)
        
        layout.addLayout(header_layout)

        # Filters
        filter_box = QGroupBox("Filters")
        filter_layout = QHBoxLayout(filter_box)

        # Event type filter
        filter_layout.addWidget(QLabel("Event Type:"))
        event_filter = NeumorphicGradientComboBox("All Events")
        event_filter.setMinimumHeight(70)
        event_filter.addItems([
            "All Events", "LOGIN_SUCCESS", "LOGIN_FAILED", "PIN_AUTO_RESET",
            "PASSWORD_CHANGE", "USER_CREATED", "USER_DELETED", "PERMISSION_CHANGE",
            "BACKUP_CREATED", "DATA_EXPORT"
        ])
        filter_layout.addWidget(event_filter)

        # Severity filter
        filter_layout.addWidget(QLabel("Severity:"))
        severity_filter = NeumorphicGradientComboBox("All")
        severity_filter.setMinimumHeight(70)
        severity_filter.addItems(["All", "INFO", "WARNING", "ERROR", "CRITICAL"])
        filter_layout.addWidget(severity_filter)

        # Username filter
        filter_layout.addWidget(QLabel("Username:"))
        username_filter = NeumorphicGradientLineEdit("Filter by username")
        username_filter.setMinimumHeight(70)
        filter_layout.addWidget(username_filter)

        # Date range
        filter_layout.addWidget(QLabel("From:"))
        date_from = DatePicker()
        date_from.setDate(QDate.currentDate().addDays(-30))
        filter_layout.addWidget(date_from)

        filter_layout.addWidget(QLabel("To:"))
        date_to = DatePicker()
        date_to.setDate(QDate.currentDate())
        filter_layout.addWidget(date_to)

        # Refresh button
        refresh_btn = ModernAnimatedButton("🔄 Refresh")
        filter_layout.addWidget(refresh_btn)

        layout.addWidget(filter_box)

        # Table for security audit
        from PySide6.QtWidgets import QTableWidget, QTableWidgetItem, QHeaderView
        table = QTableWidget()
        table.setColumnCount(8)
        table.setHorizontalHeaderLabels([
            "ID", "Timestamp", "Event Type", "Username", "Severity",
            "Computer", "Details", "Hash (truncated)"
        ])
        table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        table.horizontalHeader().setStretchLastSection(True)
        table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        table.setSelectionBehavior(QAbstractItemView.SelectRows)
        table.setAlternatingRowColors(True)

        def load_logs():
            """Load security logs with current filters"""
            event_type = None if event_filter.combo_box.currentText() == "All Events" else event_filter.combo_box.currentText()
            severity = None if severity_filter.combo_box.currentText() == "All" else severity_filter.combo_box.currentText()
            username = username_filter.line_edit.text().strip() or None
            start_date = date_from.date().toString("yyyy-MM-dd")
            end_date = date_to.date().toString("yyyy-MM-dd")

            logs = self.db.get_security_audit(
                limit=500,
                event_type=event_type,
                username=username,
                severity=severity,
                start_date=start_date,
                end_date=end_date
            )

            table.setRowCount(len(logs))
            for i, log in enumerate(logs):
                table.setItem(i, 0, QTableWidgetItem(str(log.get('id', ''))))
                table.setItem(i, 1, QTableWidgetItem(log.get('timestamp', '')[:19]))  # Trim microseconds
                
                # Color code event types
                event_item = QTableWidgetItem(log.get('event_type', ''))
                event_text = log.get('event_type', '')
                if 'SUCCESS' in event_text:
                    event_item.setForeground(QColor('#4CAF50'))  # Green
                elif 'FAILED' in event_text:
                    event_item.setForeground(QColor('#F44336'))  # Red
                elif 'RESET' in event_text:
                    event_item.setForeground(QColor('#FF9800'))  # Orange
                elif 'CHANGE' in event_text:
                    event_item.setForeground(QColor('#2196F3'))  # Blue
                table.setItem(i, 2, event_item)
                
                table.setItem(i, 3, QTableWidgetItem(log.get('username', '') or '—'))
                
                # Color code severity
                severity_item = QTableWidgetItem(log.get('severity', 'INFO'))
                severity_text = log.get('severity', 'INFO')
                if severity_text == 'CRITICAL':
                    severity_item.setForeground(QColor('#F44336'))
                    severity_item.setBackground(QColor('#F44336', 30))
                elif severity_text == 'ERROR':
                    severity_item.setForeground(QColor('#FF5722'))
                elif severity_text == 'WARNING':
                    severity_item.setForeground(QColor('#FF9800'))
                else:
                    severity_item.setForeground(QColor('#4CAF50'))
                table.setItem(i, 4, severity_item)
                
                table.setItem(i, 5, QTableWidgetItem(log.get('computer_name', '') or '—'))
                
                # Details with tooltip
                details_text = log.get('details', '') or ''
                details_item = QTableWidgetItem(details_text[:80] + ('...' if len(details_text) > 80 else ''))
                details_item.setToolTip(details_text)
                table.setItem(i, 6, details_item)
                
                # Truncated hash for display
                entry_hash = log.get('entry_hash', '')
                hash_item = QTableWidgetItem(entry_hash[:16] + '...' if entry_hash else '—')
                hash_item.setToolTip(f"Full hash: {entry_hash}")
                hash_item.setForeground(QColor('#888'))
                table.setItem(i, 7, hash_item)

            table.resizeColumnsToContents()
            update_stats()

        def verify_integrity():
            """Verify hash chain integrity"""
            integrity_status.setText("⏳ Verifying...")
            integrity_status.setStyleSheet("color: #888;")
            QApplication.processEvents()
            
            result = self.db.verify_security_audit_integrity()
            
            if result.get('valid'):
                integrity_status.setText(f"✅ Integrity OK ({result.get('total_entries', 0)} entries)")
                integrity_status.setStyleSheet("color: #4CAF50; font-weight: bold;")
                show_success_toast(dlg, "Security audit log integrity verified - no tampering detected")
            else:
                integrity_status.setText("❌ INTEGRITY FAILURE!")
                integrity_status.setStyleSheet("color: #F44336; font-weight: bold;")
                details = result.get('details', 'Unknown error')
                if isinstance(details, list):
                    details = '\n'.join(details)
                show_error_toast(dlg, f"Security audit integrity check FAILED!\n{details}")

        # Initial load
        load_logs()
        
        # Verify integrity on open
        QTimer.singleShot(500, verify_integrity)

        # Connect filters
        refresh_btn.clicked.connect(load_logs)
        verify_btn.clicked.connect(verify_integrity)
        event_filter.currentTextChanged.connect(load_logs)
        severity_filter.currentTextChanged.connect(load_logs)

        layout.addWidget(table)

        # Stats
        stats_label = QLabel()
        def update_stats():
            stats_label.setText(f"Showing {table.rowCount()} security events")
        update_stats()
        layout.addWidget(stats_label)

        # Button row
        btn_layout = QHBoxLayout()
        
        # Export button
        export_btn = ModernAnimatedButton("📥 Export to JSON")
        def export_logs():
            filename, _ = QFileDialog.getSaveFileName(
                dlg, "Save Security Audit", "security_audit.json", "JSON Files (*.json)"
            )
            if filename:
                try:
                    logs = self.db.get_security_audit(limit=10000)
                    integrity = self.db.verify_security_audit_integrity()
                    export_data = {
                        'exported_at': datetime.now().isoformat(),
                        'integrity_verified': integrity.get('valid', False),
                        'total_entries': len(logs),
                        'entries': logs
                    }
                    with open(filename, 'w') as f:
                        json.dump(export_data, f, indent=2)
                    show_success_toast(dlg, f"Security audit exported to {filename}")
                except Exception as e:
                    show_error_toast(dlg, f"Failed to export: {str(e)}")
        export_btn.clicked.connect(export_logs)
        btn_layout.addWidget(export_btn)
        
        btn_layout.addStretch()

        # Close button
        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(dlg.close)
        btn_layout.addWidget(close_btn)
        
        layout.addLayout(btn_layout)

        dlg.exec()

    def _show_reports(self):
        """Show reports dialog"""
        # v4.4.1: Animated dialog for reports
        from employee_vault.ui.widgets import AnimatedDialogBase
        dlg = AnimatedDialogBase(self, animation_style="fade")
        dlg.setWindowTitle("📊 Reports")
        dlg.resize(500, 400)

        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel("<h2>📊 Generate Reports</h2>"))

        # Report options
        reports = QGroupBox("Available Reports")
        reports_layout = QVBoxLayout(reports)

        btn_contract_expiry = ModernAnimatedButton("📄 Contract Expiry Report")
        btn_contract_expiry.clicked.connect(lambda: self._generate_contract_report(dlg))
        reports_layout.addWidget(btn_contract_expiry)

        btn_department_summary = ModernAnimatedButton("🏢 Department Summary")
        btn_department_summary.clicked.connect(lambda: self._generate_department_report(dlg))
        reports_layout.addWidget(btn_department_summary)

        btn_employee_list = ModernAnimatedButton("👥 Complete Employee List")
        btn_employee_list.clicked.connect(lambda: self._generate_employee_list_report(dlg))
        reports_layout.addWidget(btn_employee_list)

        btn_agency_report = ModernAnimatedButton("🏢 Agency Report")
        btn_agency_report.clicked.connect(lambda: self._generate_agency_report(dlg))
        reports_layout.addWidget(btn_agency_report)

        # NEW: Advanced Reports with Charts & Visualizations
        btn_advanced_reports = ModernAnimatedButton("📊 Advanced Analytics & Charts")
        btn_advanced_reports.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #9C27B0, stop:1 #BA68C8);
                color: white;
                font-weight: bold;
                border: none;
                border-radius: 6px;
                padding: 12px 16px;
                min-height: 40px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #BA68C8, stop:1 #9C27B0);
            }
        """)
        btn_advanced_reports.clicked.connect(lambda: self._show_advanced_reports(dlg))
        reports_layout.addWidget(btn_advanced_reports)

        # Salary analysis removed per user request

        layout.addWidget(reports)

        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(dlg.close)
        layout.addWidget(close_btn)

        dlg.exec()

    def _generate_contract_report(self, parent):
        """Generate contract expiry report"""
        expired = [e for e in self.employees if (contract_days_left(e) or 999999) < 0]
        expiring_soon = [e for e in self.employees if 0 <= (contract_days_left(e) or 999999) <= ALERT_DAYS]
        valid = [e for e in self.employees if (contract_days_left(e) or 999999) > ALERT_DAYS]

        report = f"""
<h2>📄 Contract Expiry Report</h2>
<p><b>Generated:</b> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
<hr>

<h3 style='color:#ff6b6b;'>🔴 Expired Contracts ({len(expired)})</h3>
"""
        for e in expired:
            days = contract_days_left(e)
            report += f"<p>• <b>{e['name']}</b> ({e['emp_id']}) - Expired {-days} days ago<br>"
            report += f"&nbsp;&nbsp;Department: {e.get('department', 'N/A')} | Agency: {e.get('agency', 'N/A')}</p>"

        report += f"<h3 style='color:#ffcc66;'>🟠 Expiring Soon ({len(expiring_soon)})</h3>"
        for e in expiring_soon:
            days = contract_days_left(e)
            report += f"<p>• <b>{e['name']}</b> ({e['emp_id']}) - Expires in {days} days<br>"
            report += f"&nbsp;&nbsp;Department: {e.get('department', 'N/A')} | Agency: {e.get('agency', 'N/A')}</p>"

        report += f"<h3 style='color:#9ad17a;'>🟢 Valid Contracts ({len(valid)})</h3>"
        report += f"<p>{len(valid)} employees with valid contracts (>{ALERT_DAYS} days remaining)</p>"

        self._show_report_dialog(parent, "Contract Expiry Report", report)

    def _generate_department_report(self, parent):
        """Generate department summary"""
        dept_stats = {}
        for e in self.employees:
            dept = e.get('department', 'Unassigned')
            if dept not in dept_stats:
                dept_stats[dept] = {'total': 0, 'active': 0, 'resigned': 0, 'total_salary': 0}
            dept_stats[dept]['total'] += 1
            if not e.get('resign_date'):
                dept_stats[dept]['active'] += 1
            else:
                dept_stats[dept]['resigned'] += 1
            dept_stats[dept]['total_salary'] += e.get('salary', 0)

        report = f"""
<h2>🏢 Department Summary Report</h2>
<p><b>Generated:</b> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
<hr>
"""
        for dept, stats in sorted(dept_stats.items()):
            avg_salary = stats['total_salary'] / stats['total'] if stats['total'] > 0 else 0
            report += f"""
<h3>{dept}</h3>
<p>
• <b>Total Employees:</b> {stats['total']}<br>
• <b>Active:</b> {stats['active']}<br>
• <b>Resigned:</b> {stats['resigned']}<br>
• <b>Avg Salary/Day:</b> ₱{avg_salary:,.2f}<br>
• <b>Total Daily Payroll:</b> ₱{stats['total_salary']:,.2f}
</p>
"""

        self._show_report_dialog(parent, "Department Summary", report)

    def _generate_employee_list_report(self, parent):
        """Generate complete employee list"""
        active = [e for e in self.employees if not e.get('resign_date')]
        resigned = [e for e in self.employees if e.get('resign_date')]

        report = f"""
<h2>👥 Complete Employee List</h2>
<p><b>Generated:</b> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
<p><b>Total:</b> {len(self.employees)} | <b>Active:</b> {len(active)} | <b>Resigned:</b> {len(resigned)}</p>
<hr>

<h3>Active Employees ({len(active)})</h3>
"""
        for e in sorted(active, key=lambda x: x.get('name', '')):
            report += f"""
<p><b>{e['name']}</b> ({e['emp_id']})<br>
&nbsp;&nbsp;{e.get('position', 'N/A')} - {e.get('department', 'N/A')}<br>
&nbsp;&nbsp;Hired: {e.get('hire_date', 'N/A')}</p>
"""

        self._show_report_dialog(parent, "Employee List", report)

    def _generate_agency_report(self, parent):
        """Generate agency report"""
        agency_stats = {}
        for e in self.employees:
            agency = e.get('agency', 'Direct Hire')
            if agency not in agency_stats:
                agency_stats[agency] = []
            agency_stats[agency].append(e)

        report = f"""
<h2>🏢 Agency Report</h2>
<p><b>Generated:</b> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
<hr>
"""
        for agency, emps in sorted(agency_stats.items(), key=lambda x: x[0] or 'ZZZ_No_Agency'):
            active = sum(1 for e in emps if not e.get('resign_date'))
            report += f"""
<h3>{agency}</h3>
<p><b>Total:</b> {len(emps)} | <b>Active:</b> {active}</p>
<ul>
"""
            for e in emps:
                status = "✅" if not e.get('resign_date') else "❌"
                report += f"<li>{status} {e['name']} - {e.get('position', 'N/A')}</li>"
            report += "</ul>"

        self._show_report_dialog(parent, "Agency Report", report)

    def _generate_salary_report(self, parent):
        """Generate salary analysis"""
        salaries = [e.get('salary', 0) for e in self.employees if not e.get('resign_date')]
        if not salaries:
            show_success_toast(parent, "No salary data available.")
            return

        total = sum(salaries)
        avg = total / len(salaries)
        min_sal = min(salaries)
        max_sal = max(salaries)

        report = f"""
<h2>💰 Salary Analysis</h2>
<p><b>Generated:</b> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
<p><b>Based on:</b> {len(salaries)} active employees</p>
<hr>

<h3>Summary Statistics</h3>
<p>
• <b>Total Daily Payroll:</b> ₱{total:,.2f}<br>
• <b>Average Salary/Day:</b> ₱{avg:,.2f}<br>
• <b>Lowest Salary:</b> ₱{min_sal:,.2f}<br>
• <b>Highest Salary:</b> ₱{max_sal:,.2f}<br>
• <b>Monthly Estimate:</b> ₱{total * 30:,.2f}
</p>

<h3>Salary Ranges</h3>
"""
        ranges = [
            (0, 500, "Under ₱500"),
            (500, 1000, "₱500 - ₱1,000"),
            (1000, 1500, "₱1,000 - ₱1,500"),
            (1500, 2000, "₱1,500 - ₱2,000"),
            (2000, 999999, "Over ₱2,000")
        ]

        for min_r, max_r, label in ranges:
            count = sum(1 for s in salaries if min_r <= s < max_r)
            pct = (count / len(salaries) * 100) if salaries else 0
            report += f"<p>• {label}: {count} employees ({pct:.1f}%)</p>"

        self._show_report_dialog(parent, "Salary Analysis", report)

    def _show_report_dialog(self, parent, title, html_content):
        """Show report in a dialog with export option"""
        # v4.4.1: Animated dialog for report display
        from employee_vault.ui.widgets import AnimatedDialogBase
        dlg = AnimatedDialogBase(parent, animation_style="fade")
        dlg.setWindowTitle(title)
        dlg.resize(700, 600)

        layout = QVBoxLayout(dlg)

        # Report content
        text_browser = NeumorphicGradientTextEdit("Report Content", min_height=500)
        text_browser.setMinimumHeight(520)
        text_browser.text_edit.setHtml(html_content)
        text_browser.text_edit.setReadOnly(True)
        layout.addWidget(text_browser)

        # Buttons
        btn_layout = QHBoxLayout()

        export_btn = ModernAnimatedButton("💾 Export as HTML")
        export_btn.clicked.connect(lambda: self._export_report_html(html_content, title))
        btn_layout.addWidget(export_btn)

        print_btn = ModernAnimatedButton("🖨️ Print")
        print_btn.clicked.connect(lambda: text_browser.text_edit.document().print_(QPrinter()))
        btn_layout.addWidget(print_btn)

        close_btn = ModernAnimatedButton("Close")
        close_btn.clicked.connect(dlg.close)
        btn_layout.addWidget(close_btn)

        layout.addLayout(btn_layout)
        dlg.exec()

    def _export_report_html(self, content, title):
        """Export report as HTML file"""
        filename = title.replace(" ", "_").lower() + ".html"
        path, _ = QFileDialog.getSaveFileName(self, "Export Report", filename, "HTML (*.html)")
        if not path:
            return

        html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {
            margin: 0;
            font-family: Arial, sans-serif;
            font-size: 10px;
        }
        /* Defines the size of a single ID card (standard CR80 size) */
        .card-container {
            width: 3.375in;
            height: 2.125in;
            position: relative;
            overflow: hidden;
            float: left;
            margin: 10px;
            border: 1px solid #ccc;
        }
        .card-front, .card-back {
            width: 100%;
            height: 100%;
            position: absolute;
            top: 0;
            left: 0;
        }
        .bg-image {
            width: 100%;
            height: 100%;
            position: absolute;
            top: 0;
            left: 0;
            z-index: 1;
        }
        .content {
            position: relative;
            z-index: 2;
            padding: 0.15in;
        }

        /* ----- FRONT OF CARD ----- */
        .front-logo {
            width: 0.5in;
            position: absolute;
            top: 0.1in;
            left: 0.15in;
        }
        .front-company-info {
            position: absolute;
            top: 0.15in;
            left: 0.7in;
            font-size: 6px;
            line-height: 1.3;
        }
        .employee-photo {
            position: absolute;
            top: 0.7in;
            left: 0.15in;
            width: 1in;
            height: 1in;
            border: 1px solid #333;
            background: #fff;
            object-fit: cover;
        }
        .signature-box {
            position: absolute;
            top: 1.75in;
            left: 0.15in;
            width: 1in;
            height: 0.25in;
            border: 1px solid #333;
            background: #fff;
        }
        .employee-name {
            position: absolute;
            top: 1.0in;
            left: 1.25in;
            font-size: 16px;
            font-weight: bold;
            color: #333;
        }
        .employee-position {
            position: absolute;
            top: 1.4in;
            left: 1.25in;
            font-size: 14px;
        }
        .id-number {
            position: absolute;
            bottom: 0.1in;
            left: 0.15in;
            font-size: 12px;
            font-weight: bold;
        }
        .front-apruva-logo {
            position: absolute;
            bottom: 0.4in;
            right: 0.2in;
            width: 1.2in;
        }

        /* ----- BACK OF CARD ----- */
        .card-back {
            background: #fff;
            padding: 0.15in;
            box-sizing: border-box; /* Important for padding */
        }
        .agency-info {
            background: #ffffa0; /* Light yellow */
            padding: 8px;
            border-radius: 5px;
            font-size: 10px;
            font-weight: bold;
            line-height: 1.4;
        }
        .contact-info {
            margin-top: 0.1in;
            font-size: 9px;
            line-height: 1.5;
        }
        .contact-info img {
            width: 12px;
            height: 12px;
            vertical-align: middle;
            margin-right: 5px;
        }
        .terms-box {
            background: #ffe0e0; /* Light pink */
            padding: 8px;
            font-size: 7px;
            line-height: 1.4;
            text-align: justify;
            margin-top: 0.1in;
        }
        .terms-box b {
            font-size: 8px;
            color: #00008b; /* Dark Blue */
        }
        .signature-line-img {
            position: absolute;
            bottom: 0.3in;
            left: 0.15in;
            width: 1.5in;
        }
        .back-apruva-logo {
            position: absolute;
            bottom: 0.1in;
            right: 0.15in;
            width: 1.2in;
        }
    </style>
</head>
<body>

    <div class="card-container">
        <div class="card-front">
            <img src="{{BG_PATH}}" class="bg-image">

            <img src="{{LOGO_PATH}}" class="front-logo">
            <div class="front-company-info">
                <b>CUDDLY INTERNATIONAL CORPORATION</b><br>
                #650 Jesus Ext. cor. Beata, Pandacan, Manila<br>
                Tel. Nos. 588-0324 to 25 / Fax No. 588-0327
            </div>

            <img src="{{PHOTO_PATH}}" class="employee-photo">
            <div class="signature-box"></div>

            <div class="employee-name">
                {{NAME_FIRST_MIDDLE}}<br>
                {{NAME_LAST}}
            </div>
            <div class="employee-position">{{POSITION}}</div>

            <div class="id-number">ID No. {{ID_NUMBER}}</div>
            <img src="{{APRUVA_LOGO_PATH}}" class="front-apruva-logo">
        </div>
    </div>

    <div class="card-container">
        <div class="card-back">
            <div class="agency-info">
                Agency: {{AGENCY_NAME}}<br>
                {{AGENCY_ADDRESS}}
            </div>

            <div class="contact-info">
                <img src="{{PHONE_ICON_PATH}}"> {{CONTACT_NUMBER}}<br>
                <img src="{{HOUSE_ICON_PATH}}"> #18 Apo St. Marikina Village, Nangka, Marikina City
            </div>

            <div class="terms-box">
                This card is issued by CUDDLY INTERNATIONAL CORPORATION for proper identification of cardholder designated hereon and would serve for any and all legal intents and purposes.
                Any sign of tampering / erasure will render this card null and void. Cardholder shall be held liable for misuse of this card.
                This card shall be surrendered upon separation from service.
                <br><b>ID Validity: {{VALIDITY}}</b>
            </div>

            <img src="{{SIGNATURE_PATH}}" class="signature-line-img">
            <img src="{{APRUVA_LOGO_PATH}}" class="back-apruva-logo">
        </div>
    </div>

</body>
</html>
"""
        try:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(html)
            show_success_toast(self, f"Report exported to:\n{path}")
        except Exception as e:
            show_error_toast(self, f"Failed to export report:\n{str(e)}")

    def _show_advanced_reports(self, parent=None):
        """Show advanced reports dialog with charts and visualizations"""
        from employee_vault.ui.widgets import AdvancedReportsDialog

        dlg = AdvancedReportsDialog(parent or self, self.db)
        dlg.exec()

    def _toggle_theme(self): self.setStyleSheet("" if self.styleSheet() else APP_QSS)
    def _show_theme_selector(self):
        """Show theme selector dialog"""
        # v4.4.1: Animated dialog for theme selector
        from employee_vault.ui.widgets import AnimatedDialogBase
        dialog = AnimatedDialogBase(self, animation_style="fade")
        dialog.setWindowTitle("🎨 Choose Theme")
        dialog.resize(500, 400)
        
        layout = QVBoxLayout(dialog)
        
        # Title
        title = QLabel("<h2>Choose Your Theme</h2>")
        layout.addWidget(title)
        
        info = QLabel("Select a theme to change the appearance of the application.")
        info.setWordWrap(True)
        info.setStyleSheet("color: #888; margin-bottom: 10px;")
        layout.addWidget(info)
        
        # Theme list
        theme_list = QListWidget()
        
        # Get current theme
        current_theme = load_theme_preference()
        
        for theme_id, theme_data in MODERN_THEMES.items():
            item = QListWidgetItem(f"🎨 {theme_data['name']}")
            item.setData(Qt.UserRole, theme_id)
            theme_list.addItem(item)
            
            # Select current theme
            if theme_id == current_theme:
                theme_list.setCurrentItem(item)
        
        layout.addWidget(theme_list)
        
        # Preview label
        preview_label = QLabel("<i>Click a theme to preview its colors</i>")
        preview_label.setStyleSheet("color: #888; font-size: 12px; padding: 10px;")
        layout.addWidget(preview_label)
        
        # Buttons
        btn_layout = QHBoxLayout()
        
        apply_btn = ModernAnimatedButton("✓ Apply Theme")
        apply_btn.setStyleSheet("background-color: #4CAF50;")
        apply_btn.clicked.connect(dialog.accept)
        
        cancel_btn = ModernAnimatedButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        
        btn_layout.addStretch()
        btn_layout.addWidget(apply_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)
        
        # Execute dialog
        if dialog.exec() == QDialog.Accepted:
            selected_item = theme_list.currentItem()
            if selected_item:
                new_theme = selected_item.data(Qt.UserRole)
                
                # Save theme preference
                try:
                    with open(THEME_PREFERENCE_FILE, 'w') as f:
                        f.write(new_theme)
                except (IOError, OSError):
                    # Failed to save theme preference
                    pass
                
                # Ask for confirmation and restart
                reply = QMessageBox.question(
                    self,
                    "Apply Theme",
                    f"Apply {MODERN_THEMES[new_theme]['name']} theme?\n\n"
                    f"The application will restart automatically.",
                    QMessageBox.Yes | QMessageBox.No
                )
                
                if reply == QMessageBox.Yes:
                    # Restart the application
                    QTimer.singleShot(500, lambda: self._restart_application())
    
    def _restart_application(self):
        """Restart the application to apply theme"""
        import subprocess
        
        # Close database connection
        try:
            if hasattr(self, 'db'):
                self.db.conn.close()
        except (AttributeError, sqlite3.Error):
            # Database already closed or connection error
            pass
        
        # Restart using python
        python = sys.executable
        subprocess.Popen([python] + sys.argv)
        
        # Close current application
        QApplication.quit()

    def _show_about(self):
        """Show about dialog with version and system info"""
        stats = self.db.get_database_stats()
        
        # Get current theme name
        current_theme = load_theme_preference()
        theme_name = MODERN_THEMES[current_theme]['name']

        about_text = f"""
        <h2>{APP_TITLE}</h2>
        <p><b>Version:</b> {VERSION}<br>
        <b>Build Date:</b> {BUILD_DATE}<br>
        <b>Database Version:</b> {DATABASE_VERSION}</p>

        <hr>

        <h3>UI Theme</h3>
        <p><b>Current Theme:</b> {theme_name}<br>
        <b>Available Themes:</b> {len(MODERN_THEMES)}<br>
        <i>Click "Change Theme" in Settings to switch themes</i></p>

        <hr>

        <h3>System Information</h3>
        <p><b>Total Employees:</b> {stats.get('total_employees', 0)}<br>
        <b>Active Employees:</b> {stats.get('active_employees', 0)}<br>
        <b>Total Users:</b> {stats.get('total_users', 0)}<br>
        <b>Total Agencies:</b> {stats.get('total_agencies', 0)}<br>
        <b>Database Size:</b> {stats.get('db_size_mb', 0):.2f} MB</p>

        <hr>

        <p><b>Logged in as:</b> {self.current_user} ({self.user_row['role']})<br>
        <b>Security:</b> ✅ Bcrypt password hashing<br>
        <b>Database:</b> SQLite with WAL mode<br>
        <b>UI:</b> Modern with glassmorphism & animations</p>

        <p style="color: #888; font-size: 11px;">
        © 2025 Cuddly International Corporation<br>
        Employee Information Management System - Modern UI Edition
        </p>
        """

        msg = QMessageBox(self)
        msg.setWindowTitle("About")
        msg.setTextFormat(Qt.RichText)
        msg.setText(about_text)
        msg.setIcon(QMessageBox.Information)
        msg.exec()
    
    def _load_user_photo(self):
        """v4.0.1: Load user photo with perfect circular clipping (like employee form)"""
        user_photo_path = os.path.join(PHOTOS_DIR, f"user_{self.current_user}.png")
        
        # If the user hasn't uploaded a photo, fall back to a default icon.
        if not os.path.exists(user_photo_path):
            self.user_photo_label.setPixmap(QPixmap())
            self.user_photo_label.setText("👤")
            self.user_photo_label.setAlignment(Qt.AlignCenter)
            return

        pix = QPixmap(user_photo_path)
        # Determine the size based on the current label dimensions.  This
        # allows the avatar to scale automatically if the UI designer
        # changes its size.  We pick the smaller of width and height to
        # ensure a square canvas for the circular mask.
        size = min(self.user_photo_label.width(), self.user_photo_label.height())

        # If the pixmap couldn't be loaded, show the default icon
        if pix.isNull():
            self.user_photo_label.setPixmap(QPixmap())
            self.user_photo_label.setText("👤")
            self.user_photo_label.setAlignment(Qt.AlignCenter)
            return

        # Scale the image to fit entirely within the circle.  Using Qt.KeepAspectRatio
        # avoids cropping any part of the photo, instead letterboxing inside the
        # circular mask.  This solves cases where a tall portrait photo is
        # clipped at the top/bottom when using KeepAspectRatioByExpanding.
        scaled = pix.scaled(size, size, Qt.KeepAspectRatio, Qt.SmoothTransformation)

        # Create a transparent canvas to draw the final circular image on
        result = QPixmap(size, size)
        result.fill(Qt.transparent)

        painter = QPainter(result)
        painter.setRenderHint(QPainter.Antialiasing, True)
        painter.setRenderHint(QPainter.SmoothPixmapTransform, True)

        # Define a circular clip path
        path = QPainterPath()
        path.addEllipse(0, 0, size, size)
        painter.setClipPath(path)

        # Center the scaled image inside the circle by computing offsets.
        x_offset = (size - scaled.width()) // 2
        y_offset = (size - scaled.height()) // 2
        painter.drawPixmap(x_offset, y_offset, scaled)
        painter.end()

        # Set the processed pixmap and clear any text on the label
        self.user_photo_label.setPixmap(result)
        self.user_photo_label.setText("")
    
    def _toggle_sidebar(self):
        """Animate collapsing or expanding the sidebar.  When the sidebar
        is collapsed it shows only icons; when expanded it shows the full
        section labels and nested options.  The toggle button arrow and
        tooltips update accordingly."""
        # Determine the target width and arrow direction based on the
        # current state
        if self.is_sidebar_collapsed:
            # We are currently collapsed so we will expand: arrow should
            # point left (collapse), tooltip invites collapsing after
            # expansion.
            target_width = self.sidebar_expanded_width
            new_arrow = "❮"
            new_tooltip = "Collapse sidebar"
        else:
            # We are currently expanded so we will collapse: arrow should
            # point right (expand), tooltip invites expanding after
            # collapse.
            target_width = self.sidebar_collapsed_width
            new_arrow = "❯"
            new_tooltip = "Expand sidebar"

        # Create animation on the maximumWidth property.  Storing the
        # animation on the instance prevents it from being garbage
        # collected while running.
        self.sidebar_animation = QPropertyAnimation(self.sidebar, b"maximumWidth", self)
        self.sidebar_animation.setDuration(300)
        self.sidebar_animation.setStartValue(self.sidebar.width())
        self.sidebar_animation.setEndValue(target_width)
        self.sidebar_animation.setEasingCurve(QEasingCurve.InOutCubic)
        self.sidebar_animation.start()

        # Also set the fixed width immediately so that the layout
        # managers respect the new size during the animation.  Without
        # this call the sidebar may try to expand back out before the
        # animation finishes.
        self.sidebar.setFixedWidth(target_width)

        # Update collapsed state flag
        self.is_sidebar_collapsed = not self.is_sidebar_collapsed

        # Update toggle button appearance
        self.toggle_sidebar_btn.setText(new_arrow)
        self.toggle_sidebar_btn.setToolTip(new_tooltip)

        # Update logout button appearance
        if self.is_sidebar_collapsed:
            self.logout_btn.setText("🔑")
            self.logout_btn.setToolTip("Logout")
        else:
            self.logout_btn.setText("🔓 Logout")
            self.logout_btn.setToolTip("")

        # Adjust header margins based on sidebar state
        if self.is_sidebar_collapsed:
            # Use minimal margins when collapsed to fit the toggle button
            self.header_layout.setContentsMargins(2, 5, 2, 5)
        else:
            # Restore normal margins when expanded
            self.header_layout.setContentsMargins(10, 5, 10, 5)

        # When collapsing, hide the user info and adjust section labels
        # When expanding, restore them.  The sections list and icons are
        # set up during UI construction.
        # Toggle visibility of the user and company information.  When
        # collapsed we hide the logo, company name and user info so
        # that the sidebar remains a thin icon bar.  When expanded
        # these widgets are restored.
        try:
            self.user_info_container.setVisible(not self.is_sidebar_collapsed)
        except Exception:
            pass
        if hasattr(self, "logo_label") and self.logo_label is not None:
            self.logo_label.setVisible(not self.is_sidebar_collapsed)
        if hasattr(self, "company_label") and self.company_label is not None:
            self.company_label.setVisible(not self.is_sidebar_collapsed)

        # Update each collapsible section button label to icon-only or
        # full text depending on the collapsed state.  This mapping is
        # defined during sidebar construction.
        if hasattr(self, "sidebar_sections"):
            for section in self.sidebar_sections:
                icon_char = self.section_icons.get(section, "")
                if self.is_sidebar_collapsed:
                    # Remove the arrow and text; show only icon.  The
                    # toggle button's text holds the arrow and title
                    # separated by a space, e.g., "▶ MAIN".  We ignore the
                    # arrow and set the button text to just the icon.  The
                    # tooltip holds the full title for accessibility.
                    section.toggle_button.setText(icon_char)
                    section.toggle_button.setToolTip(section.full_title)
                    # Ensure content area is fully collapsed
                    section.is_collapsed = True
                    section.content_area.setMaximumHeight(0)
                    section.content_area.setVisible(False)
                else:
                    # Restore full title with icon (no arrows per user request)
                    section.toggle_button.setText(f"{icon_char} {section.full_title}")
                    section.toggle_button.setToolTip("")
        # Force a layout update to avoid artifacts
        self.sidebar.updateGeometry()
    def _upload_user_photo(self):
        """v3.9: Upload user photo - clickable"""
        try:
            fn, _ = QFileDialog.getOpenFileName(
                self,
                "📷 Select Your Photo",
                "",
                "Images (*.png *.jpg *.jpeg *.bmp)"
            )
            
            if not fn:
                return
            
            if not os.path.exists(fn):
                show_warning_toast(self, "File not found!")
                return
            
            pix = QPixmap(fn)
            if pix.isNull():
                show_warning_toast(self, "Failed to load image!")
                return
            
            file_size = os.path.getsize(fn) / (1024 * 1024)
            if file_size > 5:
                show_warning_toast(self, f"Image too large ({file_size:.1f}MB). Max 5MB.")
                return
            
            # Save as user photo
            dest_path = os.path.join(PHOTOS_DIR, f"user_{self.current_user}.png")
            pix.save(dest_path, "PNG")
            
            # Reload photo
            self._load_user_photo()
            
            show_success_toast(self, "✅ Photo uploaded successfully!")
            logging.info(f"User photo uploaded for {self.current_user}")
            
        except Exception as e:
            show_error_toast(self, f"Failed to upload photo: {e}")
            logging.error(f"User photo upload error: {e}")

    def _logout(self):
        # Log logout action
        self.db.log_action(username=self.current_user, action="LOGOUT", details=f"User logged out")
        # Hide main window
        self.hide()
        # Show login dialog
        login = LoginDialog(self.db)
        if login.exec() == QDialog.Accepted:
            # User logged in again, create new main window
            u = login.username.text().strip()
            row = self.db.get_user(u)
            # Close this window
            self.close()
            # Create and show new main window
            new_window = MainWindow(self.db, u, row)
            new_window.show()
            # Store reference to prevent garbage collection
            QApplication.instance().main_window = new_window
        else:
            # User cancelled login, close app
            self.close()
            QApplication.quit()

    def _auto_refresh(self):
        """
        Enhanced auto-refresh for real-time multi-user synchronization.

        v2.1.2 MULTI-USER IMPROVEMENTS:
        - Optimized for 5-7 concurrent users
        - Detects database changes before refreshing (performance boost)
        - Better handling of concurrent writes
        - Notifies user when data changes from other users
        """
        try:
            # Check if database has been modified by another user
            current_mtime = db_latest_mtime(DB_FILE)
            if not hasattr(self, '_last_db_mtime'):
                self._last_db_mtime = current_mtime

            # Only refresh if database actually changed (performance optimization)
            if current_mtime > self._last_db_mtime:
                self._last_db_mtime = current_mtime

                # Store current selection to restore after refresh
                current_page_index = self.pages.currentIndex()

                # Refresh data
                self._refresh_all()

                # Restore page selection
                self.pages.setCurrentIndex(current_page_index)

                logging.debug("Auto-refresh: Database updated by another user")

                # Optional: Show subtle notification
                if hasattr(self, 'statusBar'):
                    self.statusBar().showMessage("Data updated", 2000)
            else:
                logging.debug("Auto-refresh: No changes detected (skipped refresh)")

        except sqlite3.OperationalError as e:
            if "locked" in str(e).lower():
                # Database locked is normal during write operations
                # Use exponential backoff - wait longer on next attempt
                logging.debug("Database locked during refresh - will retry next cycle")
            else:
                logging.error(f"Auto-refresh database error: {e}")
        except Exception as e:
            # Unexpected errors should be logged with full traceback
            logging.error(f"Unexpected error in auto-refresh: {e}", exc_info=True)

    def _apply_user_permissions(self):
        """Apply UI restrictions based on user permissions - v2.0 FIXED"""
        perms = self.user_permissions

        # If admin, skip restrictions (admin has all permissions)
        if self.user_row['role'] == 'admin':
            logging.info("Admin user - all permissions granted")
            return

        logging.info(f"Applying permissions for user: {self.current_user}")

        # Note: Most UI elements are already conditionally shown based on role
        # This method provides fine-grained control for non-admin users

        # The sidebar sections already check user_row['role'] and self.user_permissions
        # for conditional display, so permissions are already being enforced

        # Additional runtime checks can be added here for dynamic UI updates
        # For now, permissions are enforced at creation time (in __init__)

        logging.info(f"Permissions applied successfully for {self.current_user}")

    # ============================================================================
    # WEEK 2 FEATURE #1: SESSION IDLE TIMEOUT (AUTO-LOCK)
    # ============================================================================

    def eventFilter(self, obj, event):
        """Track user activity for idle timeout"""
        # Track mouse and keyboard events as activity
        if event.type() in [QEvent.MouseMove, QEvent.MouseButtonPress,
                           QEvent.KeyPress, QEvent.Wheel]:
            self.last_activity_time = datetime.now()
        return super().eventFilter(obj, event)

    def _check_idle_timeout(self):
        """Check if user has been idle and auto-lock if needed"""
        idle_minutes = (datetime.now() - self.last_activity_time).total_seconds() / 60

        # Warning at 2 minutes before timeout
        if idle_minutes >= (self.idle_timeout_minutes - 2) and idle_minutes < self.idle_timeout_minutes:
            if not hasattr(self, '_warning_shown'):
                self._warning_shown = True
                remaining = int(self.idle_timeout_minutes - idle_minutes)
                self.statusBar().showMessage(
                    f"⚠️ Session will lock in {remaining} minute(s) due to inactivity",
                    120000  # Show for 2 minutes
                )

        # Lock session after timeout
        if idle_minutes >= self.idle_timeout_minutes:
            logging.info(f"Session timeout: User {self.current_user} idle for {idle_minutes:.1f} minutes")
            self._lock_session()

    def _lock_session(self):
        """Lock the session and require re-authentication"""
        # Stop timers
        self.idle_timer.stop()
        if hasattr(self, 'timer'):
            self.timer.stop()

        # Hide main window
        self.hide()

        # Show login dialog for re-authentication
        login = LoginDialog(self.db)
        login.setWindowTitle("Session Locked - Re-authenticate")
        login.info.setText(f"Session locked due to inactivity.\nPlease re-authenticate as: {self.current_user}")
        login.username.setText(self.current_user)
        login.username.setReadOnly(True)
        login.password.setFocus()

        if login.exec() == QDialog.Accepted:
            # Re-authenticated successfully
            self.last_activity_time = datetime.now()
            self._warning_shown = False
            self.idle_timer.start()
            if hasattr(self, 'timer'):
                self.timer.start()
            self.show()
            logging.info(f"Session unlocked: {self.current_user} re-authenticated")
        else:
            # User cancelled - logout
            self._logout()

    def _notify_contracts(self):
        expired=[e for e in self.employees if (contract_days_left(e) or 999999) < 0]
        soon=[e for e in self.employees if 0 <= (contract_days_left(e) or 999999) <= ALERT_DAYS]
        if expired: self.tray.showMessage("Contract expired", f"{len(expired)} contract(s) expired.", QSystemTrayIcon.Information, 8000)
        if soon: self.tray.showMessage("Contract expiring", f"{len(soon)} contract(s) will expire ≤ {ALERT_DAYS} days.", QSystemTrayIcon.Information, 8000)
"""
Dashboard Page
Main dashboard with statistics and overview
"""

import os
import json
import shutil
